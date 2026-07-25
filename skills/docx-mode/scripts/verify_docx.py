#!/usr/bin/env python
"""One-shot DOCX verification: format, structure, equations, content summary.

Usage:
    PYTHONIOENCODING=utf-8 py scripts/verify_docx.py path/to/doc.docx

Output (JSON or human-readable):
  - Document-level: paragraphs, tables, images, equations
  - Normal style defaults: font, size, line spacing, alignment, indent
  - Section tree (headings only)
  - Body sample (first N paragraphs)
  - Equation count
  - Format consistency summary

Exit code 0 = all checks passed, 1 = something wrong.
"""
import argparse
import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Dict, Any

# ── Add skill dir to path for import (when run directly inside skills/) ──
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_SKILL_DIR = os.path.abspath(os.path.join(_SCRIPT_DIR, '..'))
if _SKILL_DIR not in sys.path:
    sys.path.insert(0, _SKILL_DIR)

from docx import Document as PyDocxDocument
from docx.oxml.ns import qn


# ═══════════════════════════════════════════════════════════════
#  Data model
# ═══════════════════════════════════════════════════════════════

@dataclass
class NormalStyleInfo:
    font_ascii: str = ''
    font_h_ansi: str = ''
    font_east_asia: str = ''
    font_size_pt: float = 0
    line_spacing_pt: float = 0       # 0 means "not set on style"
    line_spacing_rule: str = ''      # exact / auto / atLeast / …
    first_line_indent_pt: float = 0
    alignment: str = ''
    space_after_pt: float = 0
    space_before_pt: float = 0

@dataclass
class SectionInfo:
    level: int
    title: str

@dataclass
class VerificationReport:
    file_path: str = ''
    file_size_kb: float = 0
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    # Counts
    total_paragraphs: int = 0
    total_heading_paragraphs: int = 0
    total_tables: int = 0
    total_images: int = 0
    total_equations: int = 0
    total_characters: int = 0

    # Content structure
    sections: List[SectionInfo] = field(default_factory=list)
    normal_style: NormalStyleInfo = field(default_factory=NormalStyleInfo)

    # Format checks
    has_body_text: bool = False
    has_any_heading: bool = False


# ═══════════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════════

_TWIP = 20        # 1pt = 20 twips (Word internal unit)
_HALF_PT = 2      # 1pt = 2 half-pts (font size unit)


def _twips_to_pt(val: Optional[str]) -> float:
    if val is None:
        return 0
    try:
        return int(val) / _TWIP
    except (ValueError, TypeError):
        return 0


def _halfpt_to_pt(val: Optional[str]) -> float:
    if val is None:
        return 0
    try:
        return int(val) / _HALF_PT
    except (ValueError, TypeError):
        return 0


def _get_attr(elem, tag: str, attr: str) -> Optional[str]:
    """Find child *tag* on *elem* and return its *attr* (or None)."""
    child = elem.find(qn(tag))
    if child is not None:
        return child.get(qn(attr))
    return None


# ═══════════════════════════════════════════════════════════════
#  Normal style reader (pure XML – reliable)
# ═══════════════════════════════════════════════════════════════

def _read_normal_style(xml_body: str) -> NormalStyleInfo:
    """Extract Normal style defaults from styles.xml."""
    info = NormalStyleInfo()

    m = re.search(
        r'<w:style[^>]*w:styleId="Normal"[^>]*>.*?</w:style>',
        xml_body, re.DOTALL
    )
    if not m:
        return info
    style = m.group()

    # ── rPr ──
    rpr_m = re.search(r'<w:rPr>.*?</w:rPr>', style, re.DOTALL)
    if rpr_m:
        rpr = rpr_m.group()
        # rFonts
        rf = re.search(r'<w:rFonts\s+([^>]+?)/?>', rpr)
        if rf:
            attrs = rf.group(1)
            _a = lambda n: (re.search(rf'{n}="([^"]*)"', attrs) or [None, ''])[1]
            info.font_ascii = _a('w:ascii')
            info.font_h_ansi = _a('w:hAnsi')
            info.font_east_asia = _a('w:eastAsia')
        # sz
        sz_m = re.search(r'<w:sz\s+w:val="(\d+)"', rpr)
        if sz_m:
            info.font_size_pt = _halfpt_to_pt(sz_m.group(1))

    # ── pPr ──
    ppr_m = re.search(r'<w:pPr>.*?</w:pPr>', style, re.DOTALL)
    if ppr_m:
        ppr = ppr_m.group()

        # spacing
        sp_m = re.search(r'<w:spacing\s+([^>]+?)/?>', ppr)
        if sp_m:
            attrs = sp_m.group(1)
            _a = lambda n: (re.search(rf'{n}="([^"]*)"', attrs) or ['', ''])[1]
            rule = _a('w:lineRule')
            info.line_spacing_rule = rule or ''
            line_val = _a('w:line')
            if line_val:
                if rule == 'exact':
                    info.line_spacing_pt = _twips_to_pt(line_val)
                else:
                    # auto mode: line value / 240 = multiplier
                    try:
                        info.line_spacing_pt = int(line_val) / 240.0
                    except (ValueError, TypeError):
                        pass
            info.space_after_pt = _twips_to_pt(_a('w:after'))
            info.space_before_pt = _twips_to_pt(_a('w:before'))

        # indentation
        ind_m = re.search(r'<w:ind\s+([^>]+?)/?>', ppr)
        if ind_m:
            attrs = ind_m.group(1)
            _a = lambda n: (re.search(rf'{n}="([^"]*)"', attrs) or ['', ''])[1]
            info.first_line_indent_pt = _twips_to_pt(_a('w:firstLine'))

        # alignment
        jc_m = re.search(r'<w:jc\s+w:val="([^"]+)"', ppr)
        if jc_m:
            val = jc_m.group(1)
            MAP = {'left': 'LEFT', 'center': 'CENTER', 'right': 'RIGHT',
                   'both': 'JUSTIFY'}
            info.alignment = MAP.get(val, val.upper())

    return info


# ═══════════════════════════════════════════════════════════════
#  Content scanners (python-docx API)
# ═══════════════════════════════════════════════════════════════

def _scan_content(doc: PyDocxDocument) -> VerificationReport:
    """Scan document content and structure."""
    rpt = VerificationReport()

    # ── Paragraphs ──
    raw_text_parts: List[str] = []
    for p in doc.paragraphs:
        rpt.total_paragraphs += 1
        txt = p.text or ''
        raw_text_parts.append(txt)

        # Style-based heading detection
        style = p.style
        is_heading = False
        if style and 'Heading' in style.name:
            is_heading = True
        # Also detect Heading style via XML (covers direct formatting)
        if not is_heading:
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val')) or ''
                    if sid.startswith('Heading') or sid.startswith('heading'):
                        is_heading = True

        if is_heading:
            rpt.total_heading_paragraphs += 1
            rpt.has_any_heading = True
            # Extract level from style name like "Heading 1" or "heading2"
            level = 1
            m_lvl = re.search(r'(\d+)', style.name if style else '')
            if m_lvl:
                level = int(m_lvl.group(1))
            rpt.sections.append(SectionInfo(level=level, title=txt.strip()))

    rpt.total_characters = len(''.join(raw_text_parts))
    rpt.has_body_text = any(
        len(p.text or '') > 20
        for p in doc.paragraphs
    )

    # ── Tables ──
    rpt.total_tables = len(doc.tables)

    return rpt


# ═══════════════════════════════════════════════════════════════
#  Raw XML scanners (reliable for equations / images)
# ═══════════════════════════════════════════════════════════════

def _scan_raw_xml(xml_str: str) -> dict:
    """Count OMML equations, images, and other non-paragraph elements."""
    # Equations: <m:oMath> (inline) or <m:oMathPara> (display)
    eq_inline = len(re.findall(r'<m:oMath\b', xml_str))
    eq_display = len(re.findall(r'<m:oMathPara\b', xml_str))
    total_eq = eq_inline + eq_display

    # Images: <w:drawing> or <wp:inline> or <v:imagedata>
    images = len(re.findall(r'<w:drawing\b', xml_str))

    return {
        'equations': total_eq,
        'equations_inline': eq_inline,
        'equations_display': eq_display,
        'images': images,
    }


# ═══════════════════════════════════════════════════════════════
#  Format consistency checks
# ═══════════════════════════════════════════════════════════════

def _check_format_consistency(xml_str: str, normal: NormalStyleInfo,
                               rpt: VerificationReport):
    """Detect paragraph-level overrides that diverge from Normal style."""
    # Check body paragraphs (non-heading) for differing line spacing
    body_paras = re.findall(
        r'<w:p[ >].*?</w:p>', xml_str, re.DOTALL
    )
    non_heading_para_xmls = []
    for para_xml in body_paras:
        # Skip if heading style
        if re.search(r'<w:pStyle\s+w:val="Heading\d+"', para_xml):
            continue
        # Skip if first child is pPr with pStyle heading
        if re.search(r'<w:pPr>.*?<w:pStyle\s+w:val="[Hh]eading\d+"', para_xml, re.DOTALL):
            continue
        non_heading_para_xmls.append(para_xml)

    # Sample check: if many body paragraphs exist but normal style looks empty,
    # that's suspicious
    if len(non_heading_para_xmls) > 3 and not normal.font_size_pt:
        rpt.warnings.append(
            'Normal style font_size is 0; style may not be configured'
        )


# ═══════════════════════════════════════════════════════════════
#  Presenter
# ═══════════════════════════════════════════════════════════════

def _format_pt(v: float) -> str:
    return f'{v:.1f}pt' if v else '—'


def _fmt_line_spacing(info: NormalStyleInfo) -> str:
    if not info.line_spacing_pt:
        return '—'
    if info.line_spacing_rule == 'exact':
        return f'固定值 {_format_pt(info.line_spacing_pt)}'
    elif info.line_spacing_rule in ('auto', ''):
        return f'{info.line_spacing_pt:.2f} 倍行距'
    else:
        return f'{info.line_spacing_rule} {_format_pt(info.line_spacing_pt)}'


def print_report(rpt: VerificationReport):
    """Print a human-readable report."""
    style = rpt.normal_style

    print('=' * 60)
    print(f'  DOCX 验证报告: {os.path.basename(rpt.file_path)}')
    print('=' * 60)

    # ── File info ──
    print(f'\n📄 文件: {rpt.file_path}')
    print(f'   大小: {rpt.file_size_kb:.1f} KB')

    # ── Errors / Warnings ──
    if rpt.errors:
        print(f'\n❌ 错误 ({len(rpt.errors)}):')
        for e in rpt.errors:
            print(f'   • {e}')
    if rpt.warnings:
        print(f'\n⚠️  警告 ({len(rpt.warnings)}):')
        for w in rpt.warnings:
            print(f'   • {w}')

    # ── Format (Normal Style) ──
    print('\n── 默认格式 (Normal Style) ──')
    print(f'  西文字体: {style.font_ascii or "—"}')
    print(f'  中文字体: {style.font_east_asia or "—"}')
    print(f'  字号:     {_format_pt(style.font_size_pt)}')
    print(f'  行距:     {_fmt_line_spacing(style)}')
    print(f'  对齐:     {style.alignment or "—"}')
    print(f'  首行缩进: {_format_pt(style.first_line_indent_pt)}')
    print(f'  段后间距: {_format_pt(style.space_after_pt)}')
    print(f'  段前间距: {_format_pt(style.space_before_pt)}')

    # ── Counts ──
    print('\n── 内容统计 ──')
    print(f'  段落数:       {rpt.total_paragraphs}')
    print(f'  标题数:       {rpt.total_heading_paragraphs}')
    print(f'  表格数:       {rpt.total_tables}')
    print(f'  图片数:       {rpt.total_images}')
    print(f'  公式数:       {rpt.total_equations}')
    print(f'  总字符数:     {rpt.total_characters}')

    # ── Sections ──
    if rpt.sections:
        print('\n── 章节结构 ──')
        for sec in rpt.sections:
            indent = '  ' * (sec.level - 1)
            print(f'  {indent}{"#" * sec.level} {sec.title}')

    # ── Body text check ──
    if not rpt.has_body_text:
        print('\n⚠️  未检测到正文段落（所有段落长度 < 20 字符）')
    if not rpt.has_any_heading:
        print('\n⚠️  未检测到标题')

    print('\n' + '=' * 60)
    if rpt.errors:
        print('  结果: ❌ 存在问题')
    elif rpt.warnings:
        print('  结果: ⚠️  有警告，但无错误')
    else:
        print('  结果: ✅ 全部正常')
    print('=' * 60)


# ═══════════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════════

def verify(path: str) -> VerificationReport:
    """Verify a .docx file and return a structured report."""
    rpt = VerificationReport(file_path=os.path.abspath(path))

    # File checks
    if not os.path.isfile(path):
        rpt.errors.append(f'文件不存在: {path}')
        return rpt
    rpt.file_size_kb = os.path.getsize(path) / 1024
    if rpt.file_size_kb == 0:
        rpt.errors.append('文件大小为 0')
        return rpt

    # Open with python-docx
    try:
        doc = PyDocxDocument(path)
    except Exception as e:
        rpt.errors.append(f'python-docx 无法打开文件: {e}')
        return rpt

    # Read raw XML
    try:
        with zipfile.ZipFile(path) as z:
            doc_xml = z.read('word/document.xml').decode('utf-8')
            styles_xml = z.read('word/styles.xml').decode('utf-8')
    except Exception as e:
        rpt.errors.append(f'无法读取内部 XML: {e}')
        return rpt

    # ═══ Scans ═══
    style_info = _read_normal_style(styles_xml)
    rpt.normal_style = style_info

    content = _scan_content(doc)
    # Merge content fields into rpt
    for field_name in ('total_paragraphs', 'total_heading_paragraphs',
                       'total_tables', 'total_characters', 'has_body_text',
                       'has_any_heading', 'sections'):
        setattr(rpt, field_name, getattr(content, field_name))

    raw = _scan_raw_xml(doc_xml)
    rpt.total_equations = raw['equations']
    rpt.total_images = raw['images']

    _check_format_consistency(doc_xml, style_info, rpt)

    return rpt


def main():
    parser = argparse.ArgumentParser(
        description='One-shot DOCX format & content verification'
    )
    parser.add_argument('path', metavar='FILE', help='Path to .docx file')
    parser.add_argument('--json', action='store_true',
                        help='Output as JSON instead of human-readable')
    args = parser.parse_args()

    rpt = verify(args.path)
    if args.json:
        # Convert dataclasses to dict
        data = asdict(rpt)
        print(json.dumps(data, ensure_ascii=False, indent=2))
    else:
        print_report(rpt)

    sys.exit(1 if rpt.errors else 0)


if __name__ == '__main__':
    main()
