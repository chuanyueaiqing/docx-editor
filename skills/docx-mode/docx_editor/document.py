"""Main DocxDocument facade class.

Provides the single public entry point for all DOCX operations.
Delegates to specialized sub-modules for specific concerns.
"""
import logging
import os
import re
import tempfile
from collections import Counter
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document as DocxDocumentLoader
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as docx_qn
from docx.shared import Emu, Pt

from .chapter_parser import ChapterParser
from .format_extractor import FormatExtractor
from .markdown_processor import MarkdownProcessor
from .mermaid_renderer import MermaidRenderer, MermaidNotAvailableError
from .models import (
    ChapterNode, CommentData, FormatStore, MarkdownElement,
    MarkdownElementType, ParagraphData, RevisionData,
)
from .utils import (
    ChapterNotFoundError, CommentsReadError, DocxError, Win32ComError,
)
from .win32_ops import Win32Ops
from .wps_ops import WpsComError, WpsOps

logger = logging.getLogger(__name__)


class DocxDocument:
    """Main facade for reading and editing DOCX files.

    Usage:
        doc = DocxDocument('path/to/file.docx')

        # Chapter-based access
        chapter = doc.get_chapter('3.1')
        contents = doc.get_chapter_contents(chapter)

        # Replace chapter with markdown
        doc.replace_chapter('3.1', '# Title\\n\\nContent with **bold**.')

        # Search and replace
        doc.search_replace_in_chapter('3.1', 'old', 'new')

        # Delete operations
        doc.delete_chapter('3.1')
        doc.delete_in_chapter('3.1', 'text to delete')

        # Read comments and revisions
        comments = doc.read_comments()
        revisions = doc.read_revisions()

        # Save
        doc.save('output.docx')
    """

    def __init__(
        self,
        path: str,
        use_track_changes: bool = False,
        detect_manual_headings: bool = False,
        heading_format_config: Optional[dict] = None,
    ):
        """Load a DOCX document and build internal data structures.

        Args:
            path: Path to the .docx file
            use_track_changes: When True, write operations use win32com
                               for revision tracking. Default: False.
            detect_manual_headings: When True, also detect headings via
                formatting signals (bold + large font + spacing) when no
                heading style or outline level is set. Default: False.
            heading_format_config: Custom thresholds for manual heading
                detection (see ChapterParser for defaults).

        Raises:
            FileNotFoundError: If the path does not exist
            DocxError: If the file cannot be loaded as a DOCX
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"Document not found: {path}")

        self.path = path
        self.use_track_changes = use_track_changes
        self.detect_manual_headings = detect_manual_headings
        self._heading_format_config = heading_format_config

        # Load the document
        self.document = DocxDocumentLoader(path)

        # Build format store
        self.format_store = FormatStore(
            docx_path=path,
            document=self.document,
        )

        # Extract formatting
        self.format_store.formats_json = FormatExtractor.extract_all(self.document)

        # Build paragraph data
        self._build_paragraphs_data()

        # Initialize sub-modules
        self.chapter_parser = ChapterParser(
            self.format_store,
            detect_manual_headings=detect_manual_headings,
            heading_format_config=heading_format_config,
        )
        self.markdown_processor = MarkdownProcessor(self.format_store)
        self.mermaid_renderer = MermaidRenderer()

        # Track whether the document has been modified
        self._modified = False
        self._batch_counter = 0
        self._toc_position = None  # Remembered for auto-refresh

    def _build_paragraphs_data(self):
        """Build ParagraphData list from document paragraphs and format store,
        including run-level formatting (runs_data) from formats_json."""
        from .models import FormattingData
        self.format_store.paragraphs_data = []
        for i, para in enumerate(self.document.paragraphs):
            fmt_data = self.format_store.formats_json.get(i, {})

            # Build runs_data from formats_json['runs']
            runs_data: List[Tuple[str, FormattingData]] = []
            for run_dict in fmt_data.get('runs', []):
                text = run_dict.get('text', '')
                fmt = run_dict.get('formatting')
                if fmt is not None and not isinstance(fmt, FormattingData):
                    try:
                        fmt = FormattingData(**fmt)
                    except (TypeError, ValueError):
                        fmt = FormattingData()
                fmt = fmt or FormattingData()
                runs_data.append((text, fmt))

            pd = ParagraphData(
                paragraph=para,
                index=i,
                text=para.text,
                formatting=fmt_data.get('paragraph_format'),
                runs_data=runs_data,
            )
            self.format_store.paragraphs_data.append(pd)

    def _refresh_paragraphs_text(self):
        """Update ``ParagraphData.text`` from the current run state.

        ``search_replace_in_chapter`` and ``delete_in_chapter`` modify
        ``run.text`` in-place, but ``paragraphs_data`` cached the old
        ``para.text`` at build time.  This lightweight sync avoids a
        full ``_reload()``.
        """
        for i, pd in enumerate(self.format_store.paragraphs_data):
            if i < len(self.document.paragraphs):
                pd.text = self.document.paragraphs[i].text

    def _reload(self):
        """Reload the document after structural changes.

        After modifying the XML element tree, python-docx Paragraph objects
        become stale. This saves to temp, reloads, and rebuilds state.

        If called within a ``begin_batch()`` … ``end_batch()`` block the
        reload is deferred until the batch ends.
        """
        if self._batch_counter > 0:
            return  # Defer until batch ends

        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            self.document.save(tmp.name)
            self.document = DocxDocumentLoader(tmp.name)
        finally:
            os.unlink(tmp.name)

        # Rebuild all internal state
        self.format_store.document = self.document
        self.format_store.formats_json = FormatExtractor.extract_all(self.document)
        self._build_paragraphs_data()
        self.chapter_parser = ChapterParser(
            self.format_store,
            detect_manual_headings=self.detect_manual_headings,
            heading_format_config=self._heading_format_config,
        )
        self.markdown_processor = MarkdownProcessor(self.format_store)

    def _build_dominant_heading_template(self) -> Optional[ParagraphData]:
        """Build a merged template from all heading paragraphs.

        Headings typically use larger font sizes and a different east-asian
        font (often 黑体) than body text.  Returns a ``ParagraphData``
        with dominant heading formatting, or ``None``.
        """
        from collections import Counter
        from .models import ParagraphData, ParagraphFormatData, FormattingData

        para_formats: List[ParagraphFormatData] = []
        font_names: Counter = Counter()
        font_east: Counter = Counter()
        sizes: Counter = Counter()

        for node in self.chapter_parser._iter_all_nodes():
            idx = node.heading_paragraph_index
            if idx < len(self.format_store.paragraphs_data):
                pd = self.format_store.paragraphs_data[idx]
                if pd.formatting:
                    para_formats.append(pd.formatting)
                if pd.runs_data:
                    rf = pd.runs_data[0][1]
                    if rf.font_name:
                        font_names[rf.font_name] += 1
                    if rf.font_name_east_asia:
                        font_east[rf.font_name_east_asia] += 1
                    if rf.size:
                        sizes[rf.size] += 1

        if not para_formats:
            return None

        def _mode(attr, default=None):
            vals = Counter()
            for pf in para_formats:
                v = getattr(pf, attr, None)
                if v is not None:
                    vals[v] += 1
            return vals.most_common(1)[0][0] if vals else default

        merged_pf = ParagraphFormatData(
            alignment=_mode('alignment'),
            line_spacing=_mode('line_spacing'),
            line_spacing_rule=_mode('line_spacing_rule'),
            space_before=_mode('space_before'),
            space_after=_mode('space_after'),
        )

        merged_rf = FormattingData(
            font_name=font_names.most_common(1)[0][0] if font_names else None,
            font_name_east_asia=font_east.most_common(1)[0][0] if font_east else None,
            size=sizes.most_common(1)[0][0] if sizes else None,
        )

        return ParagraphData(formatting=merged_pf, runs_data=[('', merged_rf)])

    def _build_dominant_table_template(self) -> Optional[ParagraphData]:
        """Build a merged template from all table-cell paragraphs.

        Table cells often use a smaller font (e.g. 10.5pt 五号) and
        centred alignment, distinct from body text.

        Returns a ``ParagraphData`` with dominant table-cell formatting,
        or ``None`` if no tables exist.
        """
        from collections import Counter
        from .models import ParagraphData, ParagraphFormatData, FormattingData
        from docx.oxml.ns import qn

        font_names: Counter = Counter()
        font_east: Counter = Counter()
        sizes: Counter = Counter()
        alignments: Counter = Counter()

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.alignment is not None:
                            alignments[str(para.alignment)] += 1
                        for run in para.runs:
                            f = run.font
                            if f.name:
                                font_names[f.name] += 1
                            if f.size:
                                # Convert points → half-points to match FormattingData.size
                                sizes[int(f.size.pt * 2)] += 1
                            rPr = run._element.find(qn('w:rPr'))
                            if rPr is not None:
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is not None:
                                    ea = rFonts.get(qn('w:eastAsia'))
                                    if ea:
                                        font_east[ea] += 1

        if not font_names:
            return None

        merged_pf = ParagraphFormatData(
            alignment=alignments.most_common(1)[0][0] if alignments else None,
        )
        merged_rf = FormattingData(
            font_name=font_names.most_common(1)[0][0] if font_names else None,
            font_name_east_asia=font_east.most_common(1)[0][0] if font_east else None,
            size=sizes.most_common(1)[0][0] if sizes else None,
        )

        return ParagraphData(formatting=merged_pf, runs_data=[('', merged_rf)])

    def _build_dominant_body_template(self) -> Optional[ParagraphData]:
        """Build a merged template from the document's dominant body formatting.

        Scans all body paragraphs and picks the most representative value
        for each formatting attribute:
        - Font attributes (font name, east-asian font, size) → majority vote
        - Spacing attributes (space_before, space_after) → maximum value
          (most paragraphs default to 0; the max picks up the actual
          paragraph-spacing used in the document)
        - Structural attributes (alignment, indent, line_spacing) → mode

        Returns a ``ParagraphData`` with merged ``formatting`` and
        ``runs_data``, or ``None`` if no body paragraphs exist at all.
        """
        from collections import Counter
        from .models import ParagraphData, ParagraphFormatData, FormattingData

        para_formats: List[ParagraphFormatData] = []
        font_names: Counter = Counter()
        font_east: Counter = Counter()
        sizes: Counter = Counter()

        for node in self.chapter_parser._iter_all_nodes():
            for idx in node.body_paragraph_indices:
                if idx >= len(self.format_store.paragraphs_data):
                    continue
                pd = self.format_store.paragraphs_data[idx]
                if pd.formatting:
                    para_formats.append(pd.formatting)
                if pd.runs_data:
                    rf = pd.runs_data[0][1]
                    if rf.font_name:
                        font_names[rf.font_name] += 1
                    if rf.font_name_east_asia:
                        font_east[rf.font_name_east_asia] += 1
                    if rf.size:
                        sizes[rf.size] += 1

        if not para_formats:
            return None

        def _most_common(attr: str, default=None):
            vals = Counter()
            for pf in para_formats:
                v = getattr(pf, attr, None)
                if v is not None:
                    vals[v] += 1
            return vals.most_common(1)[0][0] if vals else default

        def _most_common_including_none(attr: str, default=None):
            """Mode of an attribute across all paragraphs, treating
            ``None`` (attribute not set) as a valid value."""
            vals = Counter()
            for pf in para_formats:
                vals[getattr(pf, attr, None)] += 1
            return vals.most_common(1)[0][0] if vals else default

        merged_pf = ParagraphFormatData(
            alignment=_most_common('alignment'),
            first_line_indent=_most_common('first_line_indent'),
            left_indent=_most_common('left_indent'),
            line_spacing=_most_common('line_spacing'),
            line_spacing_rule=_most_common('line_spacing_rule'),
            space_before=_most_common_including_none('space_before'),
            space_after=_most_common_including_none('space_after'),
        )

        # Run-level: majority-vote font name / east-asia / size
        merged_rf = FormattingData(
            font_name=font_names.most_common(1)[0][0] if font_names else None,
            font_name_east_asia=font_east.most_common(1)[0][0] if font_east else None,
            size=sizes.most_common(1)[0][0] if sizes else None,
        )

        return ParagraphData(
            formatting=merged_pf,
            runs_data=[('', merged_rf)],
        )

    def begin_batch(self):
        """Enter batch mode — defer ``_reload()`` until ``end_batch()``.

        Use when making many ``replace_chapter()`` / ``delete_chapter()``
        calls in a row to avoid O(N × document) repeated reloads.
        """
        self._batch_counter += 1

    def end_batch(self):
        """Exit batch mode and perform one final reload."""
        self._batch_counter = max(0, self._batch_counter - 1)
        if self._batch_counter == 0:
            self._reload()

    # ======================== Chapter Access ========================

    def get_chapter(self, number: Union[str, Tuple[int, ...]]) -> Optional[ChapterNode]:
        """Look up a chapter by number string like '3.1.1' or tuple (3,1,1).

        Args:
            number: Chapter number as string (e.g. "3.1") or tuple (e.g. (3,1))

        Returns:
            ChapterNode if found, None otherwise
        """
        return self.chapter_parser.get_chapter_by_number(number)

    def get_chapter_contents(self, chapter: ChapterNode) -> List[ParagraphData]:
        """Get all paragraphs belonging to a chapter (heading + body).

        Args:
            chapter: ChapterNode from get_chapter()

        Returns:
            List of ParagraphData objects
        """
        return self.chapter_parser.get_chapter_contents(chapter)

    def get_chapter_text(self, chapter: ChapterNode) -> str:
        """Get concatenated text of a chapter (including tables and image markers).

        Args:
            chapter: ChapterNode from get_chapter()

        Returns:
            Combined text of heading + body paragraphs + tables + image markers
        """
        return self.chapter_parser.get_chapter_text(chapter)

    def get_chapter_images(self, chapter: ChapterNode,
                           output_dir: str) -> Dict[str, str]:
        """Extract all embedded images in a chapter to files on disk.

        Args:
            chapter: ChapterNode from get_chapter()
            output_dir: Directory to save extracted image files

        Returns:
            Dict mapping image reference (e.g. ``"rId9"``) to saved file path
        """
        return self.chapter_parser.get_chapter_images(chapter, output_dir)

    def get_chapter_tree(self) -> List[ChapterNode]:
        """Get the root-level chapter tree.

        Returns:
            List of top-level ChapterNode objects
        """
        return self.chapter_parser.root_chapters

    def print_chapter_tree(self):
        """Print a visual representation of the chapter tree."""
        print(self.chapter_parser.tree_to_string())

    def insert_toc(
        self,
        position: Union[str, int, None] = None,
        *,
        title: str = '目录',
        max_level: int = 3,
        use_word_field: bool = False,
    ):
        """Insert a table of contents into the document.

        Gets TOC entries from the current chapter tree and inserts them
        at the specified position.  Supports both static and Word TOC field
        modes (see :meth:`DocxCreator.add_toc` for details on each mode).

        Args:
            position: Where to insert the TOC.
                - ``None``: at the very beginning (before first paragraph)
                - ``str``: chapter number string, e.g. ``'1'`` or ``'2.1'``
                  — inserts before that chapter's heading
                - ``int``: direct paragraph index in ``document.paragraphs``
            title: Heading text for the TOC section (default ``'目录'``).
                Pass ``None`` to skip the title.
            max_level: Maximum heading level to include (default 3).
            use_word_field: If True, insert a native Word TOC field
                instead of static paragraphs.
        """
        from docx.oxml import OxmlElement

        # Determine the reference paragraph index
        if isinstance(position, str):
            # Chapter number → find its heading paragraph index
            chapter = self.chapter_parser.get_chapter_by_number(position)
            if chapter is None:
                raise ChapterNotFoundError(
                    f'Chapter "{position}" not found in document'
                )
            ref_idx = chapter.heading_paragraph_index
        elif isinstance(position, int):
            ref_idx = position
        else:
            ref_idx = 0  # beginning

        # ---- Build TOC paragraph elements ----
        paragraphs_added = 0

        def _make_para(text: str, style_name: Optional[str] = None):
            """Create a ``<w:p>`` element with text and optional style."""
            p = OxmlElement('w:p')
            if style_name:
                ppr = OxmlElement('w:pPr')
                ps = OxmlElement('w:pStyle')
                ps.set(docx_qn('w:val'), style_name)
                ppr.append(ps)
                p.append(ppr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.set(docx_qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)
            p.append(r)
            return p

        def _make_toc_field(max_lvl: int):
            """Create a ``<w:p>`` with a Word TOC field code."""
            p = OxmlElement('w:p')

            # begin
            r1 = OxmlElement('w:r')
            fc1 = OxmlElement('w:fldChar')
            fc1.set(docx_qn('w:fldCharType'), 'begin')
            r1.append(fc1)
            p.append(r1)

            # instrText
            r2 = OxmlElement('w:r')
            instr = OxmlElement('w:instrText')
            instr.set(docx_qn('xml:space'), 'preserve')
            instr.text = f' TOC \\o "1-{max_lvl}" \\h \\z \\u '
            r2.append(instr)
            p.append(r2)

            # separate
            r3 = OxmlElement('w:r')
            fc3 = OxmlElement('w:fldChar')
            fc3.set(docx_qn('w:fldCharType'), 'separate')
            r3.append(fc3)
            p.append(r3)

            # placeholder text
            r4 = OxmlElement('w:r')
            rpr4 = OxmlElement('w:rPr')
            sz4 = OxmlElement('w:sz')
            sz4.set(docx_qn('w:val'), '18')
            rpr4.append(sz4)
            r4.append(rpr4)
            t4 = OxmlElement('w:t')
            t4.set(docx_qn('xml:space'), 'preserve')
            t4.text = '（请在 Word 中按 Ctrl+A → F9 刷新目录）'
            r4.append(t4)
            p.append(r4)

            # end
            r5 = OxmlElement('w:r')
            fc5 = OxmlElement('w:fldChar')
            fc5.set(docx_qn('w:fldCharType'), 'end')
            r5.append(fc5)
            p.append(r5)

            return p

        # Collect elements to insert
        elements_to_insert: List[Any] = []

        # Title
        if title:
            elements_to_insert.append(
                _make_para(title, 'TOC Heading')
            )
            elements_to_insert.append(
                _make_para('')
            )
            paragraphs_added += 2

        if use_word_field:
            elements_to_insert.append(
                _make_toc_field(max_level)
            )
            paragraphs_added += 1
        else:
            entries = self.chapter_parser.generate_toc_entries(
                max_level=max_level
            )
            if not entries:
                elements_to_insert.append(
                    _make_para('（无目录条目）')
                )
                paragraphs_added += 1
            else:
                for num_str, heading_text, level in entries:
                    if level > max_level:
                        continue
                    toc_style = f'TOC {min(level, 3)}'
                    if num_str:
                        display = f'{num_str}    {heading_text}'
                    else:
                        display = heading_text
                    elements_to_insert.append(
                        _make_para(display, toc_style)
                    )
                    paragraphs_added += 1

        # ---- Insert into document body ----
        body = self.document.element.body
        if ref_idx < len(self.document.paragraphs):
            ref_element = self.document.paragraphs[ref_idx]._element
            for elem in elements_to_insert:
                body.insert(
                    list(body).index(ref_element),
                    elem,
                )
        else:
            for elem in elements_to_insert:
                body.append(elem)

        # Mark document as modified and rebuild
        self._toc_position = position
        self._modified = True
        self._rebuild_after_toc_insertion()

    # ---- TOC Detection & Auto-Refresh ----

    @staticmethod
    def _paragraph_has_toc_style(para) -> bool:
        """Check if a paragraph has a TOC-related style at the XML level."""
        try:
            ppr = para._element.find(docx_qn('w:pPr'))
            if ppr is not None:
                pstyle = ppr.find(docx_qn('w:pStyle'))
                if pstyle is not None:
                    val = pstyle.get(docx_qn('w:val')) or ''
                    return val.startswith('TOC')
        except Exception:
            pass
        return False

    def has_toc(self) -> bool:
        """Check whether the document currently contains a TOC.

        Returns:
            True if at least one paragraph has a ``TOC *`` style.
        """
        for para in self.document.paragraphs:
            if self._paragraph_has_toc_style(para):
                return True
        return False

    def remove_toc(self):
        """Remove any existing TOC paragraphs from the document.

        Removes all paragraphs with ``TOC *`` or ``TOC Heading`` styles.
        After removal, internal structures are rebuilt.
        """
        to_remove: List[Any] = []
        for para in self.document.paragraphs:
            if self._paragraph_has_toc_style(para):
                to_remove.append(para._element)

        if not to_remove:
            return

        for elem in to_remove:
            try:
                elem.getparent().remove(elem)
            except Exception:
                pass

        self._toc_position = None
        self._rebuild_after_toc_insertion()
        self._modified = True

    def refresh_toc(self) -> bool:
        """Refresh the TOC if one exists in the document.

        Removes the stale TOC and inserts a fresh one at the same
        position using the current chapter tree.  Does nothing if no
        TOC exists.

        Returns:
            True if the TOC was refreshed, False if no TOC existed.
        """
        if not self.has_toc():
            return False

        position = self._toc_position
        self.remove_toc()

        # insert_toc saves _toc_position and rebuilds
        self.insert_toc(position=position)
        return True

    def _auto_refresh_toc(self):
        """Called after structural changes (replace/delete chapter).

        Silently refreshes the TOC if one exists — no-op otherwise.
        """
        try:
            self.refresh_toc()
        except Exception as exc:
            logger.warning('TOC auto-refresh failed: %s', exc)

    def _rebuild_after_toc_insertion(self):
        """Rebuild internal data after TOC insertion.

        The paragraph list has shifted; we re-extract formatting from
        scratch and rebuild the chapter tree so subsequent operations
        are consistent and heading detection is correct.
        """
        self.format_store.formats_json = FormatExtractor.extract_all(
            self.document
        )
        self._build_paragraphs_data()
        self.chapter_parser = ChapterParser(
            self.format_store,
            detect_manual_headings=self.detect_manual_headings,
            heading_format_config=self._heading_format_config,
        )

    # ======================== Content Replacement ========================

    def replace_chapter(
        self,
        number: Union[str, Tuple[int, ...]],
        md_text: str,
    ):
        """Replace a chapter's entire content with markdown text.

        Preserves the original heading paragraph and replaces body content
        with parsed markdown elements. Supports tables, images, mermaid
        diagrams, lists, and inline formatting.

        Args:
            number: Chapter number to replace
            md_text: Markdown text for the new content

        Raises:
            ChapterNotFoundError: If the chapter number doesn't exist
            DocxError: If the replacement fails
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        if self.use_track_changes:
            self._replace_chapter_win32(chapter, md_text)
        else:
            self._replace_chapter_pythondocx(chapter, md_text)

    def _replace_chapter_pythondocx(
        self,
        chapter: ChapterNode,
        md_text: str,
    ):
        """Replace chapter content using python-docx XML manipulation.

        If the markdown starts with a heading, the original heading text is
        updated in-place (no duplicate heading created).  Body content after
        the heading (if present) is inserted as new body paragraphs.
        """
        # Parse markdown
        elements = self.markdown_processor.parse_markdown(md_text)

        # Handle mermaid elements: render to images
        elements = self._render_mermaid_elements(elements)

        # Build body template from the document's dominant (majority-vote)
        # formatting rather than from the chapter's own paragraphs.  After
        # potentially-corrupt previous replacements the chapter's own body
        # paragraphs may carry anomalous formatting (e.g. space_before from
        # a one-off heading-like paragraph) that would pollute the output.
        body_template = self._build_dominant_body_template()

        # heading_template: from the document's dominant heading format.
        # The font size in the template won't override built-in heading
        # styles (which already set correct per-level sizes), but font
        # name / east-asian font will be consistently applied.
        heading_template = self._build_dominant_heading_template()

        # Update heading text if the markdown starts with a heading element.
        # Preserve run structure: only replace text in the first text-bearing run,
        # clear other text runs, and leave non-content runs (tabs, page-refs, etc.) alone.
        if elements and elements[0].type == MarkdownElementType.HEADING:
            new_heading_text = elements[0].text
            heading_para = self.document.paragraphs[chapter.heading_paragraph_index]
            text_set = False
            for run in heading_para.runs:
                txt = run.text or ''
                if txt.strip():
                    if not text_set:
                        run.text = new_heading_text
                        text_set = True
                    else:
                        run.text = ''
            # Remove the heading from elements so it won't be duplicated
            elements = elements[1:]

        # Get the heading element reference (may have been modified above)
        heading_para = self.document.paragraphs[chapter.heading_paragraph_index]
        body = self.document.element.body
        heading_element = heading_para._element

        # Delete existing body paragraphs (work backwards)
        for idx in sorted(chapter.body_paragraph_indices, reverse=True):
            try:
                para = self.document.paragraphs[idx]
                body.remove(para._element)
            except Exception:
                logger.debug(
                    "Could not remove body paragraph %d; may already be gone", idx
                )

        # Delete existing tables in the chapter (work backwards)
        table_indices = sorted(
            self.chapter_parser._collect_table_indices(chapter),
            reverse=True
        )
        if table_indices:
            # Collect tbl elements before any removals shift the tables list
            all_tbl_elements = [
                t._element for t in self.document.tables
            ]
            for ti in table_indices:
                if ti < len(all_tbl_elements):
                    try:
                        body.remove(all_tbl_elements[ti])
                    except Exception:
                        logger.debug(
                            "Could not remove table %d; may already be gone", ti
                        )

        # Build new docx elements with content-type-specific templates
        table_template = self._build_dominant_table_template()
        new_elements = self.markdown_processor.build_elements_for_chapter(
            elements, body_template, heading_template, table_template
        )

        # Insert new elements after the heading (in reverse order to maintain sequence)
        for elem in reversed(new_elements):
            try:
                heading_element.addnext(elem)
            except Exception:
                logger.debug("Could not insert new element after heading", exc_info=True)

        # Reload to refresh internal state
        self._reload()
        self._modified = True
        self._auto_refresh_toc()

    def _replace_chapter_win32(
        self,
        chapter: ChapterNode,
        md_text: str,
    ):
        """Replace chapter content with COM automation, producing revision marks.

        Tries in order:
        1. Microsoft Word (via ``Win32Ops``)
        2. WPS Office (via ``WpsOps``, independent module)
        3. python-docx direct save (no tracked changes, as last resort)

        The COM Compare API generates proper ``<w:ins>`` / ``<w:del>``
        tracked changes by diffing the original and modified documents.
        """
        # --- 1. Save original and modified versions side-by-side ---
        original_tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        original_tmp.close()
        modified_tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        modified_tmp.close()

        try:
            # Snapshot the document BEFORE python-docx modifies it
            self.document.save(original_tmp.name)

            # Do the python-docx replacement
            self._replace_chapter_pythondocx(chapter, md_text)
            # self.document now has the modified content in python-docx's model
            self.document.save(modified_tmp.name)

            # --- 2. Try COM compare (Word first, then WPS) ---
            success = self._try_com_compare(
                original_tmp.name, modified_tmp.name,
            )
            if not success:
                logger.warning(
                    "Neither Word nor WPS could produce tracked changes. "
                    "Falling back to python-docx save (no tracked changes)."
                )
                self.document.save(self.path)

        finally:
            for tmp_path in (original_tmp.name, modified_tmp.name):
                try:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                except Exception:
                    logger.warning("Failed to clean up temp file: %s", tmp_path)

        # Reload
        self.document = DocxDocumentLoader(self.path)
        self._reload()

    def _try_com_compare(self, original_path: str, modified_path: str) -> bool:
        """Try to produce tracked changes via Word or WPS COM.

        Args:
            original_path: Path to the original (before) document
            modified_path: Path to the modified (after) document

        Returns:
            True if COM compare and save succeeded, False otherwise
        """
        # ── Attempt 1: Microsoft Word ──
        if Win32Ops.is_word_available():
            try:
                with Win32Ops() as ops:
                    ops.open_document(original_path)
                    ops.word.ActiveDocument.Compare(
                        Name=modified_path,
                        CompareTarget=2,            # wdCompareTargetNew
                        IgnoreAllComparisonWarnings=True,
                    )
                    ops.wd_doc = ops.word.ActiveDocument
                    ops.save_document(self.path)
                return True
            except Win32ComError as e:
                logger.debug("Word compare failed: %s", e)
        else:
            logger.debug("Microsoft Word not available, skipping.")

        # ── Attempt 2: WPS Office ──
        if WpsOps.is_wps_available():
            try:
                with WpsOps() as ops:
                    ops.open_document(original_path)
                    # WPS Document.Compare() is API-compatible with Word's
                    ops.word.ActiveDocument.Compare(
                        Name=modified_path,
                        CompareTarget=2,
                        IgnoreAllComparisonWarnings=True,
                    )
                    ops.wd_doc = ops.word.ActiveDocument
                    ops.save_document(self.path)
                return True
            except WpsComError as e:
                logger.debug("WPS compare failed: %s", e)
        else:
            logger.debug("WPS Office not available, skipping.")

        return False

    def _render_mermaid_elements(
        self,
        elements: List,
    ) -> List:
        """Render mermaid elements to images in-place.

        Args:
            elements: List of MarkdownElement objects

        Returns:
            Modified elements list with mermaid blocks replaced by images
        """
        rendered = []
        for elem in elements:
            if elem.type == MarkdownElementType.MERMAID:
                # Try to render mermaid
                if self.mermaid_renderer.is_available() and elem.text:
                    try:
                        img_path = self.mermaid_renderer.render(elem.text)
                        rendered.append(MarkdownElement(
                            type=MarkdownElementType.IMAGE,
                            image_path=img_path,
                            alt_text="Mermaid diagram",
                        ))
                        continue
                    except Exception:
                        pass
                # Fallback: keep as code block
                rendered.append(MarkdownElement(
                    type=MarkdownElementType.CODE_BLOCK,
                    text=elem.text or '',
                    code_language='mermaid',
                ))
            else:
                rendered.append(elem)
        return rendered

    # ======================== Search and Replace ========================

    @staticmethod
    def _prepare_paragraph_for_replace(
        para,
        old_text: str,
        case_sensitive: bool = False,
    ) -> bool:
        """Ensure *old_text* is findable within a single run.

        When ``old_text`` crosses a run boundary (e.g. bold turns off mid-word),
        Word splits the text across ``<w:r>`` elements.  This method merges all
        runs in the paragraph into the first run so that ``run.text`` contains
        the full paragraph text and a simple ``old_text in run.text`` check
        works.  Formatting from the first run is preserved; per-run formatting
        details are lost.

        Returns ``True`` if ``old_text`` can be found in the paragraph text
        (either natively or after merging).
        """
        full = para.text
        if not full:
            return False

        if case_sensitive:
            if old_text not in full:
                return False
            # Already findable in a single run?
            if any(old_text in r.text for r in para.runs if r.text):
                return True
        else:
            lower_old = old_text.lower()
            if lower_old not in full.lower():
                return False
            if any(lower_old in r.text.lower() for r in para.runs if r.text):
                return True

        # Cross-run merge: concatenate all text into the first run
        runs = para.runs
        if runs:
            runs[0].text = full
            for r in runs[1:]:
                r.text = ''
        return True

    def search_replace_in_chapter(
        self,
        number: Union[str, Tuple[int, ...]],
        old_text: str,
        new_text: str,
        case_sensitive: bool = False,
    ) -> int:
        """Search and replace text within a specific chapter.

        Iterates through the chapter's paragraphs and replaces text
        in each run. Operates at the text level, not XML level.

        .. note::
           If ``old_text`` crosses a run boundary (uncommon in well-formed
           documents), runs are merged so the replacement still happens.
           Doing so loses per-run formatting differentiation for that
           paragraph.

        Args:
            number: Chapter number
            old_text: Text to search for
            new_text: Replacement text
            case_sensitive: Whether to match case

        Returns:
            Number of replacements made
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        count = 0
        indices = [chapter.heading_paragraph_index] + list(chapter.body_paragraph_indices)

        for idx in indices:
            if idx >= len(self.document.paragraphs):
                continue
            para = self.document.paragraphs[idx]

            # Cross-run fallback
            if not self._prepare_paragraph_for_replace(para, old_text, case_sensitive):
                continue

            for run in para.runs:
                if not run.text:
                    continue
                if case_sensitive:
                    occurrences = run.text.count(old_text)
                    if occurrences:
                        run.text = run.text.replace(old_text, new_text)
                        count += occurrences
                else:
                    matches = run.text.lower().count(old_text.lower())
                    if matches:
                        run.text = self._replace_ignore_case(
                            run.text, old_text, new_text
                        )
                        count += matches

        if count > 0:
            self._modified = True
            self._refresh_paragraphs_text()

        return count

    @staticmethod
    def _replace_ignore_case(text: str, old: str, new: str) -> str:
        """Replace *old* with *new* ignoring case.

        .. caution::
           The replacement text ``new`` is used **as-is** — the capitalisation
           pattern of the matched text is **not** preserved (despite the
           earlier docstring claiming otherwise).

        Args:
            text: Original text
            old: Text to find (case-insensitive)
            new: Replacement text

        Returns:
            Text with replacements
        """
        return re.sub(re.escape(old), new, text, flags=re.IGNORECASE)

    # ======================== Deletion ========================

    def delete_chapter(self, number: Union[str, Tuple[int, ...]]):
        """Delete an entire chapter (heading + all body content).

        Args:
            number: Chapter number to delete

        Raises:
            ChapterNotFoundError: If the chapter doesn't exist
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        self.chapter_parser.delete_chapter(chapter)
        self._reload()
        self._modified = True
        self._auto_refresh_toc()

    def delete_in_chapter(
        self,
        number: Union[str, Tuple[int, ...]],
        target_text: str,
    ) -> int:
        """Delete specific text content from within a chapter.

        Args:
            number: Chapter number
            target_text: Text to delete

        Returns:
            Number of deletions made
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        count = 0
        indices = [chapter.heading_paragraph_index] + list(chapter.body_paragraph_indices)

        for idx in indices:
            if idx >= len(self.document.paragraphs):
                continue
            para = self.document.paragraphs[idx]

            # Cross-run fallback
            if not self._prepare_paragraph_for_replace(para, target_text, case_sensitive=True):
                continue

            for run in para.runs:
                if run.text and target_text in run.text:
                    occurrences = run.text.count(target_text)
                    run.text = run.text.replace(target_text, '')
                    count += occurrences

        if count > 0:
            self._modified = True
            self._refresh_paragraphs_text()

        return count

    # ======================== Comments and Revisions ========================

    def read_comments(self) -> List[CommentData]:
        """Read comments from the document.

        Tries python-docx first, falls back to win32com.

        Returns:
            List of CommentData objects
        """
        try:
            return self._read_comments_python_docx()
        except Exception:
            if Win32Ops.is_word_available():
                return self._read_comments_win32()
            raise CommentsReadError("Cannot read comments from document")

    def _read_comments_python_docx(self) -> List[CommentData]:
        """Read comments using python-docx API.

        Also parses ``w:commentRangeStart`` elements from the document body
        to populate each ``CommentData.paragraph_index``, allowing callers
        to map a comment to its containing paragraph (and chapter).

        Raises ``CommentsReadError`` if the ``.comments`` attribute is
        unavailable (python-docx < 0.8.11), allowing the caller's win32com
        fallback to fire.
        """
        try:
            comments_part = self.document.comments
        except AttributeError:
            raise CommentsReadError(
                "python-docx version does not support .comments attribute"
            )

        if comments_part is None:
            return []

        qn = docx_qn
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        results = []

        # ---- Build comment_id → paragraph_index mapping ----
        # Walk all paragraphs in document order, looking for
        # w:commentRangeStart elements whose parent paragraph tells us
        # which paragraph the comment anchors to.
        comment_id_to_para: Dict[str, int] = {}
        for para_index, para in enumerate(self.document.paragraphs):
            p_element = para._element
            for crs in p_element.iter(f'{{{W}}}commentRangeStart'):
                cid = crs.get(f'{{{W}}}id')
                if cid is not None and cid not in comment_id_to_para:
                    comment_id_to_para[cid] = para_index

        # ---- Read comments ----
        for comment in comments_part:
            try:
                cid = comment._element.get(qn('w:id'), '')
                parent_id = comment._element.get(qn('w:parent'))
                results.append(CommentData(
                    id=cid,
                    author=comment._element.get(qn('w:author'), ''),
                    date=comment._element.get(qn('w:date'), ''),
                    text=comment.text.strip() if comment.text else '',
                    paragraph_index=comment_id_to_para.get(str(cid)),
                    parent_id=str(parent_id) if parent_id is not None else None,
                ))
            except Exception:
                logger.debug("Skipping a comment that could not be read", exc_info=True)

        return results

    def _read_comments_win32(self) -> List[CommentData]:
        """Fallback: read comments via win32com."""
        with Win32Ops() as ops:
            ops.open_document(self.path)
            return ops.read_comments()

    def read_revisions(self) -> List[RevisionData]:
        """Read tracked changes from the document.

        Parses document XML directly for w:ins and w:del elements.
        Falls back to win32com if XML parsing fails.

        Returns:
            List of RevisionData sorted by document position
        """
        try:
            return self._read_revisions_xml()
        except Exception:
            if Win32Ops.is_word_available():
                return self._read_revisions_win32()
            raise

    def _read_revisions_xml(self) -> List[RevisionData]:
        """Read tracked changes by parsing document XML directly.

        Uses ``document.paragraphs`` order for indices so the returned
        ``paragraph_index`` values match ``document.paragraphs[i]``
        (unlike the old code which iterated ``body.iterchildren()`` and
        counted tables as an index slot, causing drift).
        """
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        revisions = []

        for para_index, para in enumerate(self.document.paragraphs):
            p_element = para._element

            # Find w:ins elements (insertions)
            for ins in p_element.iter(f'{{{W}}}ins'):
                text_parts = []
                for t in ins.iter(f'{{{W}}}t'):
                    if t.text:
                        text_parts.append(t.text)
                revisions.append(RevisionData(
                    rev_id=ins.get(f'{{{W}}}id', ''),
                    author=ins.get(f'{{{W}}}author', ''),
                    date=ins.get(f'{{{W}}}date', ''),
                    type='insertion',
                    text=''.join(text_parts),
                    paragraph_index=para_index,
                ))

            # Find w:del elements (deletions)
            for del_elem in p_element.iter(f'{{{W}}}del'):
                text_parts = []
                for dt in del_elem.iter(f'{{{W}}}delText'):
                    if dt.text:
                        text_parts.append(dt.text)
                revisions.append(RevisionData(
                    rev_id=del_elem.get(f'{{{W}}}id', ''),
                    author=del_elem.get(f'{{{W}}}author', ''),
                    date=del_elem.get(f'{{{W}}}date', ''),
                    type='deletion',
                    text=''.join(text_parts),
                    paragraph_index=para_index,
                ))

        return revisions

    def _read_revisions_win32(self) -> List[RevisionData]:
        """Fallback: read tracked changes via win32com."""
        with Win32Ops() as ops:
            ops.open_document(self.path)
            return ops.read_revisions()

    # ======================== Comment Utilities ========================

    def get_comment_context(
        self, comment_id: str
    ) -> Optional[Dict[str, Any]]:
        """Get chapter context and surrounding text for a comment.

        Args:
            comment_id: The ``w:id`` of the comment

        Returns:
            Dict with ``comment``, ``chapter_number``, ``chapter_title``,
            ``heading_level``, ``context_text``, and ``paragraph_index``,
            or ``None`` if the comment ID is not found.
        """
        # Find the comment by ID
        comment = None
        for c in self.read_comments():
            if c.id == comment_id:
                comment = c
                break

        if comment is None:
            return None

        result: Dict[str, Any] = {
            'comment': comment,
            'chapter_number': None,
            'chapter_title': None,
            'heading_level': None,
            'context_text': None,
            'paragraph_index': comment.paragraph_index,
        }

        if comment.paragraph_index is not None:
            # Map to chapter
            chapter = self.chapter_parser.get_chapter_for_paragraph(
                comment.paragraph_index
            )
            if chapter is not None:
                result['chapter_number'] = chapter.to_string()
                result['chapter_title'] = chapter.heading_text
                result['heading_level'] = chapter.heading_level

            # Extract surrounding paragraph text (±3)
            start = max(0, comment.paragraph_index - 3)
            end = min(
                len(self.document.paragraphs),
                comment.paragraph_index + 4,
            )
            context_lines = []
            for i in range(start, end):
                prefix = '>>> ' if i == comment.paragraph_index else '    '
                text = self.document.paragraphs[i].text.strip()
                if text:
                    context_lines.append(f'{prefix}{text}')
            result['context_text'] = '\n'.join(context_lines)

        return result

    def delete_comment(self, comment_id: str) -> bool:
        """Delete a comment by its ID.

        Removes the comment definition from the comments part and all its
        reference elements (``w:commentRangeStart``, ``w:commentRangeEnd``,
        ``w:commentReference``) from the document body.

        Args:
            comment_id: The ``w:id`` of the comment to delete

        Returns:
            ``True`` if anything was removed, ``False`` if the comment was
            not found or the document has no comments part
        """
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # 1. Remove comment definition from comments part
        try:
            comments_part = self.document.comments
        except AttributeError:
            comments_part = None

        removed = False
        if comments_part is not None:
            for comment in list(comments_part):
                try:
                    elem_id = str(comment._element.get(qn('w:id'), ''))
                except Exception:
                    continue
                if elem_id == comment_id:
                    try:
                        comment._element.getparent().remove(comment._element)
                        removed = True
                    except Exception:
                        logger.debug(
                            'Failed to remove comment %s from part', comment_id
                        )

        # 2. Remove reference elements from document body
        body = self.document.element.body
        for elem in list(body.iter()):
            tag = docx_qn('w:commentRangeStart')
            if elem.tag == tag or elem.tag == docx_qn('w:commentRangeEnd') or \
                    elem.tag == docx_qn('w:commentReference'):
                try:
                    elem_id = str(elem.get(qn('w:id'), ''))
                except Exception:
                    continue
                if elem_id == comment_id:
                    try:
                        elem.getparent().remove(elem)
                        removed = True
                    except Exception:
                        logger.debug(
                            'Failed to remove comment reference %s', comment_id
                        )

        if removed:
            self._modified = True

        return removed

    def delete_all_comments(self) -> int:
        """Delete all comments from the document.

        Returns:
            Number of comments deleted
        """
        comments = self.read_comments()
        count = 0
        for c in comments:
            if self.delete_comment(c.id):
                count += 1
        return count

    # ======================== Format Utilities ========================

    def apply_format(
        self,
        format_spec: Dict[str, Any],
        target: Union[str, List[Union[str, Tuple[int, ...]]]] = 'all',
    ) -> int:
        """Apply formatting to paragraphs matching the target criteria.

        This method modifies paragraph-level and run-level formatting across
        the document.  Supported ``format_spec`` keys:

        **Run-level** (applied to every run in each matching paragraph):
            ``font_name``, ``font_name_east_asia``, ``font_size`` (in points),
            ``bold`` (bool), ``italic`` (bool), ``underline`` (bool or str)

        **Paragraph-level**:
            ``alignment`` (``'LEFT'`` | ``'CENTER'`` | ``'RIGHT'`` | ``'JUSTIFY'``),
            ``first_line_indent`` (points), ``left_indent`` (points),
            ``line_spacing`` (float, e.g. 1.5), ``line_spacing_rule``
            (``'SINGLE'`` | ``'DOUBLE'`` | ``'MULTIPLE'``),
            ``space_before`` (points), ``space_after`` (points)

        Args:
            format_spec: Dict of formatting properties to apply
            target: ``'all'`` (default), ``'body'``, ``'heading'``, or a list
                of chapter numbers (strings like ``'3.1'`` or tuples).

        Returns:
            Number of paragraphs whose formatting was modified

        Raises:
            ValueError: If an unknown target or invalid format value is given
        """
        count = 0

        # Determine which paragraph indices to process
        indices = self._resolve_target_indices(target)

        for idx in indices:
            if idx >= len(self.document.paragraphs):
                continue
            para = self.document.paragraphs[idx]
            modified = False

            # ---- Paragraph-level formatting ----
            pf = para.paragraph_format
            if 'alignment' in format_spec:
                val = format_spec['alignment']
                align_map = {
                    'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                    'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                    'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                    'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    'BOTH': WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                mapped = align_map.get(val.upper() if isinstance(val, str) else val)
                if mapped is not None:
                    pf.alignment = mapped
                    modified = True

            if 'first_line_indent' in format_spec:
                val = format_spec['first_line_indent']
                if val is not None:
                    pf.first_line_indent = Emu(int(val * 12700))  # points → EMU
                    modified = True

            if 'left_indent' in format_spec:
                val = format_spec['left_indent']
                if val is not None:
                    pf.left_indent = Emu(int(val * 12700))
                    modified = True

            if 'line_spacing' in format_spec:
                val = format_spec['line_spacing']
                if val is not None:
                    pf.line_spacing = float(val)
                    modified = True

            if 'line_spacing_rule' in format_spec:
                val = format_spec['line_spacing_rule']
                rule_map = {
                    'SINGLE': 0,
                    'DOUBLE': 1,
                    'MULTIPLE': 2,
                    'AT_LEAST': 3,
                    'EXACTLY': 4,
                }
                mapped = rule_map.get(val.upper() if isinstance(val, str) else val)
                if mapped is not None:
                    pf.line_spacing_rule = val
                    modified = True

            if 'space_before' in format_spec:
                val = format_spec['space_before']
                if val is not None:
                    pf.space_before = Pt(float(val))
                    modified = True

            if 'space_after' in format_spec:
                val = format_spec['space_after']
                if val is not None:
                    pf.space_after = Pt(float(val))
                    modified = True

            # ---- Run-level formatting ----
            for run in para.runs:
                run_modified = False
                if 'font_name' in format_spec and format_spec['font_name'] is not None:
                    run.font.name = format_spec['font_name']
                    run_modified = True
                if 'font_name_east_asia' in format_spec \
                        and format_spec['font_name_east_asia'] is not None:
                    # Set east-asian font via XML
                    rpr = run._element.get_or_add_rPr()
                    rFonts = rpr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = run._element.makeelement(
                            qn('w:rFonts'), {}
                        )
                        rpr.insert(0, rFonts)
                    rFonts.set(qn('w:eastAsia'), format_spec['font_name_east_asia'])
                    run_modified = True
                if 'font_size' in format_spec and format_spec['font_size'] is not None:
                    run.font.size = Pt(float(format_spec['font_size']))
                    run_modified = True
                if 'bold' in format_spec:
                    run.bold = bool(format_spec['bold'])
                    run_modified = True
                if 'italic' in format_spec:
                    run.italic = bool(format_spec['italic'])
                    run_modified = True
                if 'underline' in format_spec and format_spec['underline'] is not None:
                    run.underline = format_spec['underline']
                    run_modified = True

                if run_modified:
                    modified = True

            if modified:
                count += 1

        if count > 0:
            self._modified = True
            self._refresh_paragraphs_text()

        return count

    def _resolve_target_indices(
        self, target: Union[str, List[Union[str, Tuple[int, ...]]]]
    ) -> List[int]:
        """Resolve a target specifier into a list of paragraph indices.

        Args:
            target: ``'all'``, ``'body'``, ``'heading'``, or chapter list

        Returns:
            Sorted list of paragraph indices
        """
        indices: List[int] = []

        if isinstance(target, str):
            if target == 'all':
                return list(range(len(self.document.paragraphs)))

            elif target == 'body':
                for ch in self.chapter_parser.list_chapters():
                    indices.extend(ch.body_paragraph_indices)
                return sorted(set(indices))

            elif target == 'heading':
                for ch in self.chapter_parser.list_chapters():
                    indices.append(ch.heading_paragraph_index)
                return sorted(set(indices))

            else:
                raise ValueError(
                    f"Unknown target '{target}'. "
                    "Use 'all', 'body', 'heading', or a list of chapter numbers."
                )

        # List of chapter numbers
        for num in target:
            chapter = self.chapter_parser.get_chapter_by_number(num)
            if chapter is None:
                raise ChapterNotFoundError(f"Chapter {num} not found")
            indices.append(chapter.heading_paragraph_index)
            indices.extend(chapter.body_paragraph_indices)

        return sorted(set(indices))

    # ======================== Save ========================

    def save(self, path: Optional[str] = None):
        """Save the document.

        Args:
            path: Output path. If None, overwrites the original file.

        Note:
            When saving to a new path, ``self.path`` is updated so that
            subsequent track-changes operations target the new file.
        """
        save_path = path or self.path
        self.document.save(save_path)
        self.path = save_path
        self._modified = False

    # ======================== Properties ========================

    @property
    def is_modified(self) -> bool:
        """Check if the document has been modified since load/save."""
        return self._modified
