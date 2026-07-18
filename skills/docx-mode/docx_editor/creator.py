"""Document creator module.

Provides utilities for creating new DOCX documents from scratch,
supporting markdown content and configurable default formatting.

Usage:
    from docx_editor.creator import DocxCreator

    creator = DocxCreator()
    creator.set_default_format({
        'font_name': '宋体',
        'font_name_east_asia': '宋体',
        'font_size': 12,
        'line_spacing': 1.5,
    })
    creator.add_markdown('# Title\\n\\nBody content...')
    creator.save('output.docx')
"""
import logging
import os
import tempfile
from typing import Any, Dict, List, Optional

from docx import Document as NewDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

from .markdown_processor import MarkdownProcessor
from .mermaid_renderer import MermaidRenderer
from .models import (
    FormatStore, MarkdownElement, MarkdownElementType,
)

logger = logging.getLogger(__name__)


class DocxCreator:
    """Create a new DOCX document from scratch.

    Usage:
        creator = DocxCreator()
        creator.set_default_format({'font_name': 'SimSun', 'font_size': 12})
        creator.add_markdown(md_text)
        creator.save('output.docx')

    Or as a single-shot:
        DocxCreator.create('output.docx', md_text, format_spec={...})
    """

    def __init__(self, template_path: Optional[str] = None):
        """Initialize a new document.

        Args:
            template_path: Optional path to a .docx to use as style
                template. If omitted, python-docx defaults are used.
        """
        if template_path:
            self.document = NewDocument(template_path)
        else:
            self.document = NewDocument()

        self.format_store = FormatStore(
            docx_path='',
            document=self.document,
        )
        self.markdown_processor = MarkdownProcessor(self.format_store)
        self.mermaid_renderer = MermaidRenderer()

    def set_default_format(self, format_spec: Dict[str, Any]):
        """Configure default paragraph and run formatting.

        This modifies the document's default style so that all new content
        inherits these properties.  Supported keys:

        - ``font_name``: Western font name (e.g. ``'Times New Roman'``)
        - ``font_name_east_asia``: East-Asian font name (e.g. ``'宋体'``)
        - ``font_size``: Font size in points (e.g. ``12``)
        - ``line_spacing``: Line spacing multiplier (e.g. ``1.5``)
        - ``line_spacing_rule``: ``'SINGLE'`` | ``'MULTIPLE'`` | etc.
        - ``space_before`` / ``space_after``: Paragraph spacing in points
        - ``alignment``: ``'LEFT'`` | ``'CENTER'`` | ``'RIGHT'`` | ``'JUSTIFY'``
        - ``first_line_indent``: Indent in points (e.g. ``24`` = 2 chars at 12pt)

        Args:
            format_spec: Dict of formatting properties
        """
        # Modify the document's default 'Normal' style
        style = self.document.styles['Normal']

        # Run-level defaults on the style
        font = style.font
        if 'font_name' in format_spec and format_spec['font_name'] is not None:
            font.name = format_spec['font_name']
        if 'font_name_east_asia' in format_spec \
                and format_spec['font_name_east_asia'] is not None:
            # Set east-asian font via XML on the style
            rpr = style.element.find(qn('w:rPr'))
            if rpr is None:
                rpr = style.element.makeelement(qn('w:rPr'), {})
                style.element.append(rpr)
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rpr.makeelement(qn('w:rFonts'), {})
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), format_spec['font_name_east_asia'])
        if 'font_size' in format_spec and format_spec['font_size'] is not None:
            font.size = Pt(float(format_spec['font_size']))

        # Paragraph-level defaults on the style
        pf = style.paragraph_format
        if 'line_spacing' in format_spec and format_spec['line_spacing'] is not None:
            pf.line_spacing = float(format_spec['line_spacing'])
        if 'line_spacing_rule' in format_spec \
                and format_spec['line_spacing_rule'] is not None:
            pf.line_spacing_rule = format_spec['line_spacing_rule']
        if 'space_before' in format_spec and format_spec['space_before'] is not None:
            pf.space_before = Pt(float(format_spec['space_before']))
        if 'space_after' in format_spec and format_spec['space_after'] is not None:
            pf.space_after = Pt(float(format_spec['space_after']))
        if 'alignment' in format_spec and format_spec['alignment'] is not None:
            align_map = {
                'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'BOTH': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            mapped = align_map.get(format_spec['alignment'].upper())
            if mapped is not None:
                pf.alignment = mapped
        if 'first_line_indent' in format_spec \
                and format_spec['first_line_indent'] is not None:
            pf.first_line_indent = Emu(
                int(format_spec['first_line_indent'] * 12700)
            )

    def add_markdown(self, md_text: str):
        """Parse markdown text and append its content to the document.

        Supports all markdown features that ``MarkdownProcessor`` handles:
        headings, paragraphs, tables, code blocks, mermaid (if mmdc is
        available), images, lists, blockquotes, and inline formatting.

        Args:
            md_text: Markdown text to parse and convert
        """
        if not md_text or not md_text.strip():
            return

        elements = self.markdown_processor.parse_markdown(md_text)

        # Render mermaid diagrams to images when possible
        rendered = self._render_mermaid_elements(elements)

        self.markdown_processor.apply_to_document(rendered)

    def add_heading(self, text: str, level: int = 1):
        """Add a heading paragraph.

        Args:
            text: Heading text
            level: Heading level (1-9)
        """
        self.document.add_heading(text, level=min(max(level, 1), 9))

    def add_paragraph(self, text: str):
        """Add a body paragraph.

        Args:
            text: Paragraph text (supports inline markdown formatting)
        """
        from .markdown_processor import MarkdownProcessor as MP
        segments = MP.parse_inline_formatting(text)
        para = self.document.add_paragraph()
        for seg_text, fmt_overrides in segments:
            run = para.add_run(seg_text)
            MP(None)._apply_inline_format(run, fmt_overrides)

    def add_toc(
        self,
        entries: Optional[List[Tuple[str, str, int]]] = None,
        *,
        title: str = '目录',
        max_level: int = 3,
        use_word_field: bool = False,
    ):
        """Insert a table of contents into the document.

        Two modes:
        - **Static** (default): one paragraph per heading entry, using Word's
          built-in ``TOC 1`` / ``TOC 2`` / ``TOC 3`` styles, with dot leaders.
        - **Word field** (``use_word_field=True``): inserts a native Word TOC
          field (``TOC \\o "1-3" \\h \\z \\u``).  The TOC is populated when the
          user opens the document in Word and presses Ctrl+A → F9.

        Args:
            entries: List of ``(number_string, heading_text, level)`` tuples,
                e.g. from ``ChapterParser.generate_toc_entries()``.
                If ``None``, a placeholder message is inserted instead
                (useful when the document will be populated later).
            title: Heading text for the TOC section (default ``'目录'``).
                Pass ``None`` to skip the title.
            max_level: Maximum heading level to include in the TOC field
                switch (only meaningful when ``use_word_field=True``).
            use_word_field: If True, insert a Word TOC field instead of
                static paragraphs.
        """
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        from docx.oxml import OxmlElement
        from docx.shared import Cm, Pt

        # 1. Optional title
        if title:
            self.document.add_paragraph(title, style='TOC Heading')

        if use_word_field:
            # ── Mode B: native Word TOC field ──
            para = self.document.add_paragraph()
            run = para.add_run()

            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')
            run._element.append(fld_char_begin)

            run2 = para.add_run()
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
            run2._element.append(instr)

            run3 = para.add_run()
            fld_char_sep = OxmlElement('w:fldChar')
            fld_char_sep.set(qn('w:fldCharType'), 'separate')
            run3._element.append(fld_char_sep)

            run4 = para.add_run('（请在 Word 中按 Ctrl+A → F9 刷新目录）')
            run4.font.size = Pt(9)
            run4.font.color.rgb = None  # inherit

            run5 = para.add_run()
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')
            run5._element.append(fld_char_end)

            return

        # ── Mode A: static TOC ──
        if not entries:
            self.document.add_paragraph('（无目录条目）')
            return

        # Add a blank paragraph after title for spacing
        if title:
            self.document.add_paragraph('')

        for num_str, heading_text, level in entries:
            if level > max_level:
                continue

            # Use TOC 1/2/3 style (clamp to 1-9 → 1-3, Word only ships 3 levels)
            toc_style_num = min(level, 3)
            style_name = f'TOC {toc_style_num}'

            # Build display text: "编号  标题"
            if num_str:
                display = f'{num_str}    {heading_text}'
            else:
                display = heading_text

            para = self.document.add_paragraph(display, style=style_name)

    def save(self, path: str):
        """Save the document to a file.

        Args:
            path: Output .docx path
        """
        os.makedirs(os.path.dirname(os.path.abspath(path)) or '.', exist_ok=True)
        self.document.save(path)
        logger.info('Document saved to %s', path)

    def _render_mermaid_elements(
        self, elements: List[MarkdownElement]
    ) -> List[MarkdownElement]:
        """Render mermaid elements to images in-place.

        Args:
            elements: List of MarkdownElement objects

        Returns:
            Modified elements list with mermaid blocks replaced by images
            or code blocks (fallback)
        """
        rendered = []
        for elem in elements:
            if elem.type == MarkdownElementType.MERMAID:
                if self.mermaid_renderer.is_available() and elem.text:
                    try:
                        img_path = self.mermaid_renderer.render(elem.text)
                        rendered.append(MarkdownElement(
                            type=MarkdownElementType.IMAGE,
                            image_path=img_path,
                            alt_text='Mermaid diagram',
                        ))
                        continue
                    except Exception:
                        logger.warning(
                            'Mermaid rendering failed, falling back to code block'
                        )
                rendered.append(MarkdownElement(
                    type=MarkdownElementType.CODE_BLOCK,
                    text=elem.text or '',
                    code_language='mermaid',
                ))
            else:
                rendered.append(elem)
        return rendered

    @staticmethod
    def create(
        output_path: str,
        md_text: str,
        format_spec: Optional[Dict[str, Any]] = None,
        template_path: Optional[str] = None,
    ) -> 'DocxCreator':
        """One-shot: create a new DOCX from markdown and save it.

        Args:
            output_path: Where to write the .docx file
            md_text: Markdown content
            format_spec: Optional default format (see ``set_default_format``)
            template_path: Optional style template path

        Returns:
            The DocxCreator instance (for further manipulation if needed)
        """
        creator = DocxCreator(template_path=template_path)
        if format_spec:
            creator.set_default_format(format_spec)
        creator.add_markdown(md_text)
        creator.save(output_path)
        return creator
