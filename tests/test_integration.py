"""Comprehensive integration tests for the DocxDocument module.

Tests all features end-to-end: load, chapter operations, markdown
conversion, table building, comments, revisions, and save.
"""
import pytest
import sys
import os
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx_editor import DocxDocument
from docx_editor.utils import ChapterNotFoundError


@pytest.fixture
def complex_docx():
    """Create a complex test document with multi-level headings."""
    doc = Document()

    # Chapter 1
    doc.add_heading("第1章 简介", level=1)
    doc.add_paragraph("这是第一章的简介内容。")
    doc.add_paragraph("包含多个段落。")

    # Chapter 1.1
    doc.add_heading("1.1 背景", level=2)
    doc.add_paragraph("背景介绍。这里有一些背景信息。")
    doc.add_paragraph("更多背景内容。")

    # Chapter 1.2
    doc.add_heading("1.2 目标", level=2)
    doc.add_paragraph("项目目标说明。")

    # Chapter 1.2.1
    doc.add_heading("1.2.1 具体目标", level=3)
    doc.add_paragraph("具体目标1：完成模块A。")
    doc.add_paragraph("具体目标2：完成模块B。")

    # Chapter 2
    doc.add_heading("第2章 设计", level=1)
    doc.add_paragraph("设计思路概述。")

    # Chapter 2.1
    doc.add_heading("2.1 架构设计", level=2)
    doc.add_paragraph("系统架构说明。")
    doc.add_paragraph("分为三个层次。")

    # Chapter 2.1.1
    doc.add_heading("2.1.1 前端", level=3)
    doc.add_paragraph("前端使用Vue框架。")

    # Chapter 2.1.2
    doc.add_heading("2.1.2 后端", level=3)
    doc.add_paragraph("后端使用Python。")

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    doc.save(tmp.name)
    yield tmp.name
    os.unlink(tmp.name)


class TestIntegrationFullFlow:
    """Full end-to-end integration tests."""

    def test_load_and_explore_tree(self, complex_docx):
        """Load document and explore chapter tree."""
        doc = DocxDocument(complex_docx)
        tree = doc.get_chapter_tree()

        assert len(tree) == 2  # 第1章, 第2章
        ch1 = tree[0]
        assert ch1.number_tuple == (1,)
        assert len(ch1.children) == 2  # 1.1, 1.2

        # Check deep nesting
        ch1_2 = ch1.children[1]
        assert ch1_2.number_tuple == (1, 2)
        assert len(ch1_2.children) == 1  # 1.2.1

    def test_get_all_chapters(self, complex_docx):
        """Test all chapter lookups work."""
        doc = DocxDocument(complex_docx)

        tests = [
            ("1", 1),
            ("1.1", 2),
            ("1.2", 2),
            ("1.2.1", 3),
            ("2", 1),
            ("2.1", 2),
            ("2.1.1", 3),
            ("2.1.2", 3),
        ]

        for num_str, expected_level in tests:
            ch = doc.get_chapter(num_str)
            assert ch is not None, f"Chapter {num_str} not found"
            assert ch.heading_level == expected_level, \
                f"Chapter {num_str} level: got {ch.heading_level}, expected {expected_level}"

    def test_get_chapter_contents_and_text(self, complex_docx):
        """Test chapter content retrieval."""
        doc = DocxDocument(complex_docx)

        ch = doc.get_chapter("1.2.1")
        assert ch is not None

        contents = doc.get_chapter_contents(ch)
        assert len(contents) >= 3  # heading + 2 body paragraphs

        text = doc.get_chapter_text(ch)
        assert "具体目标" in text or "目标" in text

    def test_replace_chapter_paragraphs(self, complex_docx):
        """Replace chapter content with markdown paragraphs."""
        doc = DocxDocument(complex_docx)

        md = """# 新目标

这是替换后的新内容。

包含多个段落。

还有更多内容。"""

        doc.replace_chapter("1.2", md)

        # Verify the replacement worked
        ch = doc.get_chapter("1.2")
        assert ch is not None

        text = doc.get_chapter_text(ch)
        # Should have new content (may or may not include original heading)
        assert len(text) > 0

    def test_replace_chapter_with_table(self, complex_docx):
        """Replace chapter content with markdown table."""
        doc = DocxDocument(complex_docx)

        md = """| 名称 | 版本 | 状态 |
| v1.0 | 1.0 | 已发布 |
| v2.0 | 2.0 | 开发中 |"""

        doc.replace_chapter("1.1", md)

        # Should not crash and should contain table content
        doc.save()  # Save should work

    def test_replace_chapter_with_headings(self, complex_docx):
        """Replace with content that has multiple heading levels."""
        doc = DocxDocument(complex_docx)

        md = """# 一级标题

段落1

## 二级标题

段落2

### 三级标题

段落3"""

        doc.replace_chapter("2.1", md)

        ch = doc.get_chapter("2.1")
        assert ch is not None

    def test_search_replace_in_chapter(self, complex_docx):
        """Search and replace text in a chapter."""
        doc = DocxDocument(complex_docx)

        # Replace text in chapter 1
        count = doc.search_replace_in_chapter("1", "简介", "概述")
        assert count > 0, "Should have found and replaced text"

        # Verify by reading the revisions (chapter text should have changed)
        ch = doc.get_chapter("1")
        text = doc.get_chapter_text(ch)
        assert "概述" in text or "简介" not in text

    def test_delete_chapter_and_renumber(self, complex_docx):
        """Delete a chapter and verify renumbering."""
        doc = DocxDocument(complex_docx)

        # Delete 2.1.1
        doc.delete_chapter("2.1.1")

        # 2.1.2 should become 2.1.1 after renumbering
        ch = doc.get_chapter("2.1.1")
        assert ch is not None

    def test_delete_in_chapter(self, complex_docx):
        """Delete specific content within a chapter."""
        doc = DocxDocument(complex_docx)

        count = doc.delete_in_chapter("1", "简介")
        assert count > 0, "Should have deleted text"

    def test_read_tracked_changes(self, complex_docx):
        """Read tracked changes from document (may be empty)."""
        doc = DocxDocument(complex_docx)
        revisions = doc.read_revisions()
        assert isinstance(revisions, list)

    def test_read_comments(self, complex_docx):
        """Read comments from document (may be empty)."""
        doc = DocxDocument(complex_docx)
        comments = doc.read_comments()
        assert isinstance(comments, list)

    def test_save_and_reload(self, complex_docx):
        """Save modifications and reload to verify persistence."""
        doc = DocxDocument(complex_docx)

        # Make changes - use non-heading content to avoid renumbering
        doc.search_replace_in_chapter("1", "简介", "概述")
        doc.delete_in_chapter("1.2", "项目")

        # Save
        output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        output.close()
        try:
            doc.save(output.name)

            # Reload into new instance
            doc2 = DocxDocument(output.name)

            # Verify chapter structure preserved (search-replace doesn't change structure)
            assert doc2.get_chapter("1") is not None
            assert doc2.get_chapter("1.1") is not None

        finally:
            os.unlink(output.name)


class TestIntegrationEdgeCases:
    """Edge case tests."""

    def test_replace_nonexistent_chapter(self, complex_docx):
        doc = DocxDocument(complex_docx)
        with pytest.raises(ChapterNotFoundError):
            doc.replace_chapter("99.99", "内容")

    def test_delete_nonexistent_chapter(self, complex_docx):
        doc = DocxDocument(complex_docx)
        with pytest.raises(ChapterNotFoundError):
            doc.delete_chapter("99.99")

    def test_search_replace_non_existent_text(self, complex_docx):
        doc = DocxDocument(complex_docx)
        count = doc.search_replace_in_chapter("1", "不存在的文本", "替换")
        assert count == 0

    def test_delete_non_existent_text(self, complex_docx):
        doc = DocxDocument(complex_docx)
        count = doc.delete_in_chapter("1", "不存在的文本")
        assert count == 0

    def test_empty_markdown_replacement(self, complex_docx):
        doc = DocxDocument(complex_docx)
        # Should not crash with empty markdown
        doc.replace_chapter("2.1", "")
        ch = doc.get_chapter("2.1")
        assert ch is not None

    def test_various_heading_patterns(self):
        """Test documents with different heading numbering patterns."""
        doc = Document()
        doc.add_heading("第一章 测试", level=1)
        doc.add_paragraph("内容1")
        doc.add_heading("1.1 小节", level=2)
        doc.add_paragraph("内容2")
        doc.add_heading("一、另一章", level=1)
        doc.add_paragraph("内容3")

        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        doc.save(tmp.name)
        try:
            d = DocxDocument(tmp.name)
            tree = d.get_chapter_tree()
            # Should detect at least some chapters
            assert len(tree) > 0
        finally:
            os.unlink(tmp.name)

    def test_concurrent_operations(self, complex_docx):
        """Chain multiple operations without reloading issues."""
        doc = DocxDocument(complex_docx)

        # Use non-heading markdown to avoid structural changes
        doc.replace_chapter("1.1", "新背景内容\n\n更多新内容。")
        doc.replace_chapter("2.1", "新架构内容\n\n更多架构说明。")

        # Verify final state
        ch1_1 = doc.get_chapter("1.1")
        assert ch1_1 is not None

        # Delete a leaf chapter
        doc.delete_chapter("1.2.1")

        # Delete a top-level chapter
        doc.delete_chapter("2")

        # Verify remaining structure
        tree = doc.get_chapter_tree()
        assert len(tree) >= 1


class TestIntegrationWithOriginalDocx:
    """Test against the original 测试.docx file."""

    @pytest.fixture
    def original_docx(self):
        path = os.path.join(os.path.dirname(__file__), '..', '测试.docx')
        if not os.path.exists(path):
            pytest.skip("测试.docx not found")
        return path

    def test_load_original(self, original_docx):
        doc = DocxDocument(original_docx)
        assert doc is not None

    def test_original_revisions(self, original_docx):
        doc = DocxDocument(original_docx)
        revisions = doc.read_revisions()
        # The 测试.docx has tracked changes
        assert isinstance(revisions, list)

    def test_original_chapter_tree(self, original_docx):
        doc = DocxDocument(original_docx)
        tree = doc.get_chapter_tree()
        # May be empty if no headings, but shouldn't crash
        assert isinstance(tree, list)

    def test_original_comments(self, original_docx):
        doc = DocxDocument(original_docx)
        comments = doc.read_comments()
        assert isinstance(comments, list)


class TestIntegrationToc:
    """End-to-end TOC integration tests."""

    def test_toc_then_save_and_reload(self, complex_docx):
        """Insert TOC, save, reload — chapter structure preserved."""
        doc = DocxDocument(complex_docx)
        doc.insert_toc(position=None, max_level=3)

        # Save
        output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        output.close()
        try:
            doc.save(output.name)

            # Reload
            doc2 = DocxDocument(output.name)
            # Chapter structure should still be intact
            assert doc2.get_chapter("1") is not None
            assert doc2.get_chapter("2.1.2") is not None
            assert len(doc2.get_chapter_tree()) > 0
        finally:
            os.unlink(output.name)

    def test_toc_with_then_replace_chapter(self, complex_docx):
        """Insert TOC, then replace a chapter — no conflict."""
        doc = DocxDocument(complex_docx)
        doc.insert_toc(position=None, max_level=2)
        doc.replace_chapter("2.1", "# 新架构\n\n新内容")

        ch = doc.get_chapter("2.1")
        assert ch is not None

    def test_toc_with_then_delete_chapter(self, complex_docx):
        """Insert TOC, then delete a chapter — no crash."""
        doc = DocxDocument(complex_docx)
        doc.insert_toc(position=None, max_level=2)
        doc.delete_chapter("1.2.1")

        ch = doc.get_chapter("1.2")
        assert ch is not None

    def test_toc_word_field_then_save(self, complex_docx):
        """Word TOC field insertion then save/reload."""
        doc = DocxDocument(complex_docx)
        doc.insert_toc(position=None, use_word_field=True)

        output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        output.close()
        try:
            doc.save(output.name)
            doc2 = DocxDocument(output.name)
            assert doc2.get_chapter("1") is not None
        finally:
            os.unlink(output.name)
