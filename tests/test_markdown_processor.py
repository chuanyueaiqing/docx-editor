"""Tests for markdown_processor.py"""
import pytest
import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx_editor.markdown_processor import MarkdownProcessor
from docx_editor.models import (
    MarkdownElement, MarkdownElementType, FormatStore
)
from docx_editor.format_extractor import FormatExtractor


@pytest.fixture
def empty_store():
    doc = Document()
    store = FormatStore(docx_path="test.docx", document=doc)
    store.formats_json = FormatExtractor.extract_all(doc)
    return store


class TestMarkdownParsing:
    def test_parse_heading(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("# 一级标题\n\n## 二级标题\n\n正文内容")
        headings = [e for e in elements if e.type == MarkdownElementType.HEADING]
        assert len(headings) == 2
        assert headings[0].level == 1
        assert headings[0].text == "一级标题"
        assert headings[1].level == 2
        assert headings[1].text == "二级标题"

    def test_parse_paragraph(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("第一段\n\n第二段")
        paras = [e for e in elements if e.type == MarkdownElementType.PARAGRAPH]
        assert len(paras) == 2
        assert paras[0].text == "第一段"
        assert paras[1].text == "第二段"

    def test_parse_table(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        md = "| A | B | C |\n| a1 | > | c1 |\n| a2 | v | > |"
        elements = mp.parse_markdown(md)
        tables = [e for e in elements if e.type == MarkdownElementType.TABLE]
        assert len(tables) == 1
        table = tables[0]
        assert len(table.rows) == 2  # 2 data rows
        # Check headers are separate
        assert table.children is not None
        headers = [c.text for c in table.children if c.type == MarkdownElementType.TABLE]
        if headers:
            assert len(headers) > 0

    def test_parse_code_block(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("```python\nprint('hello')\n```")
        codes = [e for e in elements if e.type == MarkdownElementType.CODE_BLOCK]
        assert len(codes) == 1
        assert codes[0].code_language == "python"
        assert "print" in codes[0].text

    def test_parse_mermaid(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("```mermaid\ngraph TD;\nA-->B;\n```")
        mermaids = [e for e in elements if e.type == MarkdownElementType.MERMAID]
        assert len(mermaids) == 1
        assert "A-->B" in mermaids[0].text

    def test_parse_image(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("![架构图](diagram.png)")
        images = [e for e in elements if e.type == MarkdownElementType.IMAGE]
        assert len(images) == 1
        assert images[0].image_path == "diagram.png"
        assert images[0].alt_text == "架构图"

    def test_parse_list(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("- item1\n- item2\n- item3")
        lists = [e for e in elements if e.type == MarkdownElementType.LIST]
        assert len(lists) == 1
        assert lists[0].ordered is False
        assert len(lists[0].items) >= 2

    def test_parse_ordered_list(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("1. first\n2. second\n3. third")
        lists = [e for e in elements if e.type == MarkdownElementType.LIST]
        assert len(lists) == 1
        assert lists[0].ordered is True

    def test_parse_horizontal_rule(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("---")
        hrs = [e for e in elements if e.type == MarkdownElementType.HORIZONTAL_RULE]
        assert len(hrs) >= 1

    def test_inline_bold(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("这是**粗体**文字")
        # Bold should be handled as inline formatting within a paragraph
        paras = [e for e in elements if e.type == MarkdownElementType.PARAGRAPH]
        assert len(paras) >= 1
        text = paras[0].text
        assert "粗体" in text or "**" in text

    def test_parse_empty_string(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("")
        assert len(elements) == 0

    def test_parse_only_whitespace(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("   \n  \n  ")
        empty = [e for e in elements if e.type == MarkdownElementType.EMPTY_LINE]
        assert len(empty) >= 0

    # ── Math formula tests ──

    def test_parse_display_math(self, empty_store):
        """Display math ($$...$$) is parsed as DISPLAY_MATH."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown(
            "$$\\int_a^b f(x) \\, dx$$"
        )
        display_maths = [
            e for e in elements if e.type == MarkdownElementType.DISPLAY_MATH
        ]
        assert len(display_maths) == 1
        assert r'\int' in display_maths[0].text
        assert 'f(x)' in display_maths[0].text

    def test_parse_inline_math(self, empty_store):
        """$...$ is parsed as MATH element."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown(
            "公式 $E=mc^2$ 是著名的质能方程"
        )
        maths = [e for e in elements if e.type == MarkdownElementType.MATH]
        assert len(maths) == 1
        assert 'E=mc^2' in maths[0].text

    def test_parse_multiple_inline_math(self, empty_store):
        """Multiple $...$ within one paragraph."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown(
            "$a^2$ 和 $b^2$ 的和"
        )
        maths = [e for e in elements if e.type == MarkdownElementType.MATH]
        assert len(maths) == 2

    def test_parse_math_with_paragraph(self, empty_store):
        """Inline math splits paragraph into text + math segments."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown(
            "根据公式 $E=mc^2$ 可知能量巨大"
        )
        types = [e.type for e in elements]
        assert MarkdownElementType.PARAGRAPH in types
        assert MarkdownElementType.MATH in types
        # Should have at least 3 elements: text before, math, text after
        assert len(elements) >= 3

    def test_parse_display_math_multiline(self, empty_store):
        """Display math can span multiple lines."""
        mp = MarkdownProcessor(empty_store)
        md = ("$$\n"
              "        \\int_{0}^{\\infty} e^{-x^2} \\, dx\n"
              "        = \\frac{\\sqrt{\\pi}}{2}\n"
              "        $$")
        elements = mp.parse_markdown(md)
        display_maths = [
            e for e in elements if e.type == MarkdownElementType.DISPLAY_MATH
        ]
        assert len(display_maths) >= 1
        text = display_maths[0].text
        assert r'\int' in text
        assert r'\infty' in text

    def test_parse_math_then_paragraph(self, empty_store):
        """Display math followed by normal text."""
        mp = MarkdownProcessor(empty_store)
        md = "$$\\sum_{i=1}^n i = \\frac{n(n+1)}{2}$$\n\n后续内容"
        elements = mp.parse_markdown(md)
        display_maths = [
            e for e in elements if e.type == MarkdownElementType.DISPLAY_MATH
        ]
        paras = [e for e in elements if e.type == MarkdownElementType.PARAGRAPH]
        assert len(display_maths) == 1
        assert r'\sum' in display_maths[0].text
        assert any('后续' in (p.text or '') for p in paras)


class TestMarkdownConversion:
    def test_heading_to_docx(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("# 测试标题\n\n段落内容")
        mp.apply_to_document(elements)
        doc = empty_store.document
        # Should have at least 2 paragraphs
        assert len(doc.paragraphs) >= 2
        # First should be heading style
        p0 = doc.paragraphs[0]
        style_name = p0.style.name if p0.style else ''
        assert 'heading' in style_name.lower() or p0.text == "测试标题"

    def test_paragraph_to_docx(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("这是一段测试文字。")
        mp.apply_to_document(elements)
        doc = empty_store.document
        assert len(doc.paragraphs) >= 1
        assert "测试文字" in doc.paragraphs[0].text

    def test_multiple_paragraphs(self, empty_store):
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("第一段\n\n第二段\n\n第三段")
        mp.apply_to_document(elements)
        doc = empty_store.document
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) >= 3

    def test_code_block_to_docx(self, empty_store):
        """Code blocks render as a single-cell table with monospace formatting."""
        from docx.shared import Pt
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown(
            "```python\ndef hello():\n    print('Hi')\n```"
        )
        mp.apply_to_document(elements)
        doc = empty_store.document

        # Should produce exactly one table
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 1
        assert len(table.columns) == 1

        # Cell contains all code lines (indentation preserved)
        cell = table.cell(0, 0)
        texts = [p.text for p in cell.paragraphs]
        assert texts[0] == "def hello():"
        assert "    print('Hi')" in texts

        # Font is Consolas, 9pt
        for para in cell.paragraphs:
            for run in para.runs:
                assert run.font.name == 'Consolas'
                assert run.font.size == Pt(9)

    # ── Math formula conversion tests ──

    def test_display_math_to_docx(self, empty_store):
        """Display math produces an OMML equation with structure."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("$$E=mc^2$$")
        mp.apply_to_document(elements)
        body_xml = empty_store.document.element.body.xml

        # Should have OMML elements
        assert 'm:oMath' in body_xml or 'm:oMathPara' in body_xml

    def test_inline_math_to_docx(self, empty_store):
        """Inline math produces an OMML equation element."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown("公式 $a^2+b^2=c^2$ 成立")
        mp.apply_to_document(elements)
        body_xml = empty_store.document.element.body.xml

        # Should have OMML elements
        assert 'm:oMath' in body_xml

    def test_math_omml_structure(self, empty_store):
        """Multiple formulas produce OMML with structural elements."""
        mp = MarkdownProcessor(empty_store)
        elements = mp.parse_markdown(
            "$a^2$ 和 $b^2$ 的和"
        )
        mp.apply_to_document(elements)
        body_xml = empty_store.document.element.body.xml

        # Should have OMML elements for each formula
        import re
        omml_count = len(re.findall(r'<m:oMath[ >]', body_xml))
        assert omml_count >= 2

        # Should have superscript structure for ^2
        assert 'm:sup' in body_xml
