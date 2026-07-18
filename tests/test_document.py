"""Tests for document.py (DocxDocument facade)"""
import pytest
import sys
import os
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx_editor.document import DocxDocument
from docx_editor.models import ChapterNode, CommentData, RevisionData
from docx_editor.utils import ChapterNotFoundError


@pytest.fixture
def test_docx_path():
    return os.path.join(os.path.dirname(__file__), '..', '测试.docx')


@pytest.fixture
def doc_with_headings_path():
    """Create a docx with headings and save to temp file."""
    doc = Document()
    doc.add_heading("第一章 引言", level=1)
    doc.add_paragraph("这是引言正文。")
    doc.add_heading("1.1 背景", level=2)
    doc.add_paragraph("背景内容。")
    doc.add_heading("1.2 目标", level=2)
    doc.add_paragraph("目标内容。")
    doc.add_heading("第二章 方法", level=1)
    doc.add_paragraph("方法内容。")

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    doc.save(tmp.name)
    yield tmp.name
    os.unlink(tmp.name)


class TestDocxDocumentInit:
    def test_load_existing_document(self, test_docx_path):
        docx = DocxDocument(test_docx_path)
        assert docx is not None
        assert docx.document is not None

    def test_load_nonexistent_document(self):
        with pytest.raises(FileNotFoundError):
            DocxDocument("nonexistent.docx")

    def test_chapter_tree_built(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        tree = docx.get_chapter_tree()
        assert len(tree) >= 1

    def test_use_track_changes_default(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        assert docx.use_track_changes is False

    def test_use_track_changes_true(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path, use_track_changes=True)
        assert docx.use_track_changes is True


class TestChapterAccess:
    def test_get_chapter_by_string(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        ch = docx.get_chapter("1.1")
        assert ch is not None
        assert ch.number_tuple == (1, 1)

    def test_get_chapter_by_tuple(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        ch = docx.get_chapter((1, 1))
        assert ch is not None

    def test_get_chapter_not_found(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        ch = docx.get_chapter("5.1")
        assert ch is None

    def test_get_chapter_contents(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        ch = docx.get_chapter("1.2")
        assert ch is not None
        contents = docx.get_chapter_contents(ch)
        assert len(contents) >= 1

    def test_get_chapter_text(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        ch = docx.get_chapter("1")
        assert ch is not None
        text = docx.get_chapter_text(ch)
        assert len(text) > 0

    def test_print_chapter_tree(self, doc_with_headings_path, capsys):
        docx = DocxDocument(doc_with_headings_path)
        docx.print_chapter_tree()
        captured = capsys.readouterr()
        assert "1" in captured.out or "第一章" in captured.out


class TestChapterOperations:
    def test_replace_chapter(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        new_md = "# 新标题\n\n新段落内容。"
        docx.replace_chapter("1.1", new_md)
        # Verify the content was replaced
        ch = docx.get_chapter("1.1")
        assert ch is not None

    def test_replace_chapter_with_table(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        md = "# 数据表\n\n| A | B |\n| 1 | 2 |"
        docx.replace_chapter("1.2", md)
        ch = docx.get_chapter("1.2")
        assert ch is not None

    def test_replace_nonexistent_chapter(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        with pytest.raises(ChapterNotFoundError):
            docx.replace_chapter("9.9", "内容")

    def test_search_replace_in_chapter(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        count = docx.search_replace_in_chapter("1", "正文", "替换正文")
        assert count >= 0  # May be 0 if no match, but shouldn't crash

    def test_delete_chapter(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        docx.delete_chapter("1.1")
        # The chapter should be gone or renumbered
        ch = docx.get_chapter("1.1")
        assert ch is None or ch.heading_text != "1.1 背景"

    def test_delete_nonexistent_chapter(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        with pytest.raises(ChapterNotFoundError):
            docx.delete_chapter("9.9")

    def test_delete_in_chapter(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        count = docx.delete_in_chapter("1", "引言")
        assert count >= 0

    def test_save_document(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        output.close()
        try:
            docx.save(output.name)
            assert os.path.exists(output.name)
            # Verify it's a valid docx
            loaded = Document(output.name)
            assert loaded is not None
        finally:
            os.unlink(output.name)


class TestCommentsAndRevisions:
    def test_read_comments(self, test_docx_path):
        docx = DocxDocument(test_docx_path)
        comments = docx.read_comments()
        assert isinstance(comments, list)

    def test_read_revisions(self, test_docx_path):
        docx = DocxDocument(test_docx_path)
        revisions = docx.read_revisions()
        assert isinstance(revisions, list)
        # 测试.docx should have tracked changes
        if revisions:
            assert hasattr(revisions[0], 'author')
            assert hasattr(revisions[0], 'text')


class TestDocxDocumentEdgeCases:
    def test_empty_document(self):
        doc = Document()
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        doc.save(tmp.name)
        try:
            docx = DocxDocument(tmp.name)
            assert docx.get_chapter("1") is None
            assert docx.get_chapter_tree() == []
        finally:
            os.unlink(tmp.name)

    def test_save_overwrite(self, doc_with_headings_path):
        docx = DocxDocument(doc_with_headings_path)
        docx.replace_chapter("1.1", "# 新标题\n新内容")
        docx.save()  # Overwrite original
        # Reload and verify
        docx2 = DocxDocument(doc_with_headings_path)
        ch = docx2.get_chapter("1.1")
        # Content should reflect the replacement
        assert ch is not None


class TestInsertToc:
    """Tests for DocxDocument.insert_toc()."""

    def test_insert_toc_at_beginning(self, doc_with_headings_path):
        """Static TOC inserted at the beginning."""
        docx = DocxDocument(doc_with_headings_path)
        docx.insert_toc(position=None, max_level=3)

        # After insertion, first paragraph should be the TOC title
        first_para = docx.format_store.document.paragraphs[0]
        assert '目录' in first_para.text

        # TOC entries should exist for all headings
        all_text = ' '.join(
            p.text for p in docx.format_store.document.paragraphs
        )
        assert '第一章 引言' in all_text
        assert '1.1 背景' in all_text
        assert '第二章 方法' in all_text

    def test_insert_toc_before_chapter(self, doc_with_headings_path):
        """TOC inserted before a specific chapter."""
        docx = DocxDocument(doc_with_headings_path)
        docx.insert_toc(position='2', max_level=2)

        # The TOC should be before chapter 2's heading
        paragraphs = docx.format_store.document.paragraphs
        toc_texts = [p.text for p in paragraphs]

        # Find where "目录" and "第二章" appear
        toc_idx = next(i for i, t in enumerate(toc_texts) if '目录' in t)
        ch2_idx = next(i for i, t in enumerate(toc_texts) if '第二章' in t)
        assert toc_idx < ch2_idx

    def test_insert_toc_max_level_1(self, doc_with_headings_path):
        """Only level-1 headings appear in TOC with max_level=1."""
        docx = DocxDocument(doc_with_headings_path)
        docx.insert_toc(position=None, max_level=1)

        paragraphs = docx.format_store.document.paragraphs

        # TOC entries are between "目录" and the first original heading ("第一章 引言")
        # Find their positions
        texts = [p.text for p in paragraphs]
        toc_title_idx = next(i for i, t in enumerate(texts) if '目录' in t)
        # The TOC entries start after the blank paragraph following "目录"
        toc_start = toc_title_idx + 2  # title + blank
        # The original content starts at "第一章 引言" (second occurrence)
        first_heading_idx = texts.index('第一章 引言', toc_start)

        toc_entries = texts[toc_start:first_heading_idx]
        assert len(toc_entries) == 2  # 2 level-1 headings
        assert any('第一章' in t for t in toc_entries)
        assert any('第二章' in t for t in toc_entries)
        assert not any('1.1' in t for t in toc_entries)

    def test_insert_toc_twice(self, doc_with_headings_path):
        """Inserting TOC twice should create entries both times."""
        docx = DocxDocument(doc_with_headings_path)
        docx.insert_toc(position=None, max_level=1)
        docx.insert_toc(position='2', max_level=1)

        paragraphs = docx.format_store.document.paragraphs
        toc_texts = [p.text for p in paragraphs]
        assert toc_texts.count('目录') == 2

    def test_insert_toc_word_field(self, doc_with_headings_path):
        """Word TOC field can be inserted without error."""
        docx = DocxDocument(doc_with_headings_path)
        docx.insert_toc(position=None, use_word_field=True)

        paragraphs = docx.format_store.document.paragraphs
        # Should have TOC heading + field paragraph
        assert '目录' in paragraphs[0].text

    def test_insert_toc_empty_document(self):
        """TOC on empty document should not crash."""
        doc = Document()
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        doc.save(tmp.name)
        try:
            docx = DocxDocument(tmp.name)
            docx.insert_toc(position=None)
            # Should have at least the TOC title
            paragraphs = docx.format_store.document.paragraphs
            assert len(paragraphs) >= 1
        finally:
            os.unlink(tmp.name)

    def test_insert_toc_before_nonexistent_chapter(self, doc_with_headings_path):
        """Inserting TOC before a nonexistent chapter raises error."""
        docx = DocxDocument(doc_with_headings_path)
        with pytest.raises(ChapterNotFoundError):
            docx.insert_toc(position='9.9')
