"""Tests for chapter_parser.py"""
import pytest
import sys
import os
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx_editor.chapter_parser import ChapterParser
from docx_editor.format_extractor import FormatExtractor
from docx_editor.models import FormatStore, ChapterNode


@pytest.fixture
def doc_with_headings():
    """Create a docx with multi-level headings for testing."""
    doc = Document()

    # Chapter 1
    doc.add_heading("第一章 引言", level=1)
    doc.add_paragraph("这是引言的正文内容。这里介绍项目背景。")
    doc.add_paragraph("第二段引言正文。")

    # Chapter 1.1
    doc.add_heading("1.1 项目背景", level=2)
    doc.add_paragraph("项目背景描述。")
    doc.add_paragraph("更多背景信息。")

    # Chapter 1.2
    doc.add_heading("1.2 项目范围", level=2)
    doc.add_paragraph("项目范围说明。")

    # Chapter 1.2.1
    doc.add_heading("1.2.1 功能范围", level=3)
    doc.add_paragraph("功能范围描述。")

    # Chapter 2
    doc.add_heading("第二章 系统架构", level=1)
    doc.add_paragraph("系统架构概述。")

    # Chapter 2.1
    doc.add_heading("2.1 总体架构", level=2)
    doc.add_paragraph("总体架构描述。")

    # Chapter 2.1.1
    doc.add_heading("2.1.1 前端架构", level=3)
    doc.add_paragraph("前端架构描述。")

    # Chapter 2.1.2
    doc.add_heading("2.1.2 后端架构", level=3)
    doc.add_paragraph("后端架构描述。")

    # No heading paragraph at the end (trailing content)
    doc.add_paragraph("文档结尾内容。")

    return doc


@pytest.fixture
def parser(doc_with_headings):
    """Create ChapterParser from the test document."""
    store = FormatStore(docx_path="test.docx", document=doc_with_headings)
    store.formats_json = FormatExtractor.extract_all(doc_with_headings)
    return ChapterParser(store)


class TestChapterTreeBuilding:
    def test_root_chapters_count(self, parser):
        assert len(parser.root_chapters) == 2  # 第一章, 第二章

    def test_chapter_structure(self, parser):
        ch1 = parser.root_chapters[0]
        assert ch1.heading_level == 1
        assert ch1.number_tuple == (1,)
        assert len(ch1.children) == 2  # 1.1, 1.2

    def test_nested_chapters(self, parser):
        ch1 = parser.root_chapters[0]
        ch1_2 = ch1.children[1]  # 1.2
        assert ch1_2.number_tuple == (1, 2)
        assert len(ch1_2.children) == 1  # 1.2.1

    def test_deeply_nested(self, parser):
        ch2 = parser.root_chapters[1]
        ch2_1 = ch2.children[0]  # 2.1
        assert ch2_1.number_tuple == (2, 1)
        assert len(ch2_1.children) == 2  # 2.1.1, 2.1.2

    def test_body_paragraph_indices(self, parser):
        """Body paragraphs should include paragraphs between headings."""
        ch1 = parser.root_chapters[0]
        # Chapter 1 has heading at index 0, body at 1,2
        assert len(ch1.body_paragraph_indices) >= 2

    def test_heading_not_in_body(self, parser):
        """Heading index should not appear in body indices."""
        ch1 = parser.root_chapters[0]
        assert ch1.heading_paragraph_index not in ch1.body_paragraph_indices

    def test_children_parent_link(self, parser):
        ch1 = parser.root_chapters[0]
        ch1_1 = ch1.children[0]
        assert ch1_1.parent is ch1


class TestLookup:
    def test_lookup_by_string(self, parser):
        ch = parser.get_chapter_by_number("1.1")
        assert ch is not None
        assert ch.number_tuple == (1, 1)
        assert "项目背景" in ch.heading_text

    def test_lookup_by_tuple(self, parser):
        ch = parser.get_chapter_by_number((2, 1))
        assert ch is not None
        assert ch.number_tuple == (2, 1)

    def test_lookup_deep_nested(self, parser):
        ch = parser.get_chapter_by_number("1.2.1")
        assert ch is not None
        assert ch.number_tuple == (1, 2, 1)
        assert "功能范围" in ch.heading_text or "范围" in ch.heading_text

    def test_lookup_not_found(self, parser):
        ch = parser.get_chapter_by_number("5.1")
        assert ch is None

    def test_lookup_top_level(self, parser):
        ch = parser.get_chapter_by_number("1")
        assert ch is not None
        assert ch.number_tuple == (1,)
        assert "引言" in ch.heading_text or "一" in ch.heading_text


class TestGetContents:
    def test_get_chapter_text(self, parser):
        ch1 = parser.root_chapters[0]
        text = parser.get_chapter_text(ch1)
        # Should contain heading text and some body text
        assert len(text) > 0
        assert ch1.heading_text in text or ch1.to_string() in text

    def test_get_chapter_paragraphs_count(self, parser):
        ch1 = parser.root_chapters[0]
        contents = parser.get_chapter_contents(ch1)
        assert len(contents) >= 3  # heading + at least 2 body


class TestDeleteChapter:
    def test_delete_leaf_chapter(self, parser):
        ch = parser.get_chapter_by_number("1.1")
        assert ch is not None
        deleted_text = ch.heading_text
        parser.delete_chapter(ch)
        # After deletion + renumbering, check that the original heading text is gone
        for node in parser.list_chapters():
            assert node.heading_text != deleted_text

    def test_delete_with_children(self, parser_with_deep):
        """Deleting a parent chapter should remove children too."""
        pass  # Will implement after basic delete is working


class TestChapterTreeEdgeCases:
    def test_empty_document(self):
        """Document with no headings should have empty tree."""
        doc = Document()
        doc.add_paragraph("No headings here.")
        store = FormatStore(docx_path="test.docx", document=doc)
        store.formats_json = FormatExtractor.extract_all(doc)
        p = ChapterParser(store)
        assert p.root_chapters == []

    def test_single_heading(self):
        """Document with single heading should have one root chapter."""
        doc = Document()
        doc.add_heading("Only Chapter", level=1)
        store = FormatStore(docx_path="test.docx", document=doc)
        store.formats_json = FormatExtractor.extract_all(doc)
        p = ChapterParser(store)
        assert len(p.root_chapters) == 1
        assert p.root_chapters[0].body_paragraph_indices == []

    def test_tree_to_string(self, parser):
        output = parser.tree_to_string()
        assert "1" in output
        assert "1.1" in output
        assert "1.2.1" in output
        assert "2.1.2" in output


@pytest.fixture
def parser_with_deep(doc_with_headings):
    """Create chapter parser with deep heading structure."""
    store = FormatStore(docx_path="test.docx", document=doc_with_headings)
    store.formats_json = FormatExtractor.extract_all(doc_with_headings)
    return ChapterParser(store)
