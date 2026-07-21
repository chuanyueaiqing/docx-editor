"""Markdown to DOCX conversion module.

Parses a subset of markdown (headings, paragraphs, tables, lists, code,
images, horizontal rules) and converts them to python-docx elements.
Supports custom table merge markers:
  - '>' in a cell = horizontal merge (merge with cell to the left)
  - 'v' in a cell = vertical merge (merge with cell above)
"""
import os
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as docx_qn

from .models import (
    FormatStore, MarkdownElement, MarkdownElementType,
    ParagraphData, FormattingData, ParagraphFormatData,
)
from .utils import qn


class MarkdownProcessor:
    """Parse markdown text and convert to python-docx elements.

    Usage:
        processor = MarkdownProcessor(format_store)
        elements = processor.parse_markdown(md_text)
        processor.apply_to_document(elements)
    """

    # Regex patterns for line types
    RE_HEADING = re.compile(r'^(#{1,6})\s+(.+)$')
    RE_FENCED_CODE_START = re.compile(r'^(`{3,}|~{3,})\s*(\w*)$')
    RE_FENCED_CODE_END = re.compile(r'^(`{3,}|~{3,})\s*$')
    RE_TABLE_ROW = re.compile(r'^\|(.+)\|$')
    RE_TABLE_SEPARATOR = re.compile(r'^\|[\s:-]+\|$')
    RE_IMAGE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    RE_LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    RE_ORDERED_LIST = re.compile(r'^\s*(\d+)\.\s+(.+)$')
    RE_UNORDERED_LIST = re.compile(r'^\s*[-*+]\s+(.+)$')
    RE_HORIZONTAL_RULE = re.compile(r'^[-*_]{3,}\s*$')
    RE_BLOCKQUOTE = re.compile(r'^>\s?(.*)$')
    RE_INLINE_BOLD = re.compile(r'\*\*(.+?)\*\*|__(.+?)__')
    RE_INLINE_ITALIC = re.compile(r'\*(.+?)\*|_(.+?)_')
    RE_INLINE_CODE = re.compile(r'`([^`]+)`')
    RE_INLINE_STRIKE = re.compile(r'~~(.+?)~~')

    def __init__(self, format_store: FormatStore):
        self.store = format_store
        self.document = format_store.document

    # ======================== Parsing ========================

    def parse_markdown(self, md_text: str) -> List[MarkdownElement]:
        """Parse markdown text into a list of MarkdownElement objects.

        Args:
            md_text: Markdown text to parse

        Returns:
            List of MarkdownElement objects
        """
        if not md_text or not md_text.strip():
            return []

        lines = md_text.split('\n')
        elements = []
        i = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Empty line
            if not stripped:
                i += 1
                continue

            # Fenced code block (including mermaid)
            m = self.RE_FENCED_CODE_START.match(stripped)
            if m:
                fence_char = m.group(1)
                lang = m.group(2).strip().lower()
                code_lines = []
                i += 1
                while i < len(lines):
                    if self.RE_FENCED_CODE_END.match(lines[i].strip()):
                        i += 1
                        break
                    code_lines.append(lines[i])
                    i += 1

                code_text = '\n'.join(code_lines)

                if lang == 'mermaid':
                    elements.append(MarkdownElement(
                        type=MarkdownElementType.MERMAID,
                        text=code_text,
                        code_language='mermaid',
                    ))
                else:
                    elements.append(MarkdownElement(
                        type=MarkdownElementType.CODE_BLOCK,
                        text=code_text,
                        code_language=lang or None,
                    ))
                continue

            # Table row
            if self.RE_TABLE_ROW.match(stripped):
                table_elem, i = self._parse_table(lines, i)
                if table_elem:
                    elements.append(table_elem)
                continue

            # Horizontal rule
            if self.RE_HORIZONTAL_RULE.match(stripped):
                elements.append(MarkdownElement(
                    type=MarkdownElementType.HORIZONTAL_RULE,
                ))
                i += 1
                continue

            # Heading
            m = self.RE_HEADING.match(stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                elements.append(MarkdownElement(
                    type=MarkdownElementType.HEADING,
                    text=text,
                    level=level,
                ))
                i += 1
                continue

            # Blockquote
            if self.RE_BLOCKQUOTE.match(stripped):
                bq_lines = []
                while i < len(lines):
                    m_bq = self.RE_BLOCKQUOTE.match(lines[i].strip())
                    if not m_bq:
                        break
                    bq_lines.append(m_bq.group(1))
                    i += 1
                elements.append(MarkdownElement(
                    type=MarkdownElementType.BLOCKQUOTE,
                    text='\n'.join(bq_lines).strip(),
                ))
                continue

            # Unordered list
            if self.RE_UNORDERED_LIST.match(stripped):
                items, i = self._parse_unordered_list(lines, i)
                elements.append(MarkdownElement(
                    type=MarkdownElementType.LIST,
                    items=items,
                    ordered=False,
                ))
                continue

            # Ordered list
            if self.RE_ORDERED_LIST.match(stripped):
                items, i = self._parse_ordered_list(lines, i)
                elements.append(MarkdownElement(
                    type=MarkdownElementType.LIST,
                    items=items,
                    ordered=True,
                ))
                continue

            # Image alone on line
            if self.RE_IMAGE.match(stripped) and len(stripped) == len(
                self.RE_IMAGE.match(stripped).group(0)
            ):
                m_img = self.RE_IMAGE.match(stripped)
                elements.append(MarkdownElement(
                    type=MarkdownElementType.IMAGE,
                    alt_text=m_img.group(1),
                    image_path=m_img.group(2),
                ))
                i += 1
                continue

            # Regular paragraph (may span multiple lines)
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    i += 1
                    break
                # Stop at block-level elements
                if (self.RE_HEADING.match(next_line) or
                    self.RE_FENCED_CODE_START.match(next_line) or
                    self.RE_HORIZONTAL_RULE.match(next_line) or
                    self.RE_TABLE_ROW.match(next_line) or
                    self.RE_UNORDERED_LIST.match(next_line) or
                    self.RE_ORDERED_LIST.match(next_line)):
                    break
                para_lines.append(next_line)
                i += 1

            para_text = ' '.join(para_lines)

            # Check if paragraph contains images
            if '![' in para_text:
                # Split by images
                remaining = para_text
                last_end = 0
                for m_img in self.RE_IMAGE.finditer(para_text):
                    # Add text before image
                    before = para_text[last_end:m_img.start()].strip()
                    if before:
                        elements.append(MarkdownElement(
                            type=MarkdownElementType.PARAGRAPH,
                            text=before,
                        ))
                    elements.append(MarkdownElement(
                        type=MarkdownElementType.IMAGE,
                        alt_text=m_img.group(1),
                        image_path=m_img.group(2),
                    ))
                    last_end = m_img.end()
                after = para_text[last_end:].strip()
                if after:
                    elements.append(MarkdownElement(
                        type=MarkdownElementType.PARAGRAPH,
                        text=after,
                    ))
            else:
                elements.append(MarkdownElement(
                    type=MarkdownElementType.PARAGRAPH,
                    text=para_text,
                ))

        return elements

    def _parse_table(self, lines: List[str], start: int
                     ) -> Tuple[Optional[MarkdownElement], int]:
        """Parse a markdown table starting at line start.

        Returns:
            (MarkdownElement or None, new_line_index)
        """
        # Read header row
        header_line = lines[start].strip()
        i = start + 1

        # Check if next line is separator
        if i >= len(lines):
            return None, i

        separator = lines[i].strip()
        has_separator = bool(self.RE_TABLE_SEPARATOR.match(separator))
        if has_separator:
            i += 1

        # Parse header cells
        headers = self._split_table_row(header_line)

        # Parse data rows
        rows = []
        merge_map = []

        while i < len(lines):
            row_line = lines[i].strip()
            if not row_line or not self.RE_TABLE_ROW.match(row_line):
                break

            cells = self._split_table_row(row_line)
            if not cells:
                i += 1
                continue

            # Detect merge markers
            merge_row = []
            actual_cells = []
            for cell in cells:
                cell_stripped = cell.strip()
                if cell_stripped in ('>', 'v', '>', 'v'):
                    merge_row.append(cell_stripped)
                    actual_cells.append('')
                else:
                    merge_row.append('')
                    actual_cells.append(cell_stripped)

            rows.append(actual_cells)
            merge_map.append(merge_row)
            i += 1

        # Remove empty trailing rows
        while rows and all(c == '' for c in rows[-1]):
            rows.pop()
            merge_map.pop()

        if not rows:
            return None, i

        return MarkdownElement(
            type=MarkdownElementType.TABLE,
            rows=rows,
            merge_map=merge_map,
            children=[MarkdownElement(
                type=MarkdownElementType.TABLE,
                text=h,
            ) for h in headers] if headers else None,
        ), i

    def _split_table_row(self, row_line: str) -> List[str]:
        """Split a markdown table row into cells.

        Args:
            row_line: Raw table row line like "| A | B | C |"

        Returns:
            List of cell content strings
        """
        # Remove leading/trailing |
        inner = row_line.strip()
        if inner.startswith('|'):
            inner = inner[1:]
        if inner.endswith('|'):
            inner = inner[:-1]

        # Split by | and trim each cell
        cells = [c.strip() for c in inner.split('|')]
        return cells

    def _parse_unordered_list(self, lines: List[str], start: int
                              ) -> Tuple[List[str], int]:
        """Parse an unordered list starting at line start.

        Returns:
            (list of item texts, new_line_index)
        """
        items = []
        i = start
        while i < len(lines):
            m = self.RE_UNORDERED_LIST.match(lines[i].strip())
            if not m:
                break
            items.append(m.group(1).strip())
            i += 1
        return items, i

    def _parse_ordered_list(self, lines: List[str], start: int
                            ) -> Tuple[List[str], int]:
        """Parse an ordered list starting at line start.

        Returns:
            (list of item texts, new_line_index)
        """
        items = []
        i = start
        while i < len(lines):
            m = self.RE_ORDERED_LIST.match(lines[i].strip())
            if not m:
                break
            items.append(m.group(2).strip())
            i += 1
        return items, i

    @staticmethod
    def parse_inline_formatting(text: str) -> List[Tuple[str, dict]]:
        """Parse inline markdown formatting into segments.

        Args:
            text: Plain text with possible markdown formatting

        Returns:
            List of (text, format_overrides) tuples.
            Format overrides is a dict with keys like 'bold', 'italic', etc.
        """
        segments = []
        remaining = text
        pos = 0

        # Simple regex-based approach
        # Pattern order: code, bold, italic, strike
        pattern = re.compile(
            r'(`[^`]+`)|'           # code
            r'(\*\*.+?\*\*|__.+?__)|' # bold
            r'(\*.+?\*|_.+?_)|'      # italic
            r'(~~.+?~~)'             # strike
        )

        last_end = 0
        for m in pattern.finditer(text):
            # Add plain text before this match
            if m.start() > last_end:
                segments.append((text[last_end:m.start()], {}))

            matched = m.group(0)

            # Code
            if m.group(1):
                inner = matched[1:-1]  # Remove backticks
                segments.append((inner, {'monospace': True}))

            # Bold
            elif m.group(2):
                inner = matched[2:-2]  # Remove ** or __
                segments.append((inner, {'bold': True}))

            # Italic
            elif m.group(3):
                inner = matched[1:-1]  # Remove * or _
                segments.append((inner, {'italic': True}))

            # Strike
            elif m.group(4):
                inner = matched[2:-2]  # Remove ~~
                segments.append((inner, {'strike': True}))

            last_end = m.end()

        # Remaining text
        if last_end < len(text):
            segments.append((text[last_end:], {}))

        return segments if segments else [(text, {})]

    @staticmethod
    def extract_images_from_markdown(text: str
                                     ) -> List[Tuple[str, str, str]]:
        """Extract image references from markdown text.

        Returns:
            List of (alt_text, path_or_url, full_markdown) tuples
        """
        return [
            (m.group(1), m.group(2), m.group(0))
            for m in MarkdownProcessor.RE_IMAGE.finditer(text)
        ]

    # ======================== DOCX Conversion ========================

    def apply_to_document(
        self,
        elements: List[MarkdownElement],
        body_template: Optional[ParagraphData] = None,
        heading_template: Optional[ParagraphData] = None,
        table_template: Optional[ParagraphData] = None,
    ) -> List:
        """Convert parsed markdown elements into docx paragraphs/tables/images.

        Different content types use different templates:
        - heading_template: for HEADING elements (from the chapter's heading paragraph)
        - body_template: for body-like elements: PARAGRAPH, CODE_BLOCK, LIST, BLOCKQUOTE, etc.

        Args:
            elements: List of MarkdownElement from parse_markdown()
            body_template: Template for body-like elements
            heading_template: Template for heading elements

        Returns:
            List of created docx paragraph/table objects
        """
        created = []

        for element in elements:
            created_elements = self._convert_element(
                element, body_template, heading_template, table_template,
            )
            if created_elements:
                if isinstance(created_elements, list):
                    created.extend(created_elements)
                else:
                    created.append(created_elements)

        return created

    def _convert_element(
        self,
        element: MarkdownElement,
        body_template: Optional[ParagraphData] = None,
        heading_template: Optional[ParagraphData] = None,
        table_template: Optional[ParagraphData] = None,
    ):
        """Convert a single MarkdownElement to docx element(s).

        Args:
            element: The markdown element to convert
            body_template: Template for body-like elements (paragraphs, lists, code, etc.)
            heading_template: Template for heading elements (from the chapter's heading)
        """
        if element.type == MarkdownElementType.HEADING:
            return self._add_heading(element.text, element.level, heading_template)

        elif element.type == MarkdownElementType.PARAGRAPH:
            return self._add_paragraph(element.text, body_template)

        elif element.type == MarkdownElementType.TABLE:
            return self._build_table(element, table_template or body_template)

        elif element.type == MarkdownElementType.IMAGE:
            return self._add_image(element)

        elif element.type == MarkdownElementType.CODE_BLOCK:
            return self._add_code_block(element.text, element.code_language, body_template)

        elif element.type == MarkdownElementType.MERMAID:
            # Mermaid rendering is handled separately - return a placeholder
            return self._add_paragraph(
                f"[Mermaid diagram: {element.text[:50]}...]", body_template
            )

        elif element.type == MarkdownElementType.LIST:
            return self._add_list(element.items, element.ordered, body_template)

        elif element.type == MarkdownElementType.HORIZONTAL_RULE:
            return self._add_horizontal_rule()

        elif element.type == MarkdownElementType.BLOCKQUOTE:
            return self._add_paragraph(f"> {element.text}", body_template)

        elif element.type == MarkdownElementType.EMPTY_LINE:
            return self._add_paragraph('', body_template)

        return None

    def _add_heading(
        self,
        text: str,
        level: int,
        template: Optional[ParagraphData] = None,
    ):
        """Add a heading paragraph to the document."""
        level = min(max(level, 1), 9)
        para = self.document.add_heading(text, level=level)
        if template and template.formatting:
            run_fmt = self._extract_run_fmt(template)
            self._apply_paragraph_format(para, template.formatting, run_fmt)
        return para

    @staticmethod
    def _extract_run_fmt(template: Optional[ParagraphData]) -> Optional['FormattingData']:
        """Extract the first run's FormattingData from a ParagraphData template."""
        if template and template.runs_data:
            return template.runs_data[0][1]
        return None

    def _add_paragraph(
        self,
        text: str,
        template: Optional[ParagraphData] = None,
    ):
        """Add a regular paragraph with optional inline formatting."""
        if not text:
            para = self.document.add_paragraph()
            return para

        run_fmt = self._extract_run_fmt(template)

        # Parse inline formatting
        segments = self.parse_inline_formatting(text)

        para = self.document.add_paragraph()

        if template and template.formatting:
            self._apply_paragraph_format(para, template.formatting, run_fmt)

        for seg_text, fmt_overrides in segments:
            run = para.add_run(seg_text)
            self._apply_inline_format(run, fmt_overrides)

        # Apply template font to runs (runs didn't exist when
        # _apply_paragraph_format was called with run_fmt above).
        if run_fmt:
            from docx.shared import Pt
            for run in para.runs:
                if run_fmt.font_name and not run.font.name:
                    run.font.name = run_fmt.font_name
                if run_fmt.font_name_east_asia:
                    self._set_run_font_east_asia(run, run_fmt.font_name_east_asia)
                if run_fmt.size and run.font.size is None:
                    run.font.size = Pt(run_fmt.size / 2)

        return para

    def _add_image(self, element: MarkdownElement):
        """Add an image paragraph to the document."""
        para = self.document.add_paragraph()

        img_path = element.image_path
        if img_path and os.path.exists(img_path):
            try:
                run = para.add_run()
                run.add_picture(img_path, width=Inches(5))
            except Exception:
                para.add_run(f"[Image: {element.alt_text or element.image_path}]")
        else:
            para.add_run(f"[Image: {element.alt_text or element.image_path}]")

        return para

    def _add_code_block(
        self,
        code: str,
        language: Optional[str] = None,
        template: Optional[ParagraphData] = None,
    ):
        """Add a code block as a styled single-cell table with monospace text.

        Uses a 1×1 table as the visual container so the entire code block
        shares one border and background — the standard Word technique.
        """
        from docx.shared import Pt, Cm
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.text import WD_LINE_SPACING

        # ── 1×1 table ──
        table = self.document.add_table(rows=1, cols=1)

        # ── Table-level properties ──
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Borders: solid box, dark gray, 0.5 pt
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            if edge in ('insideH', 'insideV'):
                el.set(qn('w:val'), 'none')
            else:
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')          # 0.5pt = 4 eighth-points
                el.set(qn('w:color'), '808080')  # Dark gray
            el.set(qn('w:space'), '0')
            borders.append(el)
        tblPr.append(borders)

        # Table width: 100 % of container
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)

        # ── Cell formatting ──
        cell = table.cell(0, 0)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Background: very light gray
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        tcPr.append(shd)

        # Cell margins (padding) — 4 pt = 80 twips on each side
        tcMar = OxmlElement('w:tcMar')
        for edge in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:w'), '80')
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)

        # Remove the default empty paragraph python-docx inserts in every cell
        default_para = cell.paragraphs[0]
        default_para._element.getparent().remove(default_para._element)

        # ── Code lines as paragraphs inside the cell ──
        lines = code.split('\n')
        for i, line in enumerate(lines):
            para = cell.add_paragraph()
            run = para.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

            pf = para.paragraph_format
            pf.left_indent = Cm(0.5)           # ≈ 2 characters
            pf.line_spacing = Pt(12)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

            # Vertical gap before the first line and after the last line
            if i == 0:
                pf.space_before = Pt(6)
            elif i == len(lines) - 1:
                pf.space_after = Pt(6)

        # Ensure _element compatibility — build_elements_for_chapter reads `_element`
        if not hasattr(table, '_element'):
            table._element = table._tbl

        return table

    def _add_list(
        self,
        items: List[str],
        ordered: bool,
        template: Optional[ParagraphData] = None,
    ):
        """Add a list as paragraphs with appropriate formatting."""
        created = []
        for i, item in enumerate(items):
            prefix = f"{i + 1}. " if ordered else "- "
            para = self._add_paragraph(f"{prefix}{item}", template)
            created.append(para)
        return created

    def _add_horizontal_rule(self):
        """Add a horizontal rule as a paragraph with bottom border."""
        para = self.document.add_paragraph()
        # Add bottom border via XML
        pPr = para._element.get_or_add_pPr()
        pBdr = docx_qn('w:pBdr')
        bottom = docx_qn('w:bottom')
        bord = para._element.makeelement(bottom, {
            docx_qn('w:val'): 'single',
            docx_qn('w:sz'): '6',
            docx_qn('w:space'): '1',
            docx_qn('w:color'): 'auto',
        })
        pPr.append(bord)
        return para

    def _build_table(self, element: MarkdownElement,
                     body_template: Optional[ParagraphData] = None):
        """Build a table from a Table MarkdownElement, applying template font."""
        from .table_builder import TableBuilder
        builder = TableBuilder(self.document)
        return builder.build_table(
            headers=[c.text for c in element.children] if element.children else None,
            rows=element.rows or [],
            merge_map=element.merge_map,
            body_template=body_template,
        )

    # ======================== Formatting Helpers ========================

    def _apply_paragraph_format(self, para, fmt: ParagraphFormatData,
                                 run_fmt: Optional['FormattingData'] = None):
        """Apply ParagraphFormatData to a docx paragraph, including fonts.

        Args:
            para: python-docx Paragraph to format
            fmt: ParagraphFormatData (alignment, indentation, spacing, etc.)
            run_fmt: Optional FormattingData for run-level font info
                     (font name, east-asian font, size).  The ``_add_paragraph``
                     caller retrieves this from template.runs_data.
        """
        from docx.shared import Emu, Pt as PtShared
        from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pf = para.paragraph_format

        if fmt.alignment:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            pf.alignment = align_map.get(fmt.alignment.lower())

        if fmt.first_line_indent:
            # fmt.first_line_indent is in twips (1/20 pt).
            # python-docx expects EMU: 1 twip = 635 EMU.
            pf.first_line_indent = Emu(fmt.first_line_indent * 635)

        if fmt.left_indent:
            pf.left_indent = Emu(fmt.left_indent * 635)

        if fmt.space_before is not None:
            # Write via OOXML directly so even 0 values are preserved
            from docx.oxml import OxmlElement
            ppr = para._element.get_or_add_pPr()
            spacing = ppr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                ppr.append(spacing)
            spacing.set(qn('w:before'), str(int(fmt.space_before)))

        if fmt.space_after is not None:
            from docx.oxml import OxmlElement
            ppr = para._element.get_or_add_pPr()
            spacing = ppr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                ppr.append(spacing)
            spacing.set(qn('w:after'), str(int(fmt.space_after)))

        if fmt.line_spacing is not None:
            # Write line spacing directly via OOXML to preserve exact values.
            # python-docx's Python-level API converts values in ways that can
            # corrupt the OOXML representation (e.g. 300 → 72000 EMU).
            from docx.oxml import OxmlElement
            ppr = para._element.get_or_add_pPr()
            spacing_el = ppr.find(qn('w:spacing'))
            if spacing_el is None:
                spacing_el = OxmlElement('w:spacing')
                ppr.append(spacing_el)
            spacing_el.set(qn('w:line'), str(int(fmt.line_spacing)))
            rule_str = (fmt.line_spacing_rule or 'auto').lower()
            spacing_el.set(qn('w:lineRule'), rule_str)

        # ── Apply run-level font formatting to all existing runs ──
        if run_fmt:
            for run in para.runs:
                if run_fmt.font_name and not run.font.name:
                    run.font.name = run_fmt.font_name
                if run_fmt.font_name_east_asia:
                    self._set_run_font_east_asia(run, run_fmt.font_name_east_asia)
                if run_fmt.size and run.font.size is None:
                    run.font.size = PtShared(run_fmt.size / 2)  # half-pts → pts
                if run_fmt.bold is not None and run.font.bold is None:
                    run.font.bold = run_fmt.bold
                if run_fmt.italic is not None and run.font.italic is None:
                    run.font.italic = run_fmt.italic

    @staticmethod
    def _set_run_font_east_asia(run, font_name: str):
        """Set the East-Asian font name on a run (w:rFonts w:eastAsia)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    def _apply_inline_format(self, run, fmt_overrides: dict):
        """Apply inline format overrides to a docx run."""
        if fmt_overrides.get('bold'):
            run.bold = True
        if fmt_overrides.get('italic'):
            run.italic = True
        if fmt_overrides.get('strike'):
            run.font.strike = True
        if fmt_overrides.get('monospace'):
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # ======================== Chapter-level Operations ========================

    def build_elements_for_chapter(
        self,
        elements: List[MarkdownElement],
        body_template: Optional[ParagraphData] = None,
        heading_template: Optional[ParagraphData] = None,
        table_template: Optional[ParagraphData] = None,
    ) -> list:
        """Build new docx paragraph/table elements ready for insertion.

        Args:
            elements: Parsed markdown elements
            body_template: Template for body-like elements (paragraphs, lists, code, etc.)
            heading_template: Template for heading elements (from the chapter's heading)

        Returns list of lxml elements that can be inserted into the document body.
        """
        # Apply to document (they get added to the end)
        created = self.apply_to_document(
            elements, body_template, heading_template, table_template,
        )

        # Extract the XML elements and remove from document
        result = []
        for item in created:
            if hasattr(item, '_element'):
                element = item._element
            elif hasattr(item, '_tbl'):
                element = item._tbl
            else:
                continue
            result.append(element)
            # Remove from document (they'll be re-inserted elsewhere)
            try:
                element.getparent().remove(element)
            except Exception:
                pass

        return result
