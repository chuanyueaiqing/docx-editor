"""Main DocxDocument facade class.

Provides the single public entry point for all DOCX operations.
Delegates to specialized sub-modules for specific concerns.
"""
import logging
import os
import re
import tempfile
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document as DocxDocumentLoader
from docx.oxml.ns import qn as docx_qn

from .chapter_parser import ChapterParser
from .format_extractor import FormatExtractor
from .markdown_processor import MarkdownProcessor
from .mermaid_renderer import MermaidRenderer, MermaidNotAvailableError
from .models import (
    ChapterNode, CommentData, FormatStore, MarkdownElement,
    MarkdownElementType, ParagraphData, RevisionData,
)
from .utils import ChapterNotFoundError, CommentsReadError, DocxError
from .win32_ops import Win32Ops

logger = logging.getLogger(__name__)


class DocxDocument:
    """Main facade for reading and editing DOCX files.

    Usage:
        doc = DocxDocument('path/to/file.docx')

        # Chapter-based access
        chapter = doc.get_chapter('3.1')
        contents = doc.get_chapter_contents(chapter)

        # Replace chapter with markdown
        doc.replace_chapter('3.1', '# Title\\n\\nContent with **bold**.')

        # Search and replace
        doc.search_replace_in_chapter('3.1', 'old', 'new')

        # Delete operations
        doc.delete_chapter('3.1')
        doc.delete_in_chapter('3.1', 'text to delete')

        # Read comments and revisions
        comments = doc.read_comments()
        revisions = doc.read_revisions()

        # Save
        doc.save('output.docx')
    """

    def __init__(
        self,
        path: str,
        use_track_changes: bool = False,
    ):
        """Load a DOCX document and build internal data structures.

        Args:
            path: Path to the .docx file
            use_track_changes: When True, write operations use win32com
                               for revision tracking. Default: False.

        Raises:
            FileNotFoundError: If the path does not exist
            DocxError: If the file cannot be loaded as a DOCX
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"Document not found: {path}")

        self.path = path
        self.use_track_changes = use_track_changes

        # Load the document
        self.document = DocxDocumentLoader(path)

        # Build format store
        self.format_store = FormatStore(
            docx_path=path,
            document=self.document,
        )

        # Extract formatting
        self.format_store.formats_json = FormatExtractor.extract_all(self.document)

        # Build paragraph data
        self._build_paragraphs_data()

        # Initialize sub-modules
        self.chapter_parser = ChapterParser(self.format_store)
        self.markdown_processor = MarkdownProcessor(self.format_store)
        self.mermaid_renderer = MermaidRenderer()

        # Track whether the document has been modified
        self._modified = False
        self._batch_counter = 0

    def _build_paragraphs_data(self):
        """Build ParagraphData list from document paragraphs and format store."""
        self.format_store.paragraphs_data = []
        for i, para in enumerate(self.document.paragraphs):
            fmt_data = self.format_store.formats_json.get(i, {})
            pd = ParagraphData(
                paragraph=para,
                index=i,
                text=para.text,
                formatting=fmt_data.get('paragraph_format'),
            )
            self.format_store.paragraphs_data.append(pd)

    def _refresh_paragraphs_text(self):
        """Update ``ParagraphData.text`` from the current run state.

        ``search_replace_in_chapter`` and ``delete_in_chapter`` modify
        ``run.text`` in-place, but ``paragraphs_data`` cached the old
        ``para.text`` at build time.  This lightweight sync avoids a
        full ``_reload()``.
        """
        for i, pd in enumerate(self.format_store.paragraphs_data):
            if i < len(self.document.paragraphs):
                pd.text = self.document.paragraphs[i].text

    def _reload(self):
        """Reload the document after structural changes.

        After modifying the XML element tree, python-docx Paragraph objects
        become stale. This saves to temp, reloads, and rebuilds state.

        If called within a ``begin_batch()`` … ``end_batch()`` block the
        reload is deferred until the batch ends.
        """
        if self._batch_counter > 0:
            return  # Defer until batch ends

        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            self.document.save(tmp.name)
            self.document = DocxDocumentLoader(tmp.name)
        finally:
            os.unlink(tmp.name)

        # Rebuild all internal state
        self.format_store.document = self.document
        self.format_store.formats_json = FormatExtractor.extract_all(self.document)
        self._build_paragraphs_data()
        self.chapter_parser = ChapterParser(self.format_store)
        self.markdown_processor = MarkdownProcessor(self.format_store)

    def begin_batch(self):
        """Enter batch mode — defer ``_reload()`` until ``end_batch()``.

        Use when making many ``replace_chapter()`` / ``delete_chapter()``
        calls in a row to avoid O(N × document) repeated reloads.
        """
        self._batch_counter += 1

    def end_batch(self):
        """Exit batch mode and perform one final reload."""
        self._batch_counter = max(0, self._batch_counter - 1)
        if self._batch_counter == 0:
            self._reload()

    # ======================== Chapter Access ========================

    def get_chapter(self, number: Union[str, Tuple[int, ...]]) -> Optional[ChapterNode]:
        """Look up a chapter by number string like '3.1.1' or tuple (3,1,1).

        Args:
            number: Chapter number as string (e.g. "3.1") or tuple (e.g. (3,1))

        Returns:
            ChapterNode if found, None otherwise
        """
        return self.chapter_parser.get_chapter_by_number(number)

    def get_chapter_contents(self, chapter: ChapterNode) -> List[ParagraphData]:
        """Get all paragraphs belonging to a chapter (heading + body).

        Args:
            chapter: ChapterNode from get_chapter()

        Returns:
            List of ParagraphData objects
        """
        return self.chapter_parser.get_chapter_contents(chapter)

    def get_chapter_text(self, chapter: ChapterNode) -> str:
        """Get concatenated text of a chapter.

        Args:
            chapter: ChapterNode from get_chapter()

        Returns:
            Combined text of heading + body paragraphs
        """
        return self.chapter_parser.get_chapter_text(chapter)

    def get_chapter_tree(self) -> List[ChapterNode]:
        """Get the root-level chapter tree.

        Returns:
            List of top-level ChapterNode objects
        """
        return self.chapter_parser.root_chapters

    def print_chapter_tree(self):
        """Print a visual representation of the chapter tree."""
        print(self.chapter_parser.tree_to_string())

    # ======================== Content Replacement ========================

    def replace_chapter(
        self,
        number: Union[str, Tuple[int, ...]],
        md_text: str,
    ):
        """Replace a chapter's entire content with markdown text.

        Preserves the original heading paragraph and replaces body content
        with parsed markdown elements. Supports tables, images, mermaid
        diagrams, lists, and inline formatting.

        Args:
            number: Chapter number to replace
            md_text: Markdown text for the new content

        Raises:
            ChapterNotFoundError: If the chapter number doesn't exist
            DocxError: If the replacement fails
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        if self.use_track_changes:
            self._replace_chapter_win32(chapter, md_text)
        else:
            self._replace_chapter_pythondocx(chapter, md_text)

    def _replace_chapter_pythondocx(
        self,
        chapter: ChapterNode,
        md_text: str,
    ):
        """Replace chapter content using python-docx XML manipulation."""
        # Parse markdown
        elements = self.markdown_processor.parse_markdown(md_text)

        # Handle mermaid elements: render to images
        elements = self._render_mermaid_elements(elements)

        # Get templates by content type
        # body_template: from first body paragraph (for paragraphs, lists, code blocks, etc.)
        body_template = None
        if chapter.body_paragraph_indices:
            first_idx = chapter.body_paragraph_indices[0]
            if first_idx < len(self.format_store.paragraphs_data):
                body_template = self.format_store.paragraphs_data[first_idx]
        # heading_template: from chapter's own heading paragraph (for headings only)
        heading_template = None
        heading_idx = chapter.heading_paragraph_index
        if heading_idx < len(self.format_store.paragraphs_data):
            heading_template = self.format_store.paragraphs_data[heading_idx]

        # Get the heading element reference
        heading_para = self.document.paragraphs[chapter.heading_paragraph_index]
        body = self.document.element.body
        heading_element = heading_para._element

        # Delete existing body paragraphs (work backwards)
        for idx in sorted(chapter.body_paragraph_indices, reverse=True):
            try:
                para = self.document.paragraphs[idx]
                body.remove(para._element)
            except Exception:
                logger.debug(
                    "Could not remove body paragraph %d; may already be gone", idx
                )

        # Build new docx elements with content-type-specific templates
        new_elements = self.markdown_processor.build_elements_for_chapter(
            elements, body_template, heading_template
        )

        # Insert new elements after the heading (in reverse order to maintain sequence)
        for elem in reversed(new_elements):
            try:
                heading_element.addnext(elem)
            except Exception:
                logger.debug("Could not insert new element after heading", exc_info=True)

        # Reload to refresh internal state
        self._reload()
        self._modified = True

    def _replace_chapter_win32(
        self,
        chapter: ChapterNode,
        md_text: str,
    ):
        """Replace chapter content with win32com, producing revision marks.

        Uses Word's ``CompareDocuments`` API to diff the original and the
        modified document, generating proper ``<w:ins>`` / ``<w:del>``
        tracked changes.
        """
        if not Win32Ops.is_word_available():
            raise DocxError(
                "Track changes mode requires Microsoft Word. "
                "Use use_track_changes=False to use python-docx instead."
            )

        # --- 1. Save original and modified versions side-by-side ---
        original_tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        original_tmp.close()
        modified_tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        modified_tmp.close()

        try:
            # Snapshot the document BEFORE python-docx modifies it
            self.document.save(original_tmp.name)

            # Do the python-docx replacement
            self._replace_chapter_pythondocx(chapter, md_text)
            self.document.save(modified_tmp.name)

            # --- 2. Open modified version in Word and compare against original ---
            with Win32Ops() as ops:
                ops.open_document(modified_tmp.name)
                # Compare() generates <w:ins>/<w:del> for every difference
                # between the opened document and the original snapshot.
                ops.word.ActiveDocument.Compare(
                    Name=original_tmp.name,
                    CompareTarget=2,            # wdCompareTargetCurrent
                    IgnoreAllComparisonWarnings=True,
                )
                ops.save_document(self.path)

        finally:
            for tmp_path in (original_tmp.name, modified_tmp.name):
                try:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                except Exception:
                    logger.warning("Failed to clean up temp file: %s", tmp_path)

        # Reload
        self.document = DocxDocumentLoader(self.path)
        self._reload()

    def _render_mermaid_elements(
        self,
        elements: List,
    ) -> List:
        """Render mermaid elements to images in-place.

        Args:
            elements: List of MarkdownElement objects

        Returns:
            Modified elements list with mermaid blocks replaced by images
        """
        rendered = []
        for elem in elements:
            if elem.type == MarkdownElementType.MERMAID:
                # Try to render mermaid
                if self.mermaid_renderer.is_available() and elem.text:
                    try:
                        img_path = self.mermaid_renderer.render(elem.text)
                        rendered.append(MarkdownElement(
                            type=MarkdownElementType.IMAGE,
                            image_path=img_path,
                            alt_text="Mermaid diagram",
                        ))
                        continue
                    except Exception:
                        pass
                # Fallback: keep as code block
                rendered.append(MarkdownElement(
                    type=MarkdownElementType.CODE_BLOCK,
                    text=elem.text or '',
                    code_language='mermaid',
                ))
            else:
                rendered.append(elem)
        return rendered

    # ======================== Search and Replace ========================

    @staticmethod
    def _prepare_paragraph_for_replace(
        para,
        old_text: str,
        case_sensitive: bool = False,
    ) -> bool:
        """Ensure *old_text* is findable within a single run.

        When ``old_text`` crosses a run boundary (e.g. bold turns off mid-word),
        Word splits the text across ``<w:r>`` elements.  This method merges all
        runs in the paragraph into the first run so that ``run.text`` contains
        the full paragraph text and a simple ``old_text in run.text`` check
        works.  Formatting from the first run is preserved; per-run formatting
        details are lost.

        Returns ``True`` if ``old_text`` can be found in the paragraph text
        (either natively or after merging).
        """
        full = para.text
        if not full:
            return False

        if case_sensitive:
            if old_text not in full:
                return False
            # Already findable in a single run?
            if any(old_text in r.text for r in para.runs if r.text):
                return True
        else:
            lower_old = old_text.lower()
            if lower_old not in full.lower():
                return False
            if any(lower_old in r.text.lower() for r in para.runs if r.text):
                return True

        # Cross-run merge: concatenate all text into the first run
        runs = para.runs
        if runs:
            runs[0].text = full
            for r in runs[1:]:
                r.text = ''
        return True

    def search_replace_in_chapter(
        self,
        number: Union[str, Tuple[int, ...]],
        old_text: str,
        new_text: str,
        case_sensitive: bool = False,
    ) -> int:
        """Search and replace text within a specific chapter.

        Iterates through the chapter's paragraphs and replaces text
        in each run. Operates at the text level, not XML level.

        .. note::
           If ``old_text`` crosses a run boundary (uncommon in well-formed
           documents), runs are merged so the replacement still happens.
           Doing so loses per-run formatting differentiation for that
           paragraph.

        Args:
            number: Chapter number
            old_text: Text to search for
            new_text: Replacement text
            case_sensitive: Whether to match case

        Returns:
            Number of replacements made
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        count = 0
        indices = [chapter.heading_paragraph_index] + list(chapter.body_paragraph_indices)

        for idx in indices:
            if idx >= len(self.document.paragraphs):
                continue
            para = self.document.paragraphs[idx]

            # Cross-run fallback
            if not self._prepare_paragraph_for_replace(para, old_text, case_sensitive):
                continue

            for run in para.runs:
                if not run.text:
                    continue
                if case_sensitive:
                    occurrences = run.text.count(old_text)
                    if occurrences:
                        run.text = run.text.replace(old_text, new_text)
                        count += occurrences
                else:
                    matches = run.text.lower().count(old_text.lower())
                    if matches:
                        run.text = self._replace_ignore_case(
                            run.text, old_text, new_text
                        )
                        count += matches

        if count > 0:
            self._modified = True
            self._refresh_paragraphs_text()

        return count

    @staticmethod
    def _replace_ignore_case(text: str, old: str, new: str) -> str:
        """Replace *old* with *new* ignoring case.

        .. caution::
           The replacement text ``new`` is used **as-is** — the capitalisation
           pattern of the matched text is **not** preserved (despite the
           earlier docstring claiming otherwise).

        Args:
            text: Original text
            old: Text to find (case-insensitive)
            new: Replacement text

        Returns:
            Text with replacements
        """
        return re.sub(re.escape(old), new, text, flags=re.IGNORECASE)

    # ======================== Deletion ========================

    def delete_chapter(self, number: Union[str, Tuple[int, ...]]):
        """Delete an entire chapter (heading + all body content).

        Args:
            number: Chapter number to delete

        Raises:
            ChapterNotFoundError: If the chapter doesn't exist
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        self.chapter_parser.delete_chapter(chapter)
        self._reload()
        self._modified = True

    def delete_in_chapter(
        self,
        number: Union[str, Tuple[int, ...]],
        target_text: str,
    ) -> int:
        """Delete specific text content from within a chapter.

        Args:
            number: Chapter number
            target_text: Text to delete

        Returns:
            Number of deletions made
        """
        chapter = self.chapter_parser.get_chapter_by_number(number)
        if chapter is None:
            raise ChapterNotFoundError(f"Chapter {number} not found")

        count = 0
        indices = [chapter.heading_paragraph_index] + list(chapter.body_paragraph_indices)

        for idx in indices:
            if idx >= len(self.document.paragraphs):
                continue
            para = self.document.paragraphs[idx]

            # Cross-run fallback
            if not self._prepare_paragraph_for_replace(para, target_text, case_sensitive=True):
                continue

            for run in para.runs:
                if run.text and target_text in run.text:
                    occurrences = run.text.count(target_text)
                    run.text = run.text.replace(target_text, '')
                    count += occurrences

        if count > 0:
            self._modified = True
            self._refresh_paragraphs_text()

        return count

    # ======================== Comments and Revisions ========================

    def read_comments(self) -> List[CommentData]:
        """Read comments from the document.

        Tries python-docx first, falls back to win32com.

        Returns:
            List of CommentData objects
        """
        try:
            return self._read_comments_python_docx()
        except Exception:
            if Win32Ops.is_word_available():
                return self._read_comments_win32()
            raise CommentsReadError("Cannot read comments from document")

    def _read_comments_python_docx(self) -> List[CommentData]:
        """Read comments using python-docx API.

        Raises ``CommentsReadError`` if the ``.comments`` attribute is
        unavailable (python-docx < 0.8.11), allowing the caller's win32com
        fallback to fire.
        """
        try:
            comments_part = self.document.comments
        except AttributeError:
            raise CommentsReadError(
                "python-docx version does not support .comments attribute"
            )

        if comments_part is None:
            return []

        qn = docx_qn
        results = []

        for comment in comments_part:
            try:
                results.append(CommentData(
                    id=comment._element.get(qn('w:id'), ''),
                    author=comment._element.get(qn('w:author'), ''),
                    date=comment._element.get(qn('w:date'), ''),
                    text=comment.text.strip() if comment.text else '',
                ))
            except Exception:
                logger.debug("Skipping a comment that could not be read", exc_info=True)

        return results

    def _read_comments_win32(self) -> List[CommentData]:
        """Fallback: read comments via win32com."""
        with Win32Ops() as ops:
            ops.open_document(self.path)
            return ops.read_comments()

    def read_revisions(self) -> List[RevisionData]:
        """Read tracked changes from the document.

        Parses document XML directly for w:ins and w:del elements.
        Falls back to win32com if XML parsing fails.

        Returns:
            List of RevisionData sorted by document position
        """
        try:
            return self._read_revisions_xml()
        except Exception:
            if Win32Ops.is_word_available():
                return self._read_revisions_win32()
            raise

    def _read_revisions_xml(self) -> List[RevisionData]:
        """Read tracked changes by parsing document XML directly.

        Uses ``document.paragraphs`` order for indices so the returned
        ``paragraph_index`` values match ``document.paragraphs[i]``
        (unlike the old code which iterated ``body.iterchildren()`` and
        counted tables as an index slot, causing drift).
        """
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        revisions = []

        for para_index, para in enumerate(self.document.paragraphs):
            p_element = para._element

            # Find w:ins elements (insertions)
            for ins in p_element.iter(f'{{{W}}}ins'):
                text_parts = []
                for t in ins.iter(f'{{{W}}}t'):
                    if t.text:
                        text_parts.append(t.text)
                revisions.append(RevisionData(
                    rev_id=ins.get(f'{{{W}}}id', ''),
                    author=ins.get(f'{{{W}}}author', ''),
                    date=ins.get(f'{{{W}}}date', ''),
                    type='insertion',
                    text=''.join(text_parts),
                    paragraph_index=para_index,
                ))

            # Find w:del elements (deletions)
            for del_elem in p_element.iter(f'{{{W}}}del'):
                text_parts = []
                for dt in del_elem.iter(f'{{{W}}}delText'):
                    if dt.text:
                        text_parts.append(dt.text)
                revisions.append(RevisionData(
                    rev_id=del_elem.get(f'{{{W}}}id', ''),
                    author=del_elem.get(f'{{{W}}}author', ''),
                    date=del_elem.get(f'{{{W}}}date', ''),
                    type='deletion',
                    text=''.join(text_parts),
                    paragraph_index=para_index,
                ))

        return revisions

    def _read_revisions_win32(self) -> List[RevisionData]:
        """Fallback: read tracked changes via win32com."""
        with Win32Ops() as ops:
            ops.open_document(self.path)
            return ops.read_revisions()

    # ======================== Save ========================

    def save(self, path: Optional[str] = None):
        """Save the document.

        Args:
            path: Output path. If None, overwrites the original file.

        Note:
            When saving to a new path, ``self.path`` is updated so that
            subsequent track-changes operations target the new file.
        """
        save_path = path or self.path
        self.document.save(save_path)
        self.path = save_path
        self._modified = False

    # ======================== Properties ========================

    @property
    def is_modified(self) -> bool:
        """Check if the document has been modified since load/save."""
        return self._modified
