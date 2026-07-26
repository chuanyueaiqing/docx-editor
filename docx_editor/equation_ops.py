"""Unified equation operations module.

Provides a unified interface for converting LaTeX math placeholders
to native editable equations in both Microsoft Word and WPS Office.

Usage:
    from docx_editor.equation_ops import EquationProcessor

    # Check availability at build time
    if EquationProcessor.is_com_available():
        elements = [...]  # MATH/DISPLAY_MATH with bookmarks
    else:
        img = EquationProcessor.render_latex_as_image('E=mc^2', 'temp.png')

    # Post-process a saved document (COM phase)
    proc = EquationProcessor()
    result = proc.post_process('document.docx', math_elements)
"""

import logging
import os
import tempfile
from enum import Enum
from typing import List, Optional, Tuple

from .utils import EquationError

logger = logging.getLogger(__name__)


class EquationResult(Enum):
    """Result of equation post-processing."""
    CONVERTED_COM = 'converted_com'      # Successfully converted via COM
    CONVERTED_WPS = 'converted_wps'      # Converted via WPS COM
    FALLBACK_IMAGE = 'fallback_image'    # Rendered as image
    SKIPPED = 'skipped'                   # No equations to process
    PARTIAL = 'partial'                  # Some succeeded, some fell back


class EquationProcessor:
    """Process math equations in a saved DOCX document.

    Two usage modes:

    1. **Build-time check** — call ``is_com_available()`` to decide whether
       to insert bookmark placeholders (COM path) or render images immediately.

    2. **Post-processing** — call ``post_process()`` on the saved .docx to
       convert bookmark placeholders into native Word/WPS equations.
    """

    # ── Engine detection ──

    @staticmethod
    def is_com_available() -> bool:
        """Check if any COM engine (Microsoft Word or WPS) is available.

        Returns:
            True if either Word or WPS is installed and accessible via COM
        """
        from .win32_ops import Win32Ops
        from .wps_ops import WpsOps
        return Win32Ops.is_word_available() or WpsOps.is_wps_available()

    @staticmethod
    def detect_engine() -> str:
        """Detect the best available equation engine.

        Returns:
            ``'word'``, ``'wps'``, or ``'image'`` (fallback)
        """
        from .win32_ops import Win32Ops
        from .wps_ops import WpsOps

        if Win32Ops.is_word_available():
            return 'word'
        if WpsOps.is_wps_available():
            return 'wps'
        return 'image'

    # ── Image fallback ──

    @staticmethod
    def render_latex_as_image(
        latex: str,
        output_path: Optional[str] = None,
        fontsize: int = 20,
        dpi: int = 150,
        display: bool = False,
    ) -> str:
        """Render a LaTeX math expression to a PNG image using matplotlib.

        Uses matplotlib's built-in mathtext parser, so no LaTeX distribution
        is required.

        Args:
            latex: LaTeX math expression (without $ delimiters)
            output_path: Path for the output PNG. If None, a temp file is created.
            fontsize: Font size in points (default: 20)
            dpi: Output image DPI (default: 150)
            display: If True, uses display-style rendering (\\displaystyle)

        Returns:
            Path to the generated PNG file
        """
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
        except ImportError:
            raise EquationError(
                "matplotlib is required for image-based equation rendering. "
                "Install with: pip install matplotlib"
            )

        if output_path is None:
            fd, output_path = tempfile.mkstemp(suffix='.png', prefix='math_')
            os.close(fd)

        # Build the TeX string
        # Note: matplotlib's mathtext doesn't support \displaystyle,
        # so we skip it and let the user choose display style via fontsize
        tex_str = rf'${latex}$'

        try:
            fig, ax = plt.subplots(figsize=(0.01, 0.01))
            ax.axis('off')

            text = ax.text(
                0, 0, tex_str,
                fontsize=fontsize,
                horizontalalignment='left',
                verticalalignment='bottom',
                usetex=False,   # use matplotlib's built-in mathtext
            )

            fig.savefig(
                output_path,
                dpi=dpi,
                bbox_inches='tight',
                pad_inches=0.05,
                facecolor='white',
                edgecolor='none',
                transparent=False,
            )
            plt.close(fig)

            logger.debug('Rendered LaTeX to image: %s', output_path)
            return output_path

        except Exception as e:
            plt.close('all')
            raise EquationError(
                f"Failed to render LaTeX image: {e}\n"
                f"LaTeX: {latex}"
            ) from e

    # ── Post-processing ──

    def post_process(
        self,
        docx_path: str,
        math_elements: List[Tuple[str, bool]],
    ) -> EquationResult:
        """Post-process a saved .docx to convert equation placeholders.

        Opens the saved document via COM (Word or WPS) and converts each
        bookmarked LaTeX placeholder into a native equation.

        Args:
            docx_path: Path to the saved .docx file
            math_elements: List of ``(latex_string, is_display)`` tuples,
                in document order, matching the ``_MathEq_1``, ``_MathEq_2``,
                ... bookmarks.

        Returns:
            EquationResult indicating what happened

        Raises:
            EquationError: If post-processing fails entirely
        """
        if not math_elements:
            return EquationResult.SKIPPED

        if not os.path.exists(docx_path):
            raise EquationError(f"Document not found: {docx_path}")

        engine = self.detect_engine()

        if engine == 'word':
            return self._process_with_word(docx_path, math_elements)
        elif engine == 'wps':
            return self._process_with_wps(docx_path, math_elements)
        else:
            return self._process_with_images(docx_path, math_elements)

    def convert_equations_via_com(
        self,
        docx_path: str,
        math_placeholders: List[Tuple[str, str, bool]],
    ) -> EquationResult:
        """Convert LaTeX text placeholders to native equations via COM.

        Opens the document in Word/WPS, finds each bookmark by name,
        replaces the LaTeX text with UnicodeMath, and converts to a native
        editable equation via ``TypeText → OMaths.Add → BuildUp``.

        Args:
            docx_path: Path to the saved .docx file
            math_placeholders: List of ``(bookmark_name, latex, is_display)``

        Returns:
            EquationResult
        """
        if not math_placeholders:
            return EquationResult.SKIPPED
        if not os.path.exists(docx_path):
            raise EquationError(f"Document not found: {docx_path}")

        engine = self.detect_engine()
        if engine == 'image':
            logger.warning(
                'No COM engine available. Equations will remain as text.'
            )
            return EquationResult.FALLBACK_IMAGE

        from .latex_to_unicodemath import latex_to_unicodemath
        from .win32_ops import Win32Ops
        from .wps_ops import WpsOps
        OpsClass = Win32Ops if engine == 'word' else WpsOps

        try:
            with OpsClass() as ops:
                ops.open_document(docx_path)

                converted = 0
                for bm_name, latex, _ in math_placeholders:
                    try:
                        # Find the bookmark
                        bm = ops.wd_doc.Bookmarks(bm_name)
                        rng = bm.Range
                        text = rng.Text.strip()
                        if not text:
                            text = latex

                        # Convert LaTeX to UnicodeMath
                        unicodemath = latex_to_unicodemath(text)

                        # Save the bookmark range's start position
                        start_pos = bm.Range.Start

                        # Select the bookmark's range
                        rng.Select()

                        # Type the UnicodeMath text (replaces bookmark content)
                        ops.word.Selection.TypeText(unicodemath)

                        # Select from the saved start to current end
                        ops.word.Selection.SetRange(
                            start_pos, ops.word.Selection.End
                        )

                        # Create equation from the selected text
                        ops.word.Selection.OMaths.Add(
                            ops.word.Selection.Range
                        )

                        if ops.word.Selection.OMaths.Count > 0:
                            eq = ops.word.Selection.OMaths(1)
                            try:
                                eq.Linearize()
                            except Exception:
                                pass
                            try:
                                eq.BuildUp()
                            except Exception:
                                pass
                            converted += 1

                        # Delete the bookmark marker
                        try:
                            ops.wd_doc.Bookmarks(bm_name).Delete()
                        except Exception:
                            pass

                    except Exception as e:
                        logger.warning(
                            'Failed to convert equation "%s": %s',
                            bm_name, e,
                        )

                ops.save_document(docx_path)

            if converted == len(math_placeholders):
                return EquationResult.CONVERTED_COM
            elif converted > 0:
                return EquationResult.PARTIAL
            else:
                return EquationResult.FALLBACK_IMAGE

        except Exception as e:
            raise EquationError(
                f"COM equation conversion failed: {e}"
            ) from e

    def _process_with_word(
        self,
        docx_path: str,
        math_elements: List[Tuple[str, bool]],
    ) -> EquationResult:
        """Word COM path: raw LaTeX → OMaths.Add() + BuildUp()

        Word 365 natively understands LaTeX syntax for equations (via the
        Math AutoCorrect feature).  We pass raw LaTeX directly.
        No UnicodeMath conversion is needed for Word.
        """
        from .win32_ops import Win32Ops, Win32ComError

        try:
            with Win32Ops() as ops:
                ops.open_document(docx_path)

                converted = 0
                for i, (latex, _is_display) in enumerate(math_elements):
                    bookmark_name = f'_MathEq_{i + 1}'
                    try:
                        # Word accepts raw LaTeX directly
                        ops.convert_bookmark_to_equation(bookmark_name, latex)
                        converted += 1
                    except Exception:
                        logger.warning(
                            'Failed to convert equation %s', bookmark_name
                        )

                ops.save_document(docx_path)

            if converted == len(math_elements):
                return EquationResult.CONVERTED_COM
            elif converted > 0:
                return EquationResult.PARTIAL
            else:
                raise EquationError("All equations failed to convert via Word COM")

        except Win32ComError as e:
            raise EquationError(f"Word COM equation conversion failed: {e}") from e

    def _process_with_wps(
        self,
        docx_path: str,
        math_elements: List[Tuple[str, bool]],
    ) -> EquationResult:
        """WPS COM path: LaTeX → UnicodeMath → OMaths.Add() + BuildUp()"""
        from .wps_ops import WpsOps, WpsComError
        from .latex_to_unicodemath import latex_to_unicodemath

        try:
            with WpsOps() as ops:
                ops.open_document(docx_path)

                converted = 0
                for i, (latex, _is_display) in enumerate(math_elements):
                    bookmark_name = f'_MathEq_{i + 1}'
                    unicodemath = latex_to_unicodemath(latex)
                    try:
                        ops.convert_bookmark_to_equation(bookmark_name, unicodemath)
                        converted += 1
                    except Exception:
                        logger.warning(
                            'Failed to convert equation %s', bookmark_name
                        )

                ops.save_document(docx_path)

            if converted == len(math_elements):
                return EquationResult.CONVERTED_WPS
            elif converted > 0:
                return EquationResult.PARTIAL
            else:
                raise EquationError("All equations failed to convert via WPS COM")

        except WpsComError as e:
            raise EquationError(f"WPS COM equation conversion failed: {e}") from e

    def _process_with_images(
        self,
        docx_path: str,
        math_elements: List[Tuple[str, bool]],
    ) -> EquationResult:
        """Image fallback: re-open docx via python-docx, replace bookmark
        paragraphs with rendered LaTeX images."""
        from docx import Document as DocxLoader
        from docx.shared import Inches

        # Load the document
        doc = DocxLoader(docx_path)
        body = doc.element.body

        converted = 0
        for i, (latex, is_display) in enumerate(math_elements):
            bookmark_name = f'_MathEq_{i + 1}'

            # Find the paragraph containing this bookmark
            target_para = None
            for para in doc.paragraphs:
                if self._para_has_bookmark(para, bookmark_name):
                    target_para = para
                    break

            if target_para is None:
                logger.warning(
                    'Bookmark paragraph "%s" not found in document', bookmark_name
                )
                continue

            try:
                # Render LaTeX to image
                img_path = self.render_latex_as_image(
                    latex, display=is_display,
                )

                # Replace the paragraph with an image paragraph
                new_para = self._replace_with_image(
                    doc, target_para, img_path, is_display,
                )
                converted += 1

                # Clean up temp image
                try:
                    os.remove(img_path)
                except Exception:
                    pass

            except Exception as e:
                logger.warning(
                    'Image fallback failed for "%s": %s', bookmark_name, e
                )

        # Save
        doc.save(docx_path)

        if converted == len(math_elements):
            return EquationResult.FALLBACK_IMAGE
        elif converted > 0:
            return EquationResult.PARTIAL
        else:
            raise EquationError("All equations failed image fallback")

    @staticmethod
    def _para_has_bookmark(para, bookmark_name: str) -> bool:
        """Check if a python-docx paragraph contains a specific bookmark."""
        from docx.oxml.ns import qn
        for child in para._element:
            if child.tag == qn('w:bookmarkStart'):
                name = child.get(qn('w:name'))
                if name == bookmark_name:
                    return True
        return False

    @staticmethod
    def _replace_with_image(doc, old_para, img_path: str, center: bool):
        """Replace a paragraph with one containing an image.

        Args:
            doc: python-docx Document
            old_para: The paragraph to replace
            img_path: Path to the image file
            center: Whether to center the image

        Returns:
            The new image paragraph
        """
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create new paragraph
        new_para = doc.add_paragraph()

        if center:
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(4))

        # Swap XML elements in the body
        old_para._element.addnext(new_para._element)
        old_para._element.getparent().remove(old_para._element)

        return new_para
