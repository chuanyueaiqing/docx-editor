"""Format extraction module.

Extracts paragraph-level and run-level formatting from a python-docx Document
by directly reading the underlying lxml XML elements. This provides access to
all OOXML formatting attributes that python-docx's high-level API may not expose.
"""
from typing import Dict, List, Optional, Tuple

from docx import Document
from lxml import etree

from .models import FormattingData, ParagraphFormatData
from .utils import qn


class FormatExtractor:
    """Extract formatting from python-docx Document elements using lxml XML access.

    Usage:
        doc = Document('path.docx')
        formats = FormatExtractor.extract_all(doc)
        # formats[0] -> { 'paragraph_format': ParagraphFormatData,
        #                  'runs': [{'text': '...', 'formatting': FormattingData}, ...] }
    """

    @staticmethod
    def extract_runs_formatting(run_element) -> FormattingData:
        """Extract run-level formatting from a w:r element.

        Args:
            run_element: lxml element representing a <w:r> run

        Returns:
            FormattingData with all detected properties
        """
        rpr = run_element.find(qn('w:rPr'))
        if rpr is None:
            return FormattingData()

        data = FormattingData()

        # ---- Font names ----
        fonts = rpr.find(qn('w:rFonts'))
        if fonts is not None:
            data.font_name = fonts.get(qn('w:ascii'))
            data.font_name_east_asia = fonts.get(qn('w:eastAsia'))
            data.font_name_h_ansi = fonts.get(qn('w:hAnsi'))
            data.font_name_cs = fonts.get(qn('w:cs'))

        # ---- Font size ----
        sz = rpr.find(qn('w:sz'))
        if sz is not None:
            val = sz.get(qn('w:val'))
            if val:
                data.size = float(val)

        sz_cs = rpr.find(qn('w:szCs'))
        if sz_cs is not None:
            val = sz_cs.get(qn('w:val'))
            if val:
                data.size_cs = float(val)

        # ---- Bold ----
        b = rpr.find(qn('w:b'))
        if b is not None:
            val = b.get(qn('w:val'))
            data.bold = val is None or val.lower() != 'false'

        b_cs = rpr.find(qn('w:bCs'))
        if b_cs is not None:
            val = b_cs.get(qn('w:val'))
            if data.bold is None:
                data.bold = val is None or val.lower() != 'false'

        # ---- Italic ----
        i_tag = rpr.find(qn('w:i'))
        if i_tag is not None:
            val = i_tag.get(qn('w:val'))
            data.italic = val is None or val.lower() != 'false'

        i_cs = rpr.find(qn('w:iCs'))
        if i_cs is not None:
            val = i_cs.get(qn('w:val'))
            if data.italic is None:
                data.italic = val is None or val.lower() != 'false'

        # ---- Underline ----
        u = rpr.find(qn('w:u'))
        if u is not None:
            u_val = u.get(qn('w:val'))
            if u_val and u_val != 'none':
                data.underline = u_val
            else:
                data.underline = True  # Simple underline

        # ---- Color ----
        color = rpr.find(qn('w:color'))
        if color is not None:
            data.color = color.get(qn('w:val'))
            data.color_theme = color.get(qn('w:themeColor'))
            data.color_tint = color.get(qn('w:tint'))
            data.color_shade = color.get(qn('w:shade'))

        # ---- Strikethrough ----
        strike = rpr.find(qn('w:strike'))
        data.strike = strike is not None

        dstrike = rpr.find(qn('w:dstrike'))
        data.double_strike = dstrike is not None

        # ---- Superscript / Subscript ----
        vertAlign = rpr.find(qn('w:vertAlign'))
        if vertAlign is not None:
            val = vertAlign.get(qn('w:val'))
            if val == 'superscript':
                data.superscript = True
            elif val == 'subscript':
                data.subscript = True

        # ---- Small caps / All caps ----
        small_caps = rpr.find(qn('w:smallCaps'))
        data.small_caps = small_caps is not None

        all_caps = rpr.find(qn('w:caps'))
        data.all_caps = all_caps is not None

        # ---- Highlight ----
        highlight = rpr.find(qn('w:highlight'))
        if highlight is not None:
            data.highlight = highlight.get(qn('w:val'))

        # ---- Language ----
        lang = rpr.find(qn('w:lang'))
        if lang is not None:
            data.lang = lang.get(qn('w:val'))
            data.east_asian_lang = lang.get(qn('w:eastAsia'))

        # ---- Character spacing ----
        spacing = rpr.find(qn('w:spacing'))
        if spacing is not None:
            val = spacing.get(qn('w:val'))
            if val:
                data.spacing = int(val)

        # ---- Position (raised/lowered) ----
        position = rpr.find(qn('w:position'))
        if position is not None:
            val = position.get(qn('w:val'))
            if val:
                data.position = int(val)

        # ---- Other boolean properties ----
        data.vanish = rpr.find(qn('w:vanish')) is not None
        data.outline = rpr.find(qn('w:outline')) is not None
        data.shadow = rpr.find(qn('w:shadow')) is not None
        data.emboss = rpr.find(qn('w:emboss')) is not None
        data.imprint = rpr.find(qn('w:imprint')) is not None
        data.no_proof = rpr.find(qn('w:noProof')) is not None
        data.spec_vanish = rpr.find(qn('w:specVanish')) is not None
        data.web_hidden = rpr.find(qn('w:webHidden')) is not None
        data.rtl = rpr.find(qn('w:rtl')) is not None
        data.complex_script = rpr.find(qn('w:cs')) is not None

        return data

    @staticmethod
    def extract_paragraph_formatting(p_element) -> ParagraphFormatData:
        """Extract paragraph-level formatting from a w:p element.

        Args:
            p_element: lxml element representing a <w:p> paragraph

        Returns:
            ParagraphFormatData with all detected properties
        """
        ppr = p_element.find(qn('w:pPr'))
        if ppr is None:
            return ParagraphFormatData()

        data = ParagraphFormatData()

        # ---- Style ----
        style = ppr.find(qn('w:pStyle'))
        if style is not None:
            data.style_id = style.get(qn('w:val'))
            # Try to derive human-readable name from styles part later
            # For now use the style id
            data.style_name = data.style_id
            # Detect heading level from style name
            if data.style_name:
                upper = data.style_name.lower()
                if upper.startswith('heading') or upper.startswith('head'):
                    try:
                        # Extract number from "Heading1", "heading 1", "Heading 1"
                        num_part = ''.join(c for c in data.style_name if c.isdigit())
                        if num_part:
                            data.heading_level = int(num_part)
                        else:
                            data.heading_level = 1
                    except (ValueError, IndexError):
                        pass

        # ---- Outline level (alternative heading detection) ----
        outline_lvl = ppr.find(qn('w:outlineLvl'))
        if outline_lvl is not None:
            val = outline_lvl.get(qn('w:val'))
            if val:
                lvl = int(val)
                data.outline_level = lvl
                # If no heading style, use outline level as heading level
                if data.heading_level is None:
                    data.heading_level = lvl + 1

        # ---- Alignment ----
        jc = ppr.find(qn('w:jc'))
        if jc is not None:
            data.alignment = jc.get(qn('w:val'))

        # ---- Spacing ----
        spacing = ppr.find(qn('w:spacing'))
        if spacing is not None:
            before = spacing.get(qn('w:before'))
            after = spacing.get(qn('w:after'))
            line = spacing.get(qn('w:line'))
            line_rule = spacing.get(qn('w:lineRule'))

            if before:
                data.space_before = int(before)
            if after:
                data.space_after = int(after)
            if line:
                data.line_spacing = float(line)
            if line_rule:
                data.line_spacing_rule = line_rule

        # ---- Indentation ----
        ind = ppr.find(qn('w:ind'))
        if ind is not None:
            left = ind.get(qn('w:left'))
            right = ind.get(qn('w:right'))
            first_line = ind.get(qn('w:firstLine'))
            hanging = ind.get(qn('w:hanging'))

            if left:
                data.left_indent = int(left)
            if right:
                data.right_indent = int(right)
            if first_line:
                data.first_line_indent = int(first_line)
            if hanging:
                data.hanging_indent = int(hanging)

        # ---- Keep with next / Keep lines / Page break before ----
        data.keep_next = ppr.find(qn('w:keepNext')) is not None
        data.keep_lines = ppr.find(qn('w:keepLines')) is not None
        data.page_break_before = ppr.find(qn('w:pageBreakBefore')) is not None

        # ---- Widow control ----
        widow = ppr.find(qn('w:widowControl'))
        if widow is not None:
            val = widow.get(qn('w:val'))
            data.widow_control = val is None or val.lower() != 'false'

        # ---- Shading (background) ----
        shading = ppr.find(qn('w:shd'))
        if shading is not None:
            data.shading = shading.get(qn('w:fill'))

        # ---- Numbering (list) ----
        numPr = ppr.find(qn('w:numPr'))
        if numPr is not None:
            num_id = numPr.find(qn('w:numId'))
            ilvl = numPr.find(qn('w:ilvl'))
            num_data = {}
            if num_id is not None:
                num_data['numId'] = num_id.get(qn('w:val'))
            if ilvl is not None:
                num_data['ilvl'] = ilvl.get(qn('w:val'))
            if num_data:
                data.numPr = num_data

        # ---- Contextual spacing ----
        data.contextual_spacing = ppr.find(qn('w:contextualSpacing')) is not None

        # ---- Mirror indents ----
        data.mirror_indents = ppr.find(qn('w:mirrorIndents')) is not None

        # ---- Text direction ----
        text_dir = ppr.find(qn('w:textDirection'))
        if text_dir is not None:
            data.text_direction = text_dir.get(qn('w:val'))

        # ---- Suppress line numbers ----
        data.suppress_line_numbers = ppr.find(qn('w:supressLineNumbers')) is not None

        return data

    @classmethod
    def extract_all(cls, document: Document) -> Dict[int, dict]:
        """Extract formatting from ALL paragraphs in a document.

        Iterates via ``document.paragraphs`` (deep search into tables) so
        the returned index is always consistent with
        ``document.paragraphs[i]`` — unlike the old ``body.iterchildren()``
        approach that counted ``<w:tbl>`` as a single index slot and lost
        alignment when tables were present.

        Args:
            document: A python-docx Document object

        Returns:
            Dict mapping **paragraph index** (matching ``document.paragraphs``)
            to:
            {
                'paragraph_format': ParagraphFormatData,
                'runs': [{'text': str, 'formatting': FormattingData}, ...],
                'is_table': bool,   # True when inside a <w:tbl>
            }
        """
        results = {}
        tbl_tag = qn('w:tbl')

        for para_index, para in enumerate(document.paragraphs):
            p_element = para._element

            # Detect whether this paragraph lives inside a table
            is_table = False
            parent = p_element.getparent()
            while parent is not None:
                if parent.tag == tbl_tag:
                    is_table = True
                    break
                parent = parent.getparent()

            # Extract paragraph formatting
            pf = cls.extract_paragraph_formatting(p_element)

            # Extract runs: iterate all w:r under this w:p
            runs = []
            for run_elem in p_element.iter(qn('w:r')):
                rf = cls.extract_runs_formatting(run_elem)

                text_parts = []
                for t in run_elem.iter(qn('w:t')):
                    if t.text:
                        text_parts.append(t.text)

                runs.append({
                    'text': ''.join(text_parts),
                    'formatting': rf,
                })

                # Handle delText for tracked deletions
                del_text_parts = []
                for dt in run_elem.iter(qn('w:delText')):
                    if dt.text:
                        del_text_parts.append(dt.text)
                if del_text_parts:
                    runs.append({
                        'text': ''.join(del_text_parts),
                        'formatting': rf,
                        'is_deletion': True,
                    })

            results[para_index] = {
                'paragraph_format': pf,
                'runs': runs,
                'is_table': is_table,
            }

        return results

    @classmethod
    def extract_document_styles(cls, document: Document) -> Dict[str, str]:
        """Extract style name to style ID mapping from document styles.

        Args:
            document: A python-docx Document object

        Returns:
            Dict mapping style ID to style name
        """
        styles = {}
        try:
            for style in document.styles:
                styles[style.style_id] = style.name
        except Exception:
            pass
        return styles

    @classmethod
    def enhance_paragraph_formats(
        cls,
        formats: Dict[int, dict],
        styles_map: Dict[str, str]
    ) -> None:
        """Enhance paragraph formats with human-readable style names.

        Modifies formats in-place to add style names from the styles map.

        Args:
            formats: Output from extract_all()
            styles_map: Output from extract_document_styles()
        """
        for idx, data in formats.items():
            pf = data['paragraph_format']
            if pf.style_id and pf.style_id in styles_map:
                pf.style_name = styles_map[pf.style_id]
                # Update heading level based on style name
                if pf.heading_level is None:
                    name = pf.style_name.lower()
                    if name.startswith('heading') or name.startswith('head'):
                        try:
                            num = ''.join(c for c in pf.style_id if c.isdigit())
                            if num:
                                pf.heading_level = int(num)
                        except (ValueError, IndexError):
                            pass
