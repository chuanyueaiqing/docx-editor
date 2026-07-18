"""Chapter parser module.

Builds a hierarchical tree of chapters from document headings,
supports lookup by chapter number string/tuple, and content access.
"""
import re
from typing import Dict, List, Optional, Tuple, Union

from .models import ChapterNode, FormatStore, ParagraphData
from .utils import ChapterNotFoundError, qn


class ChapterParser:
    """Build and query a chapter tree from document headings.

    Usage:
        parser = ChapterParser(format_store)
        chapter = parser.get_chapter_by_number('3.1')
        contents = parser.get_chapter_contents(chapter)
    """

    # Heading style patterns for detection (DEPRECATED — see Method 4 note)
    HEADING_PATTERNS: list = []

    CHINESE_NUM_MAP = {
        '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
        '十': 10, '百': 100, '〇': 0,
    }

    # Map Chinese heading unit characters to heading levels
    CHINESE_HEADING_UNIT_LEVEL = {
        '篇': 1,
        '章': 1,
        '部': 1,
        '节': 2,
        '条': 3,
    }

    def __init__(
        self,
        format_store: FormatStore,
        detect_manual_headings: bool = False,
        heading_format_config: Optional[dict] = None,
    ):
        """Initialize chapter parser.

        Args:
            format_store: The format store with document data.
            detect_manual_headings: If True, also detect headings via formatting
                signals (bold + large font + spacing) when no heading style or
                outline level is present.  Default False — only use style+outline.
            heading_format_config: Thresholds for manual heading detection:
                {
                    'min_bold': True,           # require bold
                    'min_size': 28,             # min font size in half-points (14pt)
                    'min_space_before': 100,    # min space-before in twips (5pt)
                    'level1_size': 44,          # 22pt+ → level 1
                    'level2_size': 32,          # 16pt+ → level 2
                    'level3_size': 28,          # 14pt+ → level 3
                }
        """
        self.store = format_store
        self.document = format_store.document
        self.detect_manual_headings = detect_manual_headings
        self.heading_format_config = heading_format_config or {}
        self.root_chapters: List[ChapterNode] = []
        self._chapter_map: Dict[str, ChapterNode] = {}   # "3.1.1" -> node
        self._paragraph_to_chapter: Dict[int, ChapterNode] = {}
        self._all_nodes_cache: Optional[List[ChapterNode]] = None
        self._build_tree()

    # ---- Tree Building ----

    def _populate_paragraphs_data(self):
        """Build ParagraphData list from document paragraphs and format store."""
        from .models import FormattingData
        self.store.paragraphs_data = []
        for i, para in enumerate(self.document.paragraphs):
            fmt_data = self.store.formats_json.get(i, {})

            # Build runs_data from formats_json['runs'] (list of dicts)
            runs_data: List[Tuple[str, FormattingData]] = []
            for run_dict in fmt_data.get('runs', []):
                text = run_dict.get('text', '')
                fmt = run_dict.get('formatting')
                if fmt is not None and not isinstance(fmt, FormattingData):
                    # Convert dict → FormattingData if needed
                    try:
                        fmt = FormattingData(**fmt)
                    except (TypeError, ValueError):
                        fmt = FormattingData()
                fmt = fmt or FormattingData()
                runs_data.append((text, fmt))

            pd = ParagraphData(
                paragraph=para,
                index=i,
                text=para.text,
                formatting=fmt_data.get('paragraph_format'),
                runs_data=runs_data,
            )
            self.store.paragraphs_data.append(pd)

    def _build_tree(self):
        """Walk all paragraphs, identify headings, and build chapter tree."""
        self.root_chapters = []
        self._chapter_map = {}
        self._paragraph_to_chapter = {}
        self._all_nodes_cache = None

        # Ensure paragraphs_data is populated
        if not self.store.paragraphs_data:
            self._populate_paragraphs_data()

        # Collect heading paragraphs and their levels
        headings = self._collect_headings()
        if not headings:
            return

        # Assign number tuples
        self._assign_numbers(headings)

        # Build tree structure
        stack: List[ChapterNode] = []
        for idx, level, text, num_tuple in headings:
            node = ChapterNode(
                heading_text=text,
                heading_level=level,
                number_tuple=num_tuple,
                heading_paragraph_index=idx,
            )

            # Pop stack until we find the parent
            while stack and stack[-1].heading_level >= level:
                stack.pop()

            if stack:
                parent = stack[-1]
                parent.children.append(node)
                node.parent = parent
            else:
                self.root_chapters.append(node)

            stack.append(node)

            # Register in chapter map
            if num_tuple:
                key = '.'.join(str(n) for n in num_tuple)
                self._chapter_map[key] = node

            # Map heading paragraph
            self._paragraph_to_chapter[idx] = node

        # Assign body paragraphs to chapters
        self._assign_body_paragraphs(headings)

        # Assign tables to chapters
        self._assign_tables_to_chapters()

    def _collect_headings(self) -> List[Tuple[int, int, str, Optional[Tuple[int, ...]]]]:
        """Collect all heading paragraphs from the document.

        Returns:
            List of (paragraph_index, heading_level, text, number_tuple)
            where number_tuple is provisional (will be reassigned)
        """
        headings = []

        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            level = self._detect_heading_level(para, i)
            if level is not None and level > 0 and level <= 9:
                headings.append((i, level, text, None))

        return headings

    def _detect_heading_level(self, para, index: int) -> Optional[int]:
        """Detect heading level from a paragraph.

        Uses multiple signals in order of reliability:
        1. Style is a Heading style
        2. Outline level in paragraph properties
        3. Text pattern matching (numbered headings)

        Args:
            para: python-docx Paragraph object
            index: Paragraph index in document

        Returns:
            Heading level (1-9), or None if not a heading
        """
        # Method 1: Check style
        style = para.style
        if style is not None:
            style_name = style.name or ''
            style_id = style.style_id or ''
            for name in [style_name, style_id]:
                name_lower = name.lower()
                if name_lower.startswith('heading'):
                    try:
                        num = ''.join(c for c in name if c.isdigit())
                        if num:
                            return int(num)
                    except (ValueError, IndexError):
                        pass
                    return 1  # Default heading level

        # Skip TOC / Table of Contents entries — they look like headings via
        # outlineLvl in their XML, but are NOT real content headings.
        if style is not None:
            style_name_lower = (style.name or '').lower()
            style_id_lower = (style.style_id or '').lower()
            if style_name_lower.startswith('toc') or style_id_lower.startswith('toc'):
                return None

        # Method 2: Check format store for outline level
        fmt_data = self.store.formats_json.get(index, {})
        pf = fmt_data.get('paragraph_format')
        if pf and pf.heading_level is not None:
            return pf.heading_level
        if pf and pf.outline_level is not None:
            return pf.outline_level + 1

        # Method 3: Check paragraph element properties
        p_element = para._element
        ppr = p_element.find(qn('w:pPr'))
        if ppr is not None:
            outline_lvl = ppr.find(qn('w:outlineLvl'))
            if outline_lvl is not None:
                val = outline_lvl.get(qn('w:val'))
                if val:
                    return int(val) + 1

        # Method 4: Detect manually-formatted headings via formatting signals
        # (bold + large font + spacing).  Only active when the user has opted in
        # via detect_manual_headings=True — disabled by default.
        if not self.detect_manual_headings:
            return None

        cfg = self.heading_format_config
        runs = fmt_data.get('runs', [])
        if runs and isinstance(runs[0], dict):
            run_fmt = runs[0].get('formatting')
            if run_fmt:
                is_bold = bool(getattr(run_fmt, 'bold', None))
                size = getattr(run_fmt, 'size', None) or 0    # half-points

                space_before = 0
                if pf:
                    space_before = pf.space_before or 0

                # Configurable thresholds
                need_bold = cfg.get('min_bold', True)
                min_size = cfg.get('min_size', 28)       # 14pt default
                min_space = cfg.get('min_space_before', 100)  # 5pt default
                l1 = cfg.get('level1_size', 44)          # 22pt
                l2 = cfg.get('level2_size', 32)          # 16pt
                l3 = cfg.get('level3_size', 28)          # 14pt

                if (not need_bold or is_bold) and size >= min_size and space_before >= min_space:
                    if size >= l1:
                        return 1
                    elif size >= l2:
                        return 2
                    elif size >= l3:
                        return 3

        return None

    def _extract_number_from_text(self, text: str, level: int) -> Optional[Tuple[int, ...]]:
        """Extract chapter number from heading text.

        Args:
            text: Heading text (e.g. "3.1.1 系统架构", "第一章 引言")
            level: Detected heading level

        Returns:
            Number tuple or None
        """
        # Pattern 1: "3.1.1" or "3.1.1 " or "3.1.1 xxx"
        m = re.match(r'^(\d+(?:\.\d+)*)', text)
        if m:
            parts = m.group(1).split('.')
            return tuple(int(p) for p in parts)

        # Pattern 2: "第一章" -> (1,), "第二章" -> (2,)
        m = re.match(r'^第([一二三四五六七八九十百零〇]+)[章节篇部]', text)
        if m:
            num = self._chinese_to_int(m.group(1))
            return (num,)

        # Pattern 3: "第1章" -> (1,)
        m = re.match(r'^第(\d+)[章节篇部]', text)
        if m:
            return (int(m.group(1)),)

        # Pattern 4: "A." -> try to infer from outline
        m = re.match(r'^([A-Z])\.\s', text)
        if m:
            return (ord(m.group(1)) - ord('A') + 1,)

        return None

    def _chinese_to_int(self, chinese: str) -> int:
        """Convert Chinese numeral string to integer.

        Args:
            chinese: Chinese numeral (e.g. "一", "十二", "一百零五")

        Returns:
            Integer value
        """
        if not chinese:
            return 0

        # Handle single character
        if chinese in self.CHINESE_NUM_MAP and len(chinese) == 1:
            return self.CHINESE_NUM_MAP[chinese]

        total = 0
        current = 0
        for char in chinese:
            if char in self.CHINESE_NUM_MAP:
                val = self.CHINESE_NUM_MAP[char]
                if val >= 10:
                    if current == 0:
                        current = 1
                    total += current * val
                    current = 0
                else:
                    current = val
        total += current
        return total if total > 0 else current

    def _assign_numbers(self, headings: List[Tuple[int, int, str, Optional[Tuple[int, ...]]]]):
        """Assign consistent number tuples to headings.

        Algorithm:
        - Maintain counters for levels 1-9
        - When a heading at level N is found:
          - Increment counter[N]
          - Reset counters[N+1..9] to 0
          - Assign (counter[1], ..., counter[N])

        Args:
            headings: List of (index, level, text, extracted_num)
        """
        counters = [0] * 10  # Index 1-9 used

        for i, (idx, level, text, extracted_num) in enumerate(headings):
            # Increment at this level
            counters[level] += 1
            # Reset higher levels
            for j in range(level + 1, 10):
                counters[j] = 0

            # Build the number tuple
            num_tuple = tuple(counters[1:level + 1])

            # Replace in list (list is mutable, but tuples are not)
            # We need to update the headings list element
            headings[i] = (idx, level, text, num_tuple)

    def _assign_body_paragraphs(self, headings: List[Tuple[int, int, str, Optional[Tuple[int, ...]]]]):
        """Assign body paragraphs (non-heading) to their containing chapter.

        Each heading's **immediate body** consists of paragraphs that are direct
        children of that heading — i.e. paragraphs after it but NOT inside any
        sub-heading.  Sub-headings' body paragraphs belong to those sub-headings.

        Args:
            headings: List of (index, level, text, number_tuple)
        """
        if not headings:
            return

        heading_indices: set = {h[0] for h in headings}

        for i, (h_idx, h_level, h_text, h_num) in enumerate(headings):
            chapter = self._paragraph_to_chapter.get(h_idx)
            if chapter is None:
                continue

            start = h_idx + 1

            # 遇到任何一个下一个标题（不论层级）就停止，保证只取"直属正文"
            if i + 1 < len(headings):
                end = headings[i + 1][0]
            else:
                end = len(self.document.paragraphs)

            # Assign body paragraphs (skip headings)
            for body_idx in range(start, end):
                if body_idx not in heading_indices:
                    chapter.body_paragraph_indices.append(body_idx)
                    self._paragraph_to_chapter[body_idx] = chapter

    def _assign_tables_to_chapters(self):
        """Map tables to chapters based on their position in the XML body.

        In python-docx, ``document.paragraphs`` and ``document.tables`` are
        two separate flat lists, but they interleave as siblings in the OOXML
        body (``<w:p>`` and ``<w:tbl>``).  This method walks the body children
        in document order, finds which paragraph index each table sits between,
        and assigns it to the corresponding chapter's ``body_table_indices``.
        """
        body = self.document.element.body
        para_counter = 0
        table_counter = 0
        last_chapter: Optional[ChapterNode] = None

        for child in body:
            tag = child.tag
            if tag == qn('w:p'):
                ch = self._paragraph_to_chapter.get(para_counter)
                if ch is not None:
                    last_chapter = ch
                para_counter += 1
            elif tag == qn('w:tbl'):
                if last_chapter is not None:
                    last_chapter.body_table_indices.append(table_counter)
                elif self.root_chapters:
                    # Table before any mapped paragraph → assign to first chapter
                    self.root_chapters[0].body_table_indices.append(table_counter)
                table_counter += 1

    def _iter_all_nodes(self) -> List[ChapterNode]:
        """Iterate all nodes in the tree (depth-first)."""
        result = []

        def dfs(nodes):
            for node in nodes:
                result.append(node)
                if node.children:
                    dfs(node.children)

        dfs(self.root_chapters)
        return result

    # ---- Lookup Methods ----

    def get_chapter_by_number(self, number: Union[str, Tuple[int, ...]]) -> Optional[ChapterNode]:
        """Look up a chapter by its number.

        Args:
            number: Either "3.1.1" string or (3, 1, 1) tuple

        Returns:
            ChapterNode or None if not found
        """
        if isinstance(number, str):
            # Normalize: strip leading/trailing whitespace and dots
            number = number.strip().strip('.')
            key = number
        elif isinstance(number, tuple):
            key = '.'.join(str(n) for n in number)
        else:
            return None

        return self._chapter_map.get(key)

    def get_chapter_contents(self, chapter: ChapterNode) -> List[ParagraphData]:
        """Get all paragraphs (heading + body) belonging to a chapter.

        Args:
            chapter: The chapter node

        Returns:
            List of ParagraphData objects in document order
        """
        contents = []

        # Add heading paragraph
        heading_idx = chapter.heading_paragraph_index
        if heading_idx < len(self.store.paragraphs_data):
            contents.append(self.store.paragraphs_data[heading_idx])

        # Add body paragraphs
        for body_idx in chapter.body_paragraph_indices:
            if body_idx < len(self.store.paragraphs_data):
                contents.append(self.store.paragraphs_data[body_idx])

        # Recursively add children's contents (preserving document order)
        for child in chapter.children:
            contents.extend(self.get_chapter_contents(child))

        return contents

    @staticmethod
    def _collect_table_indices(node: ChapterNode) -> set:
        """Recursively collect table indices from a node and all descendants."""
        indices = set(node.body_table_indices)
        for child in node.children:
            indices.update(ChapterParser._collect_table_indices(child))
        return indices

    def get_chapter_text(self, chapter: ChapterNode) -> str:
        """Get concatenated text of a chapter (heading + body + tables + image markers).

        Tables are rendered in their correct interleaved position relative
        to paragraphs.  Images appear as ``[图片: embed.ext]`` markers.

        Args:
            chapter: The chapter node

        Returns:
            Combined text of all paragraphs, tables, and image markers
        """
        all_table_indices = self._collect_table_indices(chapter)

        if all_table_indices:
            return self._get_text_with_tables(chapter, all_table_indices)

        # ── Simple path: no tables, just paragraphs + images ──
        parts: List[str] = []
        for pd in self.get_chapter_contents(chapter):
            if pd.text:
                parts.append(pd.text)
            markers = self._get_image_markers(pd.paragraph)
            parts.extend(markers)
        return '\n'.join(parts)

    def _get_text_with_tables(self, chapter: ChapterNode,
                              all_table_indices: set) -> str:
        """Render chapter text with tables interleaved at correct positions."""
        body = self.document.element.body
        end_para = self._find_chapter_end(chapter)

        # Pre-render tables
        table_texts: Dict[int, str] = {}
        for ti in all_table_indices:
            if ti < len(self.document.tables):
                table_texts[ti] = self._format_table(self.document.tables[ti])

        # Walk body children, collecting paragraphs and tables in order
        parts: List[str] = []
        para_counter = 0
        table_counter = 0
        started = False

        for child in body:
            tag = child.tag
            if tag == qn('w:p'):
                if para_counter == chapter.heading_paragraph_index:
                    started = True
                if para_counter == end_para:
                    break
                if started and para_counter < len(self.store.paragraphs_data):
                    pd = self.store.paragraphs_data[para_counter]
                    if pd.text:
                        parts.append(pd.text)
                    markers = self._get_image_markers(pd.paragraph)
                    parts.extend(markers)
                para_counter += 1
            elif tag == qn('w:tbl'):
                if started and table_counter in table_texts:
                    parts.append(table_texts[table_counter])
                table_counter += 1

        return '\n'.join(parts)

    def _find_chapter_end(self, chapter: ChapterNode) -> int:
        """Find the body-paragraph index where this chapter's content ends.

        The end is the paragraph index of the next heading at the same or
        higher (numerically smaller) level, or the end of the document.
        """
        end = len(self.document.paragraphs)
        for node in self._iter_all_nodes():
            if (node.heading_paragraph_index > chapter.heading_paragraph_index
                    and node.heading_level <= chapter.heading_level):
                end = node.heading_paragraph_index
                break
        return end

    @staticmethod
    def _format_table(table) -> str:
        """Render a python-docx Table as readable text.

        Returns a pipe-delimited multi-line string similar to Markdown.
        Empty rows at the end are trimmed.
        """
        rows = []
        ncols = 0
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            ncols = max(ncols, len(cells))
            if any(c for c in cells):
                rows.append(cells)

        if not rows:
            return ''

        # Determine column widths
        col_widths = [0] * ncols
        for row in rows:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(cell))

        lines = []
        for ri, row in enumerate(rows):
            # Pad row to ncols
            padded = list(row) + [''] * (ncols - len(row))
            line = ' | '.join(cell.ljust(col_widths[i])
                              for i, cell in enumerate(padded))
            lines.append(line)
            if ri == 0:
                # Separator row
                sep = '-|-'.join('-' * col_widths[i] for i in range(ncols))
                lines.append(sep)

        return '\n'.join(lines)

    # ---- Image Detection & Extraction ----

    # OOXML namespaces for drawing/blip detection
    _BLIP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    _REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    _IMAGE_EXT_MAP = {
        'image/png': '.png',
        'image/jpeg': '.jpg',
        'image/gif': '.gif',
        'image/bmp': '.bmp',
        'image/svg+xml': '.svg',
        'image/tiff': '.tiff',
        'image/x-emf': '.emf',
        'image/x-wmf': '.wmf',
    }

    @classmethod
    def _get_image_markers(cls, para) -> List[str]:
        """Detect embedded images in a paragraph and return marker strings.

        Returns list like ``["[图片: rId9.png]"]`` — one per embedded image.
        Works even if relationship data is not available (returns generic marker).
        """
        markers: List[str] = []
        if para is None:
            return markers

        try:
            blips = para._element.findall(
                f'.//{{{cls._BLIP_NS}}}blip'
            )
        except Exception:
            return markers

        for blip in blips:
            embed = blip.get(f'{{{cls._REL_NS}}}embed')
            if embed:
                # Try to get the actual extension from relationships
                ext = '.png'  # default
                try:
                    from docx import Document as _Doc
                    # Access part through the element's document
                    doc_element = para._element.getroottree()
                    # We can't easily get rels from here, use filename-based approach
                    if hasattr(para, 'part') and hasattr(para.part, 'rels'):
                        rel = para.part.rels.get(embed)
                        if rel and hasattr(rel, 'target_part'):
                            ct = getattr(rel.target_part, 'content_type', '')
                            ext = cls._IMAGE_EXT_MAP.get(ct, '.bin')
                except Exception:
                    pass
                markers.append(f'[图片: {embed}{ext}]')
            else:
                markers.append('[图片]')

        return markers

    def get_chapter_images(self, chapter: ChapterNode,
                           output_dir: str) -> Dict[str, str]:
        """Extract all images in a chapter to files on disk.

        Args:
            chapter: The chapter node
            output_dir: Directory to save extracted image files

        Returns:
            Dict mapping embed ID (e.g. ``"rId9"``) to saved file path
        """
        import os
        os.makedirs(output_dir, exist_ok=True)

        end_para = self._find_chapter_end(chapter)
        start_para = chapter.heading_paragraph_index
        saved: Dict[str, str] = {}

        for para_idx in range(start_para, end_para):
            if para_idx >= len(self.store.paragraphs_data):
                break
            pd = self.store.paragraphs_data[para_idx]
            para = pd.paragraph
            if para is None:
                continue

            try:
                blips = para._element.findall(
                    f'.//{{{self._BLIP_NS}}}blip'
                )
            except Exception:
                continue

            for blip in blips:
                embed = blip.get(f'{{{self._REL_NS}}}embed')
                if not embed:
                    continue
                if embed in saved:
                    continue  # already extracted (same image may repeat)

                rel = self.document.part.rels.get(embed)
                if rel is None or 'image' not in (rel.reltype or ''):
                    continue
                try:
                    image_bytes = rel.target_part.blob
                except Exception:
                    continue

                ct = getattr(rel.target_part, 'content_type', '')
                ext = self._IMAGE_EXT_MAP.get(ct, '.bin')
                filename = f'{embed}{ext}'
                save_path = os.path.join(output_dir, filename)
                with open(save_path, 'wb') as f:
                    f.write(image_bytes)
                saved[embed] = save_path

        return saved

    # ---- Public Utilities ----

    def get_chapter_for_paragraph(
        self, para_index: int
    ) -> Optional[ChapterNode]:
        """Get the chapter that contains a given paragraph index.

        Args:
            para_index: Index into ``document.paragraphs``

        Returns:
            ChapterNode that contains this paragraph, or None if not found
        """
        return self._paragraph_to_chapter.get(para_index)

    def list_chapters(self) -> List[ChapterNode]:
        """List all chapters in depth-first order.

        Returns:
            Flat list of all ChapterNode objects in document order
        """
        return self._iter_all_nodes()

    def generate_toc_entries(
        self, max_level: int = 9
    ) -> List[Tuple[str, str, int]]:
        """Generate table-of-contents entries from the chapter tree.

        Returns entries in document (depth-first) order, filtered by
        ``max_level``.  Each entry is ``(number_string, heading_text, level)``.

        Args:
            max_level: Maximum heading level to include (default 9 = all)

        Returns:
            List of ``(number_string, heading_text, level)`` tuples.
            ``number_string`` is something like ``"3.1.1"``, or ``""`` when
            the heading has no number tuple.
        """
        entries: List[Tuple[str, str, int]] = []
        for node in self._iter_all_nodes():
            if node.heading_level > max_level:
                continue
            num = node.to_string()
            # to_string() falls back to heading_text when no number_tuple
            # — in that case use empty string for the number
            if node.number_tuple is None:
                num = ''
            entries.append((num, node.heading_text, node.heading_level))
        return entries

    def tree_to_string(self, indent: int = 2) -> str:
        """Return a visual string representation of the chapter tree.

        Args:
            indent: Spaces per indent level (default 2)

        Returns:
            Multi-line string showing the tree hierarchy
        """
        lines: List[str] = []

        def _walk(nodes: List[ChapterNode], depth: int):
            for node in nodes:
                num = node.to_string()
                title = node.heading_text
                prefix = ' ' * (depth * indent)
                body_count = len(node.body_paragraph_indices)
                children_count = len(node.children)
                info = f'({body_count} body para'
                if children_count:
                    info += f', {children_count} children'
                info += ')'
                lines.append(f'{prefix}{num} {title} {info}')
                if node.children:
                    _walk(node.children, depth + 1)

        _walk(self.root_chapters, 0)
        return '\n'.join(lines)

    # ---- Deletion ----

    def delete_chapter(self, chapter: ChapterNode):
        """Delete a chapter and all its contents from the document.

        This removes:
        - The heading paragraph
        - All body paragraphs
        - All sub-chapters (children) and their contents

        After deletion, the tree is rebuilt.
        """
        # Collect all indices to delete (this chapter + all descendants)
        indices_to_delete = set()
        self._collect_indices(chapter, indices_to_delete)

        # Delete from highest index to lowest (to preserve indices)
        for idx in sorted(indices_to_delete, reverse=True):
            if idx < len(self.document.paragraphs):
                para = self.document.paragraphs[idx]
                elem = para._element
                elem.getparent().remove(elem)

        # Rebuild tree
        self._rebuild()

    def _collect_indices(self, chapter: ChapterNode, indices: set):
        """Recursively collect paragraph indices for a chapter and all descendants.

        Args:
            chapter: Chapter to collect from
            indices: Set to add indices into
        """
        indices.add(chapter.heading_paragraph_index)
        for body_idx in chapter.body_paragraph_indices:
            indices.add(body_idx)
        for child in chapter.children:
            self._collect_indices(child, indices)

    def _rebuild(self):
        """Rebuild the chapter tree and paragraph data after structural changes."""
        from .format_extractor import FormatExtractor
        self.store.formats_json = FormatExtractor.extract_all(self.document)
        self._populate_paragraphs_data()
        self._build_tree()

    # ---- Utility ----

    def tree_to_string(self, indent: int = 2) -> str:
        """Return a visual representation of the chapter tree.

        Args:
            indent: Number of spaces per indent level

        Returns:
            Formatted tree string
        """
        lines = []

        def dfs(nodes, depth: int):
            for node in nodes:
                prefix = ' ' * (depth * indent)
                num_str = node.to_string()
                lines.append(f"{prefix}{num_str} {node.heading_text}")
                if node.children:
                    dfs(node.children, depth + 1)

        dfs(self.root_chapters, 0)
        return '\n'.join(lines)

    def list_chapters(self) -> List[ChapterNode]:
        """Get flat list of all chapters in document order.

        Returns:
            List of ChapterNode in document order (depth-first)
        """
        if self._all_nodes_cache is None:
            self._all_nodes_cache = self._iter_all_nodes()
        return self._all_nodes_cache
